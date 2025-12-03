from fastapi import FastAPI, HTTPException, Depends, status, UploadFile, File, Request
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from datetime import datetime, timedelta
from typing import Optional, List
from contextlib import asynccontextmanager
import hashlib
import os
import re
from dotenv import load_dotenv
import pymongo      
from bson import ObjectId
from pydantic import BaseModel, Field, EmailStr, field_validator
from openai import OpenAI
import io
import mimetypes
# import jwt  # Temporarily disabled due to library conflicts
import base64
import hmac
import hashlib
import json
import re
from typing import List, Dict, Optional
from google.oauth2 import id_token
from google.auth.transport import requests
import time
from collections import defaultdict, deque
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import secrets
import string
# Try to import PDF processing libraries
PDF_AVAILABLE = False
PDF_LIBRARY = None

try:
    import pypdf
    PDF_AVAILABLE = True
    PDF_LIBRARY = "pypdf"
    print("pypdf available for PDF processing")
except ImportError:
    try:
        import PyPDF2
        PDF_AVAILABLE = True
        PDF_LIBRARY = "PyPDF2"
        print("PyPDF2 available for PDF processing")
    except ImportError:
        PDF_AVAILABLE = False
        PDF_LIBRARY = None
        print("No PDF processing library available - PDF processing will use fallback content")


load_dotenv()

class PDFContentExtractor:
    """Enhanced PDF content extractor for medical cases and MCQs"""
    
    def __init__(self):
        self.case_keywords = [
            "case", "patient", "history", "presentation", 
            "chief complaint", "diagnosis", "treatment"
        ]
        self.mcq_keywords = [
            "question", "which of the following", "what is",
            "a)", "b)", "c)", "d)", "e)", "answer:"
        ]
    
    def extract_text_from_pdf(self, pdf_content: bytes) -> Dict[str, any]:
        """Extract text and structure from PDF"""
        try:
            pdf_reader = pypdf.PdfReader(io.BytesIO(pdf_content))
            
            result = {
                "full_text": "",
                "pages": [],
                "total_pages": len(pdf_reader.pages),
                "detected_cases": [],
                "detected_mcqs": [],
                "has_cases": False,
                "has_mcqs": False
            }
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                
                if page_text.strip():
                    result["full_text"] += f"\n--- Page {page_num + 1} ---\n"
                    result["full_text"] += page_text
                    
                    result["pages"].append({
                        "page_number": page_num + 1,
                        "text": page_text,
                        "char_count": len(page_text)
                    })
            
            # Detect cases and MCQs
            result["detected_cases"] = self.detect_cases(result["full_text"])
            result["detected_mcqs"] = self.detect_mcqs(result["full_text"])
            
            result["has_cases"] = len(result["detected_cases"]) > 0
            result["has_mcqs"] = len(result["detected_mcqs"]) > 0
            
            return result
            
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    def detect_cases(self, text: str) -> List[Dict[str, str]]:
        """Detect medical cases in text"""
        cases = []
        
        # Pattern 1: "Case X:" or "CASE X:"
        case_pattern = r'(?:CASE|Case)\s+(\d+|[IVX]+)[:\.]?\s*\n?(.*?)(?=(?:CASE|Case)\s+\d+|$)'
        matches = re.finditer(case_pattern, text, re.IGNORECASE | re.DOTALL)
        
        for match in matches:
            case_num = match.group(1)
            case_text = match.group(2).strip()
            
            if len(case_text) > 50:  # Minimum length for a valid case
                cases.append({
                    "case_number": case_num,
                    "content": case_text[:1000],  # First 1000 chars
                    "full_content": case_text,
                    "type": "numbered_case"
                })
        
        # Pattern 2: Patient presentation blocks
        patient_pattern = r'(?:Patient|PATIENT)[:\s]+(.*?)(?=(?:Patient|PATIENT|Question|QUESTION)|$)'
        patient_matches = re.finditer(patient_pattern, text, re.IGNORECASE | re.DOTALL)
        
        for match in patient_matches:
            patient_text = match.group(1).strip()
            
            if len(patient_text) > 100 and self._contains_case_elements(patient_text):
                cases.append({
                    "case_number": f"P{len(cases) + 1}",
                    "content": patient_text[:1000],
                    "full_content": patient_text,
                    "type": "patient_presentation"
                })
        
        return cases
    
    def detect_mcqs(self, text: str) -> List[Dict[str, any]]:
        """Detect MCQ questions in text"""
        mcqs = []
        
        # Pattern: Question followed by options A) B) C) D) E)
        question_pattern = r'(?:Question\s+(\d+)|(\d+)[\.)]\s*)([^\n]+\?)\s*\n((?:[A-E][\.)]\s*[^\n]+\s*\n?)+)'
        
        matches = re.finditer(question_pattern, text, re.IGNORECASE | re.MULTILINE)
        
        for match in matches:
            q_num = match.group(1) or match.group(2)
            question_text = match.group(3).strip()
            options_text = match.group(4).strip()
            
            # Extract options
            options = self._extract_options(options_text)
            
            if len(options) >= 4:  # Valid MCQ should have at least 4 options
                mcqs.append({
                    "question_number": q_num or str(len(mcqs) + 1),
                    "question": question_text,
                    "options": options,
                    "raw_text": match.group(0)
                })
        
        # Alternative pattern: Questions without explicit numbering
        if len(mcqs) == 0:
            # Look for question marks followed by options
            alt_pattern = r'([^\n]+\?)\s*\n((?:[A-E][\.)]\s*[^\n]+\s*\n?){4,})'
            alt_matches = re.finditer(alt_pattern, text, re.MULTILINE)
            
            for match in alt_matches:
                question_text = match.group(1).strip()
                options_text = match.group(2).strip()
                
                options = self._extract_options(options_text)
                
                if len(options) >= 4:
                    mcqs.append({
                        "question_number": str(len(mcqs) + 1),
                        "question": question_text,
                        "options": options,
                        "raw_text": match.group(0)
                    })
        
        return mcqs
    
    def _extract_options(self, options_text: str) -> List[Dict[str, str]]:
        """Extract individual options from options text"""
        options = []
        
        # Pattern for options: A) or A. followed by text
        option_pattern = r'([A-E])[\.)]\s*([^\n]+)'
        matches = re.finditer(option_pattern, options_text, re.IGNORECASE)
        
        for match in matches:
            option_id = match.group(1).upper()
            option_text = match.group(2).strip()
            
            options.append({
                "id": option_id,
                "text": option_text
            })
        
        return options
    
    def _contains_case_elements(self, text: str) -> bool:
        """Check if text contains typical case elements"""
        text_lower = text.lower()
        
        case_indicators = [
            "age", "year", "old", "male", "female",
            "complain", "present", "history",
            "symptom", "sign", "examination",
            "diagnosis", "treatment"
        ]
        
        matches = sum(1 for indicator in case_indicators if indicator in text_lower)
        return matches >= 3

MONGODB_URL = os.getenv("MONGODB_URL")
DATABASE_NAME = os.getenv("DATABASE_NAME")
SECRET_KEY = os.getenv("SECRET_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID")
ALGORITHM = "HS256"

# Configure OpenAI AI
openai_client = None
if OPENAI_API_KEY:
    openai_client = OpenAI(api_key=OPENAI_API_KEY)

# Simple password hashing with SHA-256
security = HTTPBearer()


client = None
db = None

def connect_to_mongodb():
    global client, db
    try:
        client = pymongo.MongoClient(MONGODB_URL)
        db = client[DATABASE_NAME]
        
        client.server_info()
        print(f"Connected to MongoDB: {DATABASE_NAME}")
        
        db.users.create_index("email", unique=True)
        db.users.create_index("username", unique=True)
        db.documents.create_index("uploaded_by")
        db.documents.create_index("uploaded_at")
        db.chats.create_index("user_id")
        db.chats.create_index("updated_at")
        db.chat_messages.create_index("chat_id")
        db.chat_messages.create_index("timestamp")
        db.generated_cases.create_index("chat_id")
        db.generated_cases.create_index("document_id")
        db.generated_mcqs.create_index("chat_id")
        db.generated_mcqs.create_index("document_id")
        db.generated_concepts.create_index("chat_id")
        db.generated_concepts.create_index("document_id")
        
        # Create admin user if it doesn't exist
        create_admin_user()
        
    except Exception as e:
        print(f"Failed to connect to MongoDB: {e}")
        raise

def create_admin_user():
    """Create a default admin user if it doesn't exist"""
    try:
        if db is None:
            return
            
        # Check if admin user already exists
        admin_user = db.users.find_one({"email": "admin@casewise.com"})
        if admin_user:
            print("✅ Admin user already exists")
            return
            
        # Create admin user
        admin_password = "admin123"
        hashed_password = hash_password(admin_password)
        
        admin_user_data = {
            "email": "admin@casewise.com",
            "username": "admin",
            "full_name": "System Administrator",
            "hashed_password": hashed_password,
            "role": "admin",
            "is_active": True,
            "status": "active",
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        result = db.users.insert_one(admin_user_data)
        print(f"✅ Admin user created with ID: {result.inserted_id}")
        print("📧 Admin credentials: admin@casewise.com / admin123")
        
    except Exception as e:
        print(f"❌ Failed to create admin user: {e}")

def get_database():
    return db

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    connect_to_mongodb()
    yield
    # Shutdown
    if client:
        client.close()
        print("MongoDB connection closed")

class UserSignup(BaseModel):
    """Schema for user signup"""
    email: EmailStr
    username: str = Field(..., min_length=3, max_length=50)
    password: str = Field(..., min_length=6)
    full_name: Optional[str] = Field(None, max_length=100)
    
    @field_validator('password')
    @classmethod
    def validate_password_length(cls, v):
        # Allow longer passwords - we'll truncate them in hash_password function
        if len(v) < 6:
            raise ValueError('Password must be at least 6 characters long')
        return v

class UserLogin(BaseModel):
    """Schema for user login"""
    email: EmailStr
    password: str
    remember_me: Optional[bool] = False

class GoogleAuthRequest(BaseModel):
    """Schema for Google OAuth authentication"""
    id_token: str

class Token(BaseModel):
    """Token response schema"""
    access_token: str
    token_type: str = "bearer"
    user_id: str
    email: str

class UserResponse(BaseModel):
    """User response schema"""
    id: str
    email: str
    username: str
    full_name: Optional[str] = None
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    profile_image_url: Optional[str] = None
    bio: Optional[str] = None
    role: Optional[str] = "user"
    created_at: datetime
    updated_at: Optional[datetime] = None

class ProfileUpdateRequest(BaseModel):
    """Profile update request schema"""
    first_name: Optional[str] = Field(None, max_length=50)
    last_name: Optional[str] = Field(None, max_length=50)
    full_name: Optional[str] = Field(None, max_length=100)
    bio: Optional[str] = Field(None, max_length=500)
    email: Optional[EmailStr] = None
    username: Optional[str] = Field(None, min_length=3, max_length=50)

class PasswordChangeRequest(BaseModel):
    """Password change request schema"""
    current_password: str
    new_password: str = Field(..., min_length=6)

class AdminUserUpdateRequest(BaseModel):
    """Admin user update request schema"""
    user_id: str
    role: Optional[str] = Field(None, pattern="^(user|admin)$")
    status: Optional[str] = Field(None, pattern="^(active|inactive|suspended)$")

class AdminBulkUserUpdateRequest(BaseModel):
    """Admin bulk user update request schema"""
    updates: List[AdminUserUpdateRequest]

class ForgotPasswordRequest(BaseModel):
    """Forgot password request schema"""
    email: EmailStr

class ResetPasswordRequest(BaseModel):
    """Reset password request schema"""
    token: str
    new_password: str = Field(..., min_length=8)

class Notification(BaseModel):
    """Notification schema"""
    id: str
    user_id: str
    type: str  # "file_upload", "profile_update", "system", etc.
    title: str
    message: str
    is_read: bool = False
    created_at: datetime
    metadata: Optional[dict] = None

class NotificationResponse(BaseModel):
    """Notification response schema"""
    id: str
    type: str
    title: str
    message: str
    is_read: bool
    created_at: datetime
    metadata: Optional[dict] = None

class DocumentResponse(BaseModel):
    """Document response schema"""
    id: str
    filename: str
    content_type: str
    size: int
    uploaded_by: str
    uploaded_at: datetime
    content: Optional[str] = None

class UserAnalytics(BaseModel):
    """User analytics response schema"""
    name: str
    timeSpent: str
    casesUploaded: int
    mcqAttempted: int
    mostQuestionsType: str
    totalQuestionsCorrect: int
    totalQuestionsAttempted: int
    averageScore: int
    lastActiveDate: str

class CaseGenerationRequest(BaseModel):
    """Case generation request schema"""
    document_id: str
    prompt: str = Field(..., min_length=10, max_length=1000)
    num_scenarios: int = Field(default=3, ge=1, le=10)

class CaseScenario(BaseModel):
    """Case scenario schema"""
    title: str
    description: str
    key_points: List[str]
    difficulty: str

class CaseGenerationResponse(BaseModel):
    """Case generation response schema"""
    document_id: str
    prompt: str
    scenarios: List[CaseScenario]
    generated_at: datetime

class CaseTitleRequest(BaseModel):
    """Case title generation request schema"""
    document_id: str
    num_cases: int = Field(default=5, ge=5, le=10)

class CaseTitle(BaseModel):
    """Case title schema"""
    id: str
    title: str
    description: str
    difficulty: str

class CaseTitlesResponse(BaseModel):
    """Case titles response schema"""
    document_id: str
    cases: List[CaseTitle]
    generated_at: datetime

class MCQRequest(BaseModel):
    """MCQ generation request schema"""
    document_id: Optional[str] = None
    case_id: Optional[str] = None
    case_title: Optional[str] = None
    num_questions: int = Field(default=3, ge=1, le=10)
    include_hints: bool = Field(default=True)
    difficulty: Optional[str] = Field(default=None, description="Difficulty level: Easy, Moderate, or Hard")

class MCQOption(BaseModel):
    """MCQ option schema"""
    id: str
    text: str
    is_correct: bool

class MCQQuestion(BaseModel):
    """MCQ question schema"""
    id: str
    question: str
    options: List[MCQOption]
    explanation: str
    difficulty: str
    hint: Optional[str] = None

class MCQResponse(BaseModel):
    """MCQ response schema"""
    questions: List[MCQQuestion]
    generated_at: datetime

class ConceptRequest(BaseModel):
    """Concept identification request schema"""
    document_id: str
    num_concepts: int = Field(default=1, ge=1, le=1)  # Always generate 1 case breakdown
    case_title: Optional[str] = None  # Specific case to generate concepts for

class Concept(BaseModel):
    """Concept schema"""
    id: str
    title: str
    description: str
    importance: str
    difficulty: Optional[str] = None
    # Old format structured fields for case breakdown
    objective: Optional[str] = None
    patient_profile: Optional[str] = None
    history_of_present_illness: Optional[str] = None
    past_medical_history: Optional[str] = None
    medications: Optional[str] = None
    examination: Optional[str] = None
    initial_investigations: Optional[str] = None
    case_progression: Optional[str] = None
    final_diagnosis: Optional[str] = None
    # New format structured fields for key concepts
    case_id: Optional[str] = None
    key_concept: Optional[str] = None
    key_concept_summary: Optional[str] = None
    learning_objectives: Optional[List[str]] = None
    core_pathophysiology: Optional[str] = None
    clinical_reasoning_steps: Optional[List[str]] = None
    red_flags_and_pitfalls: Optional[List[str]] = None
    differential_diagnosis_framework: Optional[List[str]] = None
    important_labs_imaging_to_know: Optional[List[str]] = None
    why_this_case_matters: Optional[str] = None

class ConceptResponse(BaseModel):
    """Concept response schema"""
    document_id: str
    concepts: List[Concept]
    generated_at: datetime

class AutoGenerationRequest(BaseModel):
    """Auto generation request schema"""
    document_id: str
    generate_cases: bool = True
    generate_mcqs: bool = True
    generate_concepts: bool = True
    generate_titles: bool = True
    num_cases: int = Field(default=5, ge=5, le=20)
    num_mcqs: int = Field(default=5, ge=1, le=15)
    num_concepts: int = Field(default=5, ge=1, le=15)
    num_titles: int = Field(default=5, ge=1, le=15)

class AutoGenerationResponse(BaseModel):
    """Auto generation response schema"""
    document_id: str
    cases: Optional[List[CaseScenario]] = None
    mcqs: Optional[List[MCQQuestion]] = None
    concepts: Optional[List[Concept]] = None
    titles: Optional[List[CaseTitle]] = None
    generated_at: datetime
    success: bool
    message: str


def hash_password(password: str) -> str:
    """Hash a password using SHA-256"""
    print(f"HASH_PASSWORD: Hashing password of length: {len(password)}")
    # Add a salt to make it more secure
    salt = "casewise_salt_2024"
    salted_password = password + salt
    hashed = hashlib.sha256(salted_password.encode('utf-8')).hexdigest()
    print(f"HASH_PASSWORD: Hash successful")
    return hashed

def verify_password(plain_password: str, hashed_password: str) -> bool:
    """Verify a password"""
    # Hash the plain password and compare
    salt = "casewise_salt_2024"
    salted_password = plain_password + salt
    computed_hash = hashlib.sha256(salted_password.encode('utf-8')).hexdigest()
    is_valid = computed_hash == hashed_password
    return is_valid

def verify_google_token(id_token_str: str) -> dict:
    """Verify Google ID token and return user info"""
    try:
        print(f"GOOGLE_AUTH: Verifying Google ID token")
        
        if not GOOGLE_CLIENT_ID:
            print("GOOGLE_AUTH: ERROR - GOOGLE_CLIENT_ID not configured")
            raise HTTPException(status_code=500, detail="Google OAuth not configured")
        
        print(f"GOOGLE_AUTH: Using client ID: {GOOGLE_CLIENT_ID[:20]}...")
        
        # Check system clock
        import time
        current_time = int(time.time())
        print(f"GOOGLE_AUTH: Current system time: {current_time}")
        
        # Verify the token with clock tolerance
        idinfo = id_token.verify_oauth2_token(
            id_token_str, 
            requests.Request(), 
            GOOGLE_CLIENT_ID,
            clock_skew_in_seconds=120  # Allow 2 minutes of clock difference
        )
        
        # Verify the issuer
        if idinfo['iss'] not in ['accounts.google.com', 'https://accounts.google.com']:
            print(f"GOOGLE_AUTH: Wrong issuer: {idinfo['iss']}")
            raise ValueError('Wrong issuer.')
        
        print(f"GOOGLE_AUTH: Token verified for user: {idinfo.get('email')}")
        return idinfo
        
    except ValueError as e:
        print(f"GOOGLE_AUTH: Invalid token: {e}")
        raise HTTPException(status_code=401, detail=f"Invalid Google token: {str(e)}")
    except Exception as e:
        print(f"GOOGLE_AUTH: Error verifying token: {e}")
        print(f"GOOGLE_AUTH: Error type: {type(e).__name__}")
        raise HTTPException(status_code=401, detail=f"Failed to verify Google token: {str(e)}")

def create_notification(user_id: str, notification_type: str, title: str, message: str, metadata: dict = None):
    """Create a notification for a user"""
    try:
        notification_doc = {
            "user_id": user_id,
            "type": notification_type,
            "title": title,
            "message": message,
            "is_read": False,
            "created_at": datetime.utcnow().isoformat() + "Z",
            "metadata": metadata or {}
        }
        
        result = db.notifications.insert_one(notification_doc)
        print(f" Notification created: {notification_type} for user {user_id}")
        return str(result.inserted_id)
    except Exception as e:
        print(f" Error creating notification: {e}")
        return None

def send_reset_password_email(email: str, reset_token: str, user_name: str = None):
    """Send password reset email to user"""
    try:
        # Email configuration - you should set these in your environment variables
        smtp_server = os.getenv("SMTP_SERVER", "smtp.gmail.com")
        smtp_port = int(os.getenv("SMTP_PORT", "587"))
        smtp_username = os.getenv("SMTP_USERNAME", "")
        smtp_password = os.getenv("SMTP_PASSWORD", "")
        
        if not smtp_username or not smtp_password:
            print("❌ SMTP credentials not configured. Please set SMTP_USERNAME and SMTP_PASSWORD environment variables.")
            return False
        
        # Create message
        msg = MIMEMultipart()
        msg['From'] = smtp_username
        msg['To'] = email
        msg['Subject'] = "Password Reset Request - CaseWise"
        
        # Create reset link
        frontend_url = os.getenv("FRONTEND_URL", "https://casewise-beta.vercel.app")
        reset_link = f"{frontend_url}/reset-password?token={reset_token}"
        
        # Email body
        body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <h2 style="color: #2563eb;">Password Reset Request</h2>
                
                <p>Hello {user_name or 'User'},</p>
                
                <p>We received a request to reset your password for your CaseWise account. If you made this request, click the button below to reset your password:</p>
                
                <div style="text-align: center; margin: 30px 0;">
                    <a href="{reset_link}" 
                       style="background-color: #2563eb; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px; display: inline-block; font-weight: bold;">
                        Reset Password
                    </a>
                </div>
                
                <p>Or copy and paste this link into your browser:</p>
                <p style="word-break: break-all; background-color: #f3f4f6; padding: 10px; border-radius: 4px;">
                    {reset_link}
                </p>
                
                <p><strong>This link will expire in 1 hour for security reasons.</strong></p>
                
                <p>If you didn't request a password reset, please ignore this email. Your password will remain unchanged.</p>
                
                <hr style="border: none; border-top: 1px solid #e5e7eb; margin: 30px 0;">
                <p style="font-size: 12px; color: #6b7280;">
                    This email was sent from CaseWise. If you have any questions, please contact our support team.
                </p>
            </div>
        </body>
        </html>
        """
        
        msg.attach(MIMEText(body, 'html'))
        
        # Send email
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(smtp_username, smtp_password)
        text = msg.as_string()
        server.sendmail(smtp_username, email, text)
        server.quit()
        
        print(f"✅ Password reset email sent to: {email}")
        return True
        
    except Exception as e:
        print(f"❌ Error sending password reset email: {e}")
        return False

def generate_reset_token():
    """Generate a secure reset token"""
    return ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(32))

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    """Create JWT access token using simple implementation"""
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(hours=24)
    to_encode.update({"exp": expire.timestamp()})
    
    # Simple JWT implementation
    header = {"alg": "HS256", "typ": "JWT"}
    payload = to_encode
    
    # Encode header and payload
    header_encoded = base64.urlsafe_b64encode(json.dumps(header).encode()).decode().rstrip('=')
    payload_encoded = base64.urlsafe_b64encode(json.dumps(payload).encode()).decode().rstrip('=')
    
    # Create signature
    message = f"{header_encoded}.{payload_encoded}"
    signature = hmac.new(
        SECRET_KEY.encode(),
        message.encode(),
        hashlib.sha256
    ).digest()
    signature_encoded = base64.urlsafe_b64encode(signature).decode().rstrip('=')
    
    return f"{header_encoded}.{payload_encoded}.{signature_encoded}"

def decode_jwt_token(token: str) -> dict:
    """Decode JWT token using simple implementation"""
    try:
        parts = token.split('.')
        if len(parts) != 3:
            raise ValueError("Invalid token format")
        
        header_encoded, payload_encoded, signature_encoded = parts
        
        # Decode payload
        payload_padded = payload_encoded + '=' * (4 - len(payload_encoded) % 4)
        payload = json.loads(base64.urlsafe_b64decode(payload_padded))
        
        # Verify signature
        message = f"{header_encoded}.{payload_encoded}"
        expected_signature = hmac.new(
            SECRET_KEY.encode(),
            message.encode(),
            hashlib.sha256
        ).digest()
        expected_signature_encoded = base64.urlsafe_b64encode(expected_signature).decode().rstrip('=')
        
        if signature_encoded != expected_signature_encoded:
            raise ValueError("Invalid signature")
        
        # Check expiration
        if 'exp' in payload:
            if datetime.utcnow().timestamp() > payload['exp']:
                raise ValueError("Token expired")
        
        return payload
    except Exception as e:
        raise ValueError(f"Token decode error: {str(e)}")

def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Get current user from JWT token"""
    try:
        payload = decode_jwt_token(credentials.credentials)
        user_id: str = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
    except ValueError as e:
        raise HTTPException(status_code=401, detail="Invalid token")
    
    user = db.users.find_one({"_id": ObjectId(user_id)})
    if user is None:
        raise HTTPException(status_code=401, detail="User not found")
    
    user["id"] = str(user["_id"])
    return user

# Rate Limiter Class
class RateLimiter:
    def __init__(self, max_requests: int = 10, time_window: int = 1):
        self.max_requests = max_requests
        self.time_window = time_window
        self.requests = defaultdict(deque)
    
    def is_allowed(self, client_ip: str) -> bool:
        now = time.time()
        client_requests = self.requests[client_ip]
        
        # Remove old requests outside the time window
        while client_requests and client_requests[0] <= now - self.time_window:
            client_requests.popleft()
        
        # Check if under the limit
        if len(client_requests) < self.max_requests:
            client_requests.append(now)
            return True
        
        return False

# Initialize rate limiter
rate_limiter = RateLimiter(max_requests=10, time_window=1)  # 10 requests per second

def get_rate_limiter():
    return rate_limiter

app = FastAPI(
    title="Medical AI - Auth API",
    description="User authentication using PyMongo",
    version="1.0.0",
    lifespan=lifespan
)

# CORS middleware - backup to custom middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://casewise-beta.vercel.app",  # Production frontend
        "http://localhost:3000",
        "http://localhost:3001", 
        "http://127.0.0.1:3000",
        "http://127.0.0.1:3001",
        "http://localhost:8000",
        "http://127.0.0.1:8000"
    ],
    allow_credentials=True,
    allow_methods=[
        "GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH", "HEAD"
    ],
    allow_headers=[
        "Accept",
        "Accept-Language",
        "Content-Language",
        "Content-Type",
        "Authorization",
        "X-Requested-With",
        "Origin",
        "Access-Control-Request-Method",
        "Access-Control-Request-Headers",
        "Cache-Control",
        "Pragma",
        "Expires",
        "X-Google-Auth-User"
    ],
    expose_headers=[
        "Content-Length",
        "Content-Type",
        "Date",
        "Server",
        "Access-Control-Allow-Origin",
        "Access-Control-Allow-Credentials"
    ],
    max_age=3600
)

@app.middleware("http")
async def cors_handler(request: Request, call_next):
    # Define allowed origins for production and development
    allowed_origins = [
        "https://casewise-beta.vercel.app",  # Production frontend
        "http://localhost:3000",              # Local development
        "http://localhost:3001",              # Alternative local port
        "http://127.0.0.1:3000",             # Local development
        "http://127.0.0.1:3001",             # Alternative local port
    ]
    
    # Get the origin from the request
    origin = request.headers.get("origin")
    print(f"🌐 CORS: Request origin: {origin}")
    print(f"🌐 CORS: Request method: {request.method}")
    print(f"🌐 CORS: Request URL: {request.url}")
    
    # Check if origin is allowed
    if origin and origin in allowed_origins:
        allowed_origin = origin
        print(f"✅ CORS: Allowed origin: {allowed_origin}")
    else:
        # Default to production frontend for unknown origins
        allowed_origin = "https://casewise-beta.vercel.app"
        print(f"⚠️ CORS: Unknown origin: {origin}, defaulting to production frontend")
    
    # Handle preflight OPTIONS requests
    if request.method == "OPTIONS":
        print("🔄 CORS: Handling preflight OPTIONS request")
        response = Response()
        response.headers["Access-Control-Allow-Origin"] = allowed_origin
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS, PATCH, HEAD"
        response.headers["Access-Control-Allow-Headers"] = "Accept, Accept-Language, Content-Language, Content-Type, Authorization, X-Requested-With, Origin, Access-Control-Request-Method, Access-Control-Request-Headers, Cache-Control, Pragma, Expires, X-Google-Auth-User"
        response.headers["Access-Control-Allow-Credentials"] = "true"
        response.headers["Access-Control-Max-Age"] = "3600"
        print(f"✅ CORS: Preflight response headers: {dict(response.headers)}")
        return response
    
    # Process the request
    print("🔄 CORS: Processing request")
    response = await call_next(request)
    
    # Add CORS headers to all responses
    response.headers["Access-Control-Allow-Origin"] = allowed_origin
    response.headers["Access-Control-Allow-Credentials"] = "true"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS, PATCH, HEAD"
    response.headers["Access-Control-Allow-Headers"] = "Accept, Accept-Language, Content-Language, Content-Type, Authorization, X-Requested-With, Origin, Access-Control-Request-Method, Access-Control-Request-Headers, Cache-Control, Pragma, Expires, X-Google-Auth-User"
    
    print(f"✅ CORS: Response headers added: {dict(response.headers)}")
    return response

@app.middleware("http")
async def catch_exceptions_middleware(request: Request, call_next):
    try:
        return await call_next(request)
    except Exception as e:
        # Log the error
        print(f"ERROR: Unhandled exception: {e}")
        import traceback
        traceback.print_exc()
        
        # Return error response with CORS headers
        origin = request.headers.get("origin", "https://casewise-beta.vercel.app")
        response = Response(
            content=f'{{"detail": "Internal server error: {str(e)}"}}',
            status_code=500,
            headers={
                "Content-Type": "application/json",
                "Access-Control-Allow-Origin": origin,
                "Access-Control-Allow-Credentials": "true"
            }
        )
        return response


@app.options("/{path:path}")
async def options_handler(path: str, request: Request):
    """Handle all OPTIONS requests for CORS preflight"""
    allowed_origins = [
        "https://casewise-beta.vercel.app",  # Production frontend
        "http://localhost:3000",              # Local development
        "http://localhost:3001",              # Alternative local port
        "http://127.0.0.1:3000",             # Local development
        "http://127.0.0.1:3001",             # Alternative local port
    ]
    
    origin = request.headers.get("origin")
    
    # Validate origin
    if origin and origin in allowed_origins:
        allowed_origin = origin
    else:
        # Default to production frontend for unknown origins
        allowed_origin = "https://casewise-beta.vercel.app"
    
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": allowed_origin,
            "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS, PATCH, HEAD",
            "Access-Control-Allow-Headers": "Accept, Accept-Language, Content-Language, Content-Type, Authorization, X-Requested-With, Origin, Access-Control-Request-Method, Access-Control-Request-Headers, Cache-Control, Pragma, Expires, X-Google-Auth-User",
            "Access-Control-Allow-Credentials": "true",
            "Access-Control-Max-Age": "3600"
        }
    )

@app.get("/test")
def test_endpoint():
    return {"message": "Test endpoint working", "status": "ok"}

@app.get("/health")
def health_check():
    """Health check endpoint with CORS headers"""
    return {
        "status": "healthy",
        "message": "Backend is running",
        "cors": "enabled",
        "timestamp": datetime.utcnow().isoformat()
    }

@app.get("/admin/test")
def admin_test():
    return {
        "message": "Admin endpoint is working",
        "timestamp": datetime.utcnow().isoformat()
    }

@app.post("/signup", response_model=UserResponse, status_code=status.HTTP_201_CREATED)
def signup(user_data: UserSignup):
    print("SIGNUP ENDPOINT CALLED")
    print(f"Email: {user_data.email}")
    print(f"Username: {user_data.username}")
    
    try:
        # Check if user already exists
        existing_user = db.users.find_one({
            "$or": [
                {"email": user_data.email}, 
                {"username": user_data.username}
            ]
        })
        
        if existing_user:
            if existing_user.get("email") == user_data.email:
                print(" Email already registered")
                raise HTTPException(status_code=400, detail="Email already registered")
            else:
                print(" Username already taken")
                raise HTTPException(status_code=400, detail="Username already taken")
        
        # Create new user
        print(" Hashing password...")
        hashed_password = hash_password(user_data.password)
        
        user_doc = {
            "email": user_data.email,
            "username": user_data.username,
            "hashed_password": hashed_password,
            "full_name": user_data.full_name,
            "is_active": True,
            "status": "active",
            "role": "user",
            "created_at": datetime.utcnow()
        }
        
        print(" Inserting user into database...")
        result = db.users.insert_one(user_doc)
        print(f" User created with ID: {result.inserted_id}")
        
        # Prepare response
        user_doc["id"] = str(result.inserted_id)
        del user_doc["hashed_password"]  
        del user_doc["_id"]  
        
        return user_doc
        
    except HTTPException as he:
        print(f" HTTP Exception: {he.detail}")
        raise
    except Exception as e:
        print(f" Unexpected error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Signup error: {str(e)}")

@app.post("/login", response_model=Token)
def login(user_credentials: UserLogin):
    # Removed debug prints for performance
    
    try:
        # Check if database is connected
        if db is None:
            print(" Database not connected")
            raise HTTPException(status_code=500, detail="Database connection error")
        
        # Find user by email
        user = db.users.find_one({"email": user_credentials.email})
        if not user:
            raise HTTPException(status_code=401, detail="Incorrect email or password")
        
        # Verify password
        password_valid = verify_password(user_credentials.password, user["hashed_password"])
        
        if not password_valid:
            raise HTTPException(status_code=401, detail="Incorrect email or password")
        
        # Check if user is active
        if not user.get("is_active", True):
            raise HTTPException(status_code=400, detail="Inactive user")
        
        # Create access token with different expiration based on remember_me
        # If remember_me is True, token expires in 30 days, otherwise 24 hours
        expires_delta = timedelta(days=30) if user_credentials.remember_me else timedelta(hours=24)
        access_token = create_access_token(data={"sub": str(user["_id"])}, expires_delta=expires_delta)
        # Removed debug print for performance
        
        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user_id": str(user["_id"]),
            "email": user["email"]
        }
        
    except HTTPException as he:
        print(f" HTTP Exception: {he.detail}")
        raise
    except Exception as e:
        print(f" Unexpected error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Login error: {str(e)}")

@app.post("/admin/login", response_model=Token)
def admin_login(user_credentials: UserLogin):
    print(" ADMIN LOGIN ENDPOINT CALLED")
    print(f" Email: {user_credentials.email}")
    print(f" Password length: {len(user_credentials.password)}")
    
    try:
        # Check if database is connected
        if db is None:
            print(" Database not connected")
            raise HTTPException(status_code=500, detail="Database connection error")
        
        # Find admin user by email
        user = db.users.find_one({"email": user_credentials.email, "role": "admin"})
        print(f" Admin user found: {user is not None}")
        
        if not user:
            print(" Admin user not found")
            raise HTTPException(status_code=401, detail="Incorrect email or password")
        
        print(f" Admin user details: {user.get('username')} - Active: {user.get('is_active', True)}")
        
        # Verify password
        password_valid = verify_password(user_credentials.password, user["hashed_password"])
        print(f" Password verification result: {password_valid}")
        
        if not password_valid:
            print(" Invalid password")
            raise HTTPException(status_code=401, detail="Incorrect email or password")
        
        # Check if user is active
        if not user.get("is_active", True):
            print(" Admin user inactive")
            raise HTTPException(status_code=400, detail="Inactive user")
        
        # Create access token
        access_token = create_access_token(data={"sub": str(user["_id"])})
        print(" Admin login successful - token created")
        
        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user_id": str(user["_id"]),
            "email": user["email"]
        }
        
    except HTTPException as he:
        print(f" HTTP Exception: {he.detail}")
        raise
    except Exception as e:
        print(f" Unexpected error during admin login: {str(e)}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.post("/auth/google", response_model=Token)
def google_auth(request: GoogleAuthRequest):
    """Authenticate user with Google OAuth"""
    print("GOOGLE_AUTH ENDPOINT CALLED")
    print(f"GOOGLE_AUTH: Token length: {len(request.id_token)}")
    print(f"GOOGLE_AUTH: Token preview: {request.id_token[:50]}...")
    
    try:
        # Verify the Google ID token
        print("GOOGLE_AUTH: Starting token verification...")
        idinfo = verify_google_token(request.id_token)
        print("GOOGLE_AUTH: Token verification successful")
        
        email = idinfo.get('email')
        name = idinfo.get('name', '')
        picture = idinfo.get('picture', '')
        
        if not email:
            raise HTTPException(status_code=400, detail="Email not provided by Google")
        
        print(f"GOOGLE_AUTH: Processing user: {email}")
        
        # Check if user already exists
        existing_user = db.users.find_one({"email": email})
        
        if existing_user:
            print(f"GOOGLE_AUTH: Existing user found: {email}")
            # Block login if user is inactive
            if not existing_user.get("is_active", True):
                raise HTTPException(status_code=400, detail="Inactive user")
            user = existing_user
        else:
            print(f"GOOGLE_AUTH: Creating new user: {email}")
            # Create new user
            username = email.split('@')[0]  # Use email prefix as username
            
            # Check if username is taken
            counter = 1
            original_username = username
            while db.users.find_one({"username": username}):
                username = f"{original_username}{counter}"
                counter += 1
            
            user_doc = {
                "email": email,
                "username": username,
                "full_name": name,
                "first_name": name.split(' ')[0] if name else '',
                "last_name": ' '.join(name.split(' ')[1:]) if name and len(name.split(' ')) > 1 else '',
                "profile_image_url": picture,
                "is_active": True,
                "created_at": datetime.utcnow(),
                "auth_provider": "google"
            }
            
            result = db.users.insert_one(user_doc)
            user = db.users.find_one({"_id": result.inserted_id})
            print(f"GOOGLE_AUTH: New user created with ID: {result.inserted_id}")
        
        # Create access token
        access_token = create_access_token(data={"sub": str(user["_id"])})
        print("GOOGLE_AUTH: Login successful - token created")
        
        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user_id": str(user["_id"]),
            "email": user["email"]
        }
        
    except HTTPException as he:
        print(f"GOOGLE_AUTH: HTTP Exception: {he.detail}")
        print(f"GOOGLE_AUTH: HTTP Exception status: {he.status_code}")
        raise
    except Exception as e:
        print(f"GOOGLE_AUTH: Unexpected error: {e}")
        print(f"GOOGLE_AUTH: Error type: {type(e).__name__}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Google authentication error: {str(e)}")

@app.post("/forgot-password")
async def forgot_password(request: ForgotPasswordRequest):
    """Send password reset email to user"""
    print(f"FORGOT_PASSWORD: Request for email: {request.email}")
    
    try:
        # Check if user exists
        user = db.users.find_one({"email": request.email})
        if not user:
            # Don't reveal if email exists or not for security
            return {"message": "If an account with that email exists, we've sent a password reset link."}
        
        # Generate reset token
        reset_token = generate_reset_token()
        expires_at = datetime.utcnow() + timedelta(hours=1)  # Token expires in 1 hour
        
        # Store reset token in database
        db.password_reset_tokens.insert_one({
            "user_id": str(user["_id"]),
            "email": request.email,
            "token": reset_token,
            "expires_at": expires_at,
            "used": False,
            "created_at": datetime.utcnow()
        })
        
        # Send email
        user_name = user.get("full_name") or user.get("first_name", "")
        email_sent = send_reset_password_email(request.email, reset_token, user_name)
        
        if email_sent:
            print(f"✅ Password reset email sent to: {request.email}")
            return {"message": "If an account with that email exists, we've sent a password reset link."}
        else:
            print(f"❌ Failed to send password reset email to: {request.email}")
            raise HTTPException(status_code=500, detail="Failed to send password reset email. Please try again later.")
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ FORGOT_PASSWORD: Unexpected error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="An error occurred. Please try again later.")

@app.post("/reset-password")
async def reset_password(request: ResetPasswordRequest):
    """Reset user password using token"""
    print(f"RESET_PASSWORD: Request with token: {request.token[:8]}...")
    
    try:
        # Find valid reset token
        reset_record = db.password_reset_tokens.find_one({
            "token": request.token,
            "used": False,
            "expires_at": {"$gt": datetime.utcnow()}
        })
        
        if not reset_record:
            raise HTTPException(status_code=400, detail="Invalid or expired reset token.")
        
        # Hash new password
        hashed_password = hashlib.sha256(request.new_password.encode()).hexdigest()
        
        # Update user password
        result = db.users.update_one(
            {"_id": ObjectId(reset_record["user_id"])},
            {
                "$set": {
                    "hashed_password": hashed_password,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        if result.modified_count == 0:
            raise HTTPException(status_code=404, detail="User not found.")
        
        # Mark token as used
        db.password_reset_tokens.update_one(
            {"_id": reset_record["_id"]},
            {"$set": {"used": True, "used_at": datetime.utcnow()}}
        )
        
        print(f"✅ Password reset successful for user: {reset_record['user_id']}")
        return {"message": "Password has been reset successfully. You can now log in with your new password."}
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ RESET_PASSWORD: Unexpected error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="An error occurred. Please try again later.")

@app.get("/me", response_model=UserResponse)
def get_current_user_info(current_user: dict = Depends(get_current_user)):
    user_response = {
        "id": current_user["id"],
        "email": current_user["email"],
        "username": current_user["username"],
        "full_name": current_user.get("full_name"),
        "first_name": current_user.get("first_name"),
        "last_name": current_user.get("last_name"),
        "profile_image_url": current_user.get("profile_image_url"),
        "bio": current_user.get("bio"),
        "role": current_user.get("role", "user"),
        "created_at": current_user["created_at"],
        "updated_at": current_user.get("updated_at")
    }
    return user_response

@app.post("/logout")
def logout():
    """Logout endpoint - token invalidation is handled client-side"""
    return {"message": "Logged out successfully"}

@app.get("/admin/users")
def get_all_users(current_user: dict = Depends(get_current_user)):
    """Get all users for admin analytics - admin only"""
    # Check if user is admin
    if current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    try:
        # Get all users from database
        users_cursor = db.users.find({}, {
            "_id": 1,
            "email": 1,
            "username": 1,
            "full_name": 1,
            "role": 1,
            "is_active": 1,
            "created_at": 1,
            "updated_at": 1,
            "analytics": 1,
            "mcq_attempted": 1,
            "total_questions_correct": 1,
            "total_questions_attempted": 1
        })
        
        users = []
        for user in users_cursor:
            # Get analytics data for each user
            analytics = user.get("analytics", {})
            
            # Count documents uploaded by user
            cases_uploaded = db.documents.count_documents({"uploaded_by": str(user["_id"])})
            
            # Get MCQ statistics from user document (direct fields, not analytics)
            mcq_attempted = user.get("mcq_attempted", 0)
            total_questions_correct = user.get("total_questions_correct", 0)
            total_questions_attempted = user.get("total_questions_attempted", 0)
            
            # Get chat sessions to calculate time spent
            chat_count = db.chats.count_documents({"user_id": str(user["_id"])})
            
            # Calculate time spent (simplified)
            time_spent_minutes = chat_count * 5  # Assume 5 minutes per chat session
            if time_spent_minutes < 60:
                time_spent = f"{time_spent_minutes} mins"
            else:
                hours = time_spent_minutes // 60
                minutes = time_spent_minutes % 60
                time_spent = f"{hours}h {minutes}m"
            
            # Calculate average score
            average_score = 0
            if total_questions_attempted > 0:
                average_score = int((total_questions_correct / total_questions_attempted) * 100)
            
            # Get last active date - use most recent activity from chats or documents
            activity_dates = []
            
            # Start with user's updated_at or created_at
            user_updated = user.get("updated_at")
            if user_updated and isinstance(user_updated, datetime):
                activity_dates.append(user_updated)
            else:
                user_created = user.get("created_at")
                if user_created and isinstance(user_created, datetime):
                    activity_dates.append(user_created)
            
            # Check for most recent chat activity
            most_recent_chat = db.chats.find_one(
                {"user_id": str(user["_id"])},
                sort=[("updated_at", -1)]
            )
            if most_recent_chat and most_recent_chat.get("updated_at"):
                chat_date = most_recent_chat["updated_at"]
                if isinstance(chat_date, datetime):
                    activity_dates.append(chat_date)
            
            # Check for most recent document upload
            most_recent_doc = db.documents.find_one(
                {"uploaded_by": str(user["_id"])},
                sort=[("uploaded_at", -1)]
            )
            if most_recent_doc and most_recent_doc.get("uploaded_at"):
                doc_date = most_recent_doc["uploaded_at"]
                if isinstance(doc_date, datetime):
                    activity_dates.append(doc_date)
            
            # Use the most recent date
            if activity_dates:
                last_active_date = max(activity_dates)
            else:
                last_active_date = datetime.now()
            
            # Format the date
            if isinstance(last_active_date, datetime):
                last_active_date = last_active_date.strftime("%Y-%m-%d")
            else:
                last_active_date = datetime.now().strftime("%Y-%m-%d")
            
            # Determine status based on is_active and status fields
            is_active = user.get("is_active", True)
            status_field = user.get("status")
            
            if status_field:
                user_status = status_field
            elif is_active:
                user_status = "active"
            else:
                user_status = "inactive"

            user_data = {
                "id": str(user["_id"]),
                "email": user["email"],
                "username": user["username"],
                "full_name": user.get("full_name"),
                "role": user.get("role", "user"),
                "status": user_status,
                "created_at": user["created_at"].strftime("%Y-%m-%d") if user.get("created_at") else "N/A",
                "last_active": last_active_date,
                "time_spent": time_spent,
                "cases_uploaded": cases_uploaded,
                "mcq_attempted": mcq_attempted,
                "most_questions_type": "Easy",  # Placeholder
                "total_cases": cases_uploaded,
                "total_mcqs": mcq_attempted,
                "total_questions_correct": total_questions_correct,
                "total_questions_attempted": total_questions_attempted,
                "average_score": average_score
            }
            users.append(user_data)
        
        return {"users": users}
        
    except Exception as e:
        print(f"Error fetching users: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch users data")

@app.put("/admin/users/update")
async def admin_update_users(
    request: AdminBulkUserUpdateRequest,
    current_user: dict = Depends(get_current_user)
):
    """Update multiple users' roles and status - admin only"""
    
    print(f"🔍 Admin update endpoint called by user: {current_user.get('id', 'unknown')}")
    print(f"🔍 User role: {current_user.get('role', 'unknown')}")
    print(f"🔍 Request data: {request}")
    
    # Check if current user is admin
    if current_user.get("role") != "admin":
        print("❌ Access denied - user is not admin")
        raise HTTPException(status_code=403, detail="Admin access required")
    
    print(f"🔄 Admin bulk user update request from user: {current_user['id']}")
    print(f"📊 Updates to process: {len(request.updates)}")
    
    try:
        updated_users = []
        errors = []
        
        for update in request.updates:
            try:
                # Validate user exists
                user = db.users.find_one({"_id": ObjectId(update.user_id)})
                if not user:
                    errors.append(f"User {update.user_id} not found")
                    continue
                
                # Prepare update data
                update_data = {}
                if update.role is not None:
                    update_data["role"] = update.role
                if update.status is not None:
                    # Allow only active/inactive; inactive disables login
                    if update.status == "active":
                        update_data["is_active"] = True
                        update_data["status"] = "active"
                    elif update.status == "inactive":
                        update_data["is_active"] = False
                        update_data["status"] = "inactive"
                
                if not update_data:
                    errors.append(f"No valid fields to update for user {update.user_id}")
                    continue
                
                update_data["updated_at"] = datetime.utcnow()
                
                # Update user in database
                result = db.users.update_one(
                    {"_id": ObjectId(update.user_id)},
                    {"$set": update_data}
                )
                
                if result.modified_count > 0:
                    updated_users.append({
                        "user_id": update.user_id,
                        "email": user["email"],
                        "username": user["username"],
                        "updates": update_data
                    })
                    print(f"✅ Updated user {update.user_id}: {update_data}")
                else:
                    errors.append(f"Failed to update user {update.user_id}")
                    
            except Exception as e:
                error_msg = f"Error updating user {update.user_id}: {str(e)}"
                errors.append(error_msg)
                print(f"❌ {error_msg}")
        
        print(f"📊 Update summary: {len(updated_users)} successful, {len(errors)} errors")
        
        return {
            "message": f"Updated {len(updated_users)} users successfully",
            "updated_users": updated_users,
            "errors": errors,
            "total_processed": len(request.updates)
        }
        
    except Exception as e:
        print(f"❌ Error in admin bulk user update: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to update users")

@app.put("/profile", response_model=UserResponse)
async def update_profile(
    profile_data: ProfileUpdateRequest,
    current_user: dict = Depends(get_current_user)
):
    """Update user profile information"""
    
    print(f" Profile update request for user: {current_user['id']}")
    
    # Check if email is being changed and if it's already taken
    if profile_data.email and profile_data.email != current_user["email"]:
        existing_user = db.users.find_one({"email": profile_data.email})
        if existing_user:
            raise HTTPException(
                status_code=400, 
                detail="Email already exists"
            )
    
    # Check if username is being changed and if it's already taken
    if profile_data.username and profile_data.username != current_user["username"]:
        existing_user = db.users.find_one({"username": profile_data.username})
        if existing_user:
            raise HTTPException(
                status_code=400, 
                detail="Username already exists"
            )
    
    # Prepare update data
    update_data = {}
    
    if profile_data.first_name is not None:
        update_data["first_name"] = profile_data.first_name
    if profile_data.last_name is not None:
        update_data["last_name"] = profile_data.last_name
    if profile_data.full_name is not None:
        update_data["full_name"] = profile_data.full_name
    if profile_data.bio is not None:
        update_data["bio"] = profile_data.bio
    if profile_data.email is not None:
        update_data["email"] = profile_data.email
    if profile_data.username is not None:
        update_data["username"] = profile_data.username
    
    if not update_data:
        raise HTTPException(
            status_code=400, 
            detail="No fields to update"
        )
    
    update_data["updated_at"] = datetime.utcnow()
    
    # Update user in database
    try:
        result = db.users.update_one(
            {"_id": ObjectId(current_user["id"])},
            {"$set": update_data}
        )
        
        if result.modified_count == 0:
            raise HTTPException(
                status_code=404, 
                detail="User not found"
            )
        
        print(f" Profile updated successfully for user: {current_user['id']}")
        
        # Create notification for profile update
        create_notification(
            user_id=current_user["id"],
            notification_type="profile_update",
            title="Profile Updated",
            message="Your profile has been successfully updated.",
            metadata={"updated_fields": list(update_data.keys())}
        )
        
        # Fetch updated user data
        updated_user = db.users.find_one({"_id": ObjectId(current_user["id"])})
        
        if not updated_user:
            raise HTTPException(
                status_code=404, 
                detail="User not found after update"
            )
        
        # Prepare response
        user_response = {
            "id": str(updated_user["_id"]),
            "email": updated_user["email"],
            "username": updated_user["username"],
            "full_name": updated_user.get("full_name"),
            "first_name": updated_user.get("first_name"),
            "last_name": updated_user.get("last_name"),
            "profile_image_url": updated_user.get("profile_image_url"),
            "bio": updated_user.get("bio"),
            "created_at": updated_user["created_at"],
            "updated_at": updated_user.get("updated_at")
        }
        
        return user_response
        
    except Exception as e:
        print(f" Error updating profile: {str(e)}")
        raise HTTPException(
            status_code=500, 
            detail="Failed to update profile"
        )

@app.put("/profile/password")
async def change_password(
    password_data: PasswordChangeRequest,
    current_user: dict = Depends(get_current_user)
):
    """Change user password"""
    
    print(f" Password change request for user: {current_user['id']}")
    
    # Verify current password
    user = db.users.find_one({"_id": ObjectId(current_user["id"])})
    if not user:
        raise HTTPException(
            status_code=404, 
            detail="User not found"
        )
    
    # Verify current password
    if not verify_password(password_data.current_password, user["hashed_password"]):
        raise HTTPException(
            status_code=400, 
            detail="Current password is incorrect"
        )
    
    # Hash new password
    new_hashed_password = hash_password(password_data.new_password)
    
    # Update password in database
    try:
        result = db.users.update_one(
            {"_id": ObjectId(current_user["id"])},
            {
                "$set": {
                    "hashed_password": new_hashed_password,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        if result.modified_count == 0:
            raise HTTPException(
                status_code=404, 
                detail="User not found"
            )
        
        print(f" Password updated successfully for user: {current_user['id']}")
        
        return {"message": "Password updated successfully"}
        
    except Exception as e:
        print(f" Error updating password: {str(e)}")
        raise HTTPException(
            status_code=500, 
            detail="Failed to update password"
        )

@app.post("/profile/upload-image")
async def upload_profile_image(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload profile image"""
    
    print(f" Profile image upload for user: {current_user['id']}")
    
    # Validate file type
    if not file.content_type or not file.content_type.startswith('image/'):
        raise HTTPException(
            status_code=400, 
            detail="File must be an image"
        )
    
    # Validate file size (2MB limit)
    max_size = 2 * 1024 * 1024  # 2MB
    if file.size and file.size > max_size:
        raise HTTPException(
            status_code=400, 
            detail="File size must be less than 2MB"
        )
    
    try:
        # Read file content
        content = await file.read()
        
        # For now, we'll store the image as base64 in the database
        # In production, you'd want to use a cloud storage service like AWS S3
        import base64
        image_base64 = base64.b64encode(content).decode('utf-8')
        image_url = f"data:{file.content_type};base64,{image_base64}"
        
        # Update user profile with image URL
        result = db.users.update_one(
            {"_id": ObjectId(current_user["id"])},
            {
                "$set": {
                    "profile_image_url": image_url,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        if result.modified_count == 0:
            raise HTTPException(
                status_code=404, 
                detail="User not found"
            )
        
        print(f" Profile image uploaded successfully for user: {current_user['id']}")
        
        # Create notification for profile image upload
        create_notification(
            user_id=current_user["id"],
            notification_type="profile_update",
            title="Profile Image Updated",
            message="Your profile image has been successfully updated.",
            metadata={"file_size": len(content), "content_type": file.content_type}
        )
        
        return {
            "message": "Profile image uploaded successfully",
            "image_url": image_url
        }
        
    except Exception as e:
        print(f" Error uploading profile image: {str(e)}")
        raise HTTPException(
            status_code=500, 
            detail="Failed to upload profile image"
        )

# Notification endpoints
@app.get("/notifications", response_model=List[NotificationResponse])
async def get_notifications(
    current_user: dict = Depends(get_current_user),
    limit: int = 50,
    unread_only: bool = False
):
    """Get user notifications"""
    
    print(f" Getting notifications for user: {current_user['id']}")
    
    try:
        query = {"user_id": current_user["id"]}
        if unread_only:
            query["is_read"] = False
        
        notifications = list(db.notifications.find(query)
                           .sort("created_at", -1)
                           .limit(limit))
        
        notification_list = []
        for notification in notifications:
            # Handle both datetime objects and ISO strings
            created_at = notification["created_at"]
            if isinstance(created_at, str):
                # Already an ISO string
                created_at_str = created_at
            elif hasattr(created_at, 'isoformat'):
                # Convert datetime to ISO string
                created_at_str = created_at.isoformat()
            else:
                # Fallback to string representation
                created_at_str = str(created_at)
            
            notification_list.append({
                "id": str(notification["_id"]),
                "type": notification["type"],
                "title": notification["title"],
                "message": notification["message"],
                "is_read": notification["is_read"],
                "created_at": created_at_str,
                "metadata": notification.get("metadata")
            })
        
        print(f" Found {len(notification_list)} notifications")
        return notification_list
        
    except Exception as e:
        print(f" Error getting notifications: {str(e)}")
        raise HTTPException(
            status_code=500, 
            detail="Failed to get notifications"
        )

@app.put("/notifications/{notification_id}/read")
async def mark_notification_read(
    notification_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Mark a notification as read"""
    
    print(f" Marking notification {notification_id} as read for user: {current_user['id']}")
    
    try:
        result = db.notifications.update_one(
            {
                "_id": ObjectId(notification_id),
                "user_id": current_user["id"]
            },
            {"$set": {"is_read": True}}
        )
        
        if result.modified_count == 0:
            raise HTTPException(
                status_code=404, 
                detail="Notification not found"
            )
        
        print(f" Notification {notification_id} marked as read")
        return {"message": "Notification marked as read"}
        
    except Exception as e:
        print(f" Error marking notification as read: {str(e)}")
        raise HTTPException(
            status_code=500, 
            detail="Failed to mark notification as read"
        )

@app.put("/notifications/read-all")
async def mark_all_notifications_read(
    current_user: dict = Depends(get_current_user)
):
    """Mark all notifications as read for the current user"""
    
    print(f" Marking all notifications as read for user: {current_user['id']}")
    
    try:
        result = db.notifications.update_many(
            {
                "user_id": current_user["id"],
                "is_read": False
            },
            {"$set": {"is_read": True}}
        )
        
        print(f" Marked {result.modified_count} notifications as read")
        return {"message": f"Marked {result.modified_count} notifications as read"}
        
    except Exception as e:
        print(f" Error marking all notifications as read: {str(e)}")
        raise HTTPException(
            status_code=500, 
            detail="Failed to mark all notifications as read"
        )

@app.get("/notifications/unread-count")
async def get_unread_count(
    current_user: dict = Depends(get_current_user)
):
    """Get count of unread notifications"""
    
    try:
        count = db.notifications.count_documents({
            "user_id": current_user["id"],
            "is_read": False
        })
        
        return {"unread_count": count}
        
    except Exception as e:
        print(f" Error getting unread count: {str(e)}")
        raise HTTPException(
            status_code=500, 
            detail="Failed to get unread count"
        )

@app.get("/analytics/user", response_model=UserAnalytics)
async def get_user_analytics(current_user: dict = Depends(get_current_user)):
    """Get user analytics and statistics"""
    
    try:
        user_id = current_user["id"]
        
        # Get user info
        user = db.users.find_one({"_id": ObjectId(user_id)})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Calculate statistics
        # Count documents uploaded by user
        cases_uploaded = db.documents.count_documents({"uploaded_by": user_id})
        
        # Get MCQ statistics from user document
        mcq_attempted = user.get("mcq_attempted", 0)
        total_questions_correct = user.get("total_questions_correct", 0)
        total_questions_attempted = user.get("total_questions_attempted", 0)
        
        # Get chat sessions to calculate time spent (optimized)
        chat_count = db.chats.count_documents({"user_id": user_id})
        
        # Calculate time spent (simplified - could be more sophisticated)
        time_spent_minutes = chat_count * 5  # Assume 5 minutes per chat session
        if time_spent_minutes < 60:
            time_spent = f"{time_spent_minutes} mins"
        else:
            hours = time_spent_minutes // 60
            minutes = time_spent_minutes % 60
            time_spent = f"{hours}h {minutes}m"
        
        # Determine most common question type based on attempted questions by difficulty
        diff_counts = user.get("difficulty_counts", {}) or {}
        if diff_counts:
            # pick max of easy/moderate/hard
            normalized = {k.lower(): int(v) for k, v in diff_counts.items()}
            choice = max([("Easy", normalized.get("easy", 0)),
                          ("Moderate", normalized.get("moderate", 0)),
                          ("Hard", normalized.get("hard", 0))], key=lambda x: x[1])[0]
            most_questions_type = choice
        else:
            most_questions_type = "Easy"
        
        # Calculate average score
        average_score = 0
        if total_questions_attempted > 0:
            average_score = int((total_questions_correct / total_questions_attempted) * 100)
        
        # Get last active date
        last_active_date = user.get("updated_at", user.get("created_at", datetime.now()))
        if isinstance(last_active_date, datetime):
            last_active_date = last_active_date.strftime("%Y-%m-%d")
        else:
            last_active_date = datetime.now().strftime("%Y-%m-%d")
        
        # Build response
        analytics = UserAnalytics(
            name=user.get("full_name", user.get("username", "User")),
            timeSpent=time_spent,
            casesUploaded=cases_uploaded,
            mcqAttempted=mcq_attempted,
            mostQuestionsType=most_questions_type,
            totalQuestionsCorrect=total_questions_correct,
            totalQuestionsAttempted=total_questions_attempted,
            averageScore=average_score,
            lastActiveDate=last_active_date
        )
        
        return analytics
        
    except Exception as e:
        print(f"Error getting user analytics: {str(e)}")
        raise HTTPException(
            status_code=500, 
            detail="Failed to get user analytics"
        )

class MCQCompletionRequest(BaseModel):
    """MCQ completion request schema"""
    correct_answers: int
    total_questions: int
    case_id: Optional[str] = None
    case_difficulty: Optional[str] = None

@app.post("/analytics/mcq-completion")
async def update_mcq_analytics(
    request: MCQCompletionRequest,
    current_user: dict = Depends(get_current_user)
):
    """Update user analytics after MCQ completion"""
    
    print(f"🔄 MCQ Analytics Update Request:")
    print(f"   User ID: {current_user.get('id', 'N/A')}")
    print(f"   Correct Answers: {request.correct_answers}")
    print(f"   Total Questions: {request.total_questions}")
    print(f"   Case ID: {request.case_id}")
    
    try:
        user_id = current_user["id"]
        
        # Get user info
        user = db.users.find_one({"_id": ObjectId(user_id)})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Update user's MCQ statistics
        # For now, we'll store this in the user document
        # In a real app, you might want a separate analytics collection
        
        current_mcq_attempted = user.get("mcq_attempted", 0)
        current_correct = user.get("total_questions_correct", 0)
        current_attempted = user.get("total_questions_attempted", 0)
        
        # Update the statistics
        # mcq_attempted should count total questions attempted, not cases
        new_mcq_attempted = current_mcq_attempted + request.total_questions
        new_correct = current_correct + request.correct_answers
        new_attempted = current_attempted + request.total_questions
        
        # Update difficulty counters based on attempted questions
        difficulty_field = (request.case_difficulty or "").lower()
        if difficulty_field in ["easy", "moderate", "hard"]:
            diff_counts = user.get("difficulty_counts", {})
            current = int(diff_counts.get(difficulty_field, 0))
            diff_counts[difficulty_field] = current + request.total_questions
        else:
            diff_counts = user.get("difficulty_counts", {})

        # Calculate new average score
        new_average_score = int((new_correct / new_attempted) * 100) if new_attempted > 0 else 0
        
        # Update user document
        update_result = db.users.update_one(
            {"_id": ObjectId(user_id)},
            {
                "$set": {
                    "mcq_attempted": new_mcq_attempted,
                    "total_questions_correct": new_correct,
                    "total_questions_attempted": new_attempted,
                    "average_score": new_average_score,
                    "difficulty_counts": diff_counts,
                    "updated_at": datetime.now()
                }
            }
        )
        
        print(f"📊 Database update result: {update_result.modified_count} documents modified")
        print(f"📊 New stats: MCQ={new_mcq_attempted}, Correct={new_correct}, Total={new_attempted}, Avg={new_average_score}%")
        
        return {
            "message": "Analytics updated successfully",
            "updated_stats": {
                "mcq_attempted": new_mcq_attempted,
                "total_questions_correct": new_correct,
                "total_questions_attempted": new_attempted,
                "average_score": new_average_score
            }
        }
        
    except Exception as e:
        print(f"Error updating MCQ analytics: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Failed to update analytics"
        )

@app.post("/analytics/clear")
async def clear_user_analytics(current_user: dict = Depends(get_current_user)):
    """Clear analytics data for the current user"""
    
    try:
        user_id = current_user["id"]
        
        # Clear analytics fields for the user
        result = db.users.update_one(
            {"_id": ObjectId(user_id)},
            {
                "$unset": {
                    "mcq_attempted": "",
                    "total_questions_correct": "",
                    "total_questions_attempted": "",
                    "average_score": ""
                }
            }
        )
        
        if result.modified_count > 0:
            return {"message": "Analytics data cleared successfully"}
        else:
            return {"message": "No analytics data found to clear"}
        
    except Exception as e:
        print(f"Error clearing analytics: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Failed to clear analytics"
        )

@app.post("/documents/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload a document and store it in MongoDB with enhanced processing"""
    
    print(f" Document upload started: {file.filename} ({file.content_type})")
    
    # Validate file size (20MB limit)
    max_size = 20 * 1024 * 1024  # 20MB
    if file.size and file.size > max_size:
        raise HTTPException(
            status_code=400, 
            detail=f"File size ({file.size} bytes) exceeds 20MB limit"
        )
    
    # Validate file type
    allowed_types = [
        "text/plain", "application/pdf", "application/msword", 
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/markdown"
    ]
    
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=400, 
            detail=f"File type {file.content_type} not supported. Allowed types: {', '.join(allowed_types)}"
        )
    
    # Read file content
    try:
        content = await file.read()
        content_str = None
        
        if file.content_type in ["text/plain", "text/markdown"]:
            content_str = content.decode('utf-8')
            print(f" Text file processed: {len(content_str)} characters")
        elif file.content_type == "application/pdf":
            if PDF_AVAILABLE and PDF_LIBRARY:
                # Extract actual content from PDF using available library
                try:
                    if PDF_LIBRARY == "pypdf":
                        import pypdf
                        pdf_reader = pypdf.PdfReader(io.BytesIO(content))
                    else:  # PyPDF2
                        import PyPDF2
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                    
                    content_str = ""
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            content_str += f"\n--- Page {page_num + 1} ---\n"
                            content_str += page_text
                    
                    if not content_str.strip():
                        # Fallback if no text could be extracted
                        content_str = f"PDF file: {file.filename}\n\nNote: This PDF file could not be processed for text extraction. Please ensure the PDF contains selectable text."
                    
                    print(f" PDF processed with {PDF_LIBRARY}: {len(content_str)} characters extracted from {len(pdf_reader.pages)} pages")
                    
                except Exception as e:
                    print(f" PDF processing error with {PDF_LIBRARY}: {e}")
                    content_str = f"PDF file: {file.filename}\n\nError: Could not extract text from PDF using {PDF_LIBRARY}. Please ensure the file is not corrupted and contains selectable text."
            else:
                # Fallback content when no PDF library is available - provide sample medical content for testing
                content_str = f"""Medical Case Study - {file.filename}

PATIENT INFORMATION:
Name: Ramesh Kumar
Age: 45 years
Gender: Male
Chief Complaint: Chest pain and shortness of breath

HISTORY OF PRESENT ILLNESS:
Patient presents with acute onset chest pain that started 2 hours ago. The pain is described as crushing, substernal, and radiating to the left arm. Associated symptoms include diaphoresis, nausea, and shortness of breath. Patient has a history of hypertension and diabetes mellitus type 2.

PHYSICAL EXAMINATION:
Vital Signs: BP 160/95 mmHg, HR 110 bpm, RR 24/min, Temp 98.6°F, O2 Sat 92% on room air
Cardiovascular: Regular rhythm, no murmurs, JVD not elevated
Respiratory: Bilateral rales in lower lung fields
Abdomen: Soft, non-tender, no organomegaly

DIAGNOSTIC WORKUP:
ECG: ST elevation in leads II, III, aVF (inferior STEMI)
Troponin I: 15.2 ng/mL (elevated)
CK-MB: 45 U/L (elevated)
Chest X-ray: Mild pulmonary edema

ASSESSMENT AND PLAN:
1. Acute ST-elevation myocardial infarction (STEMI) - inferior wall
2. Hypertension
3. Diabetes mellitus type 2

Treatment plan includes:
- Immediate cardiac catheterization for primary PCI
- Dual antiplatelet therapy (aspirin + clopidogrel)
- Beta-blocker (metoprolol)
- ACE inhibitor (lisinopril)
- Statin therapy (atorvastatin)
- Blood glucose monitoring and insulin management

PROGNOSIS:
Patient requires immediate intervention. Door-to-balloon time should be less than 90 minutes for optimal outcomes.

Note: This is sample medical content for testing purposes. In a real scenario, this would be extracted from the actual PDF document."""
                print(f" PDF file processed with enhanced fallback content: {len(content_str)} characters")
        elif file.content_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            # Process Word documents (DOC and DOCX)
            try:
                if file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    # DOCX file processing
                    import docx
                    doc = docx.Document(io.BytesIO(content))
                    content_str = ""
                    
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            content_str += paragraph.text + "\n"
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    content_str += cell.text + " "
                            content_str += "\n"
                    
                    if not content_str.strip():
                        content_str = f"DOCX file: {file.filename}\n\nNote: This DOCX file could not be processed for text extraction. Please ensure the document contains readable text."
                    
                    print(f" DOCX processed: {len(content_str)} characters extracted")
                    
                elif file.content_type == "application/msword":
                    # DOC file processing using docx2txt
                    import docx2txt
                    content_str = docx2txt.process(io.BytesIO(content))
                    
                    if not content_str.strip():
                        content_str = f"DOC file: {file.filename}\n\nNote: This DOC file could not be processed for text extraction. Please ensure the document contains readable text."
                    
                    print(f" DOC processed: {len(content_str)} characters extracted")
                    
            except ImportError as e:
                print(f" Missing library for Word document processing: {e}")
                content_str = f"Word document: {file.filename}\n\nError: Required libraries for Word document processing are not installed. Please install python-docx and docx2txt."
            except Exception as e:
                print(f" Error processing Word document: {e}")
                content_str = f"Word document: {file.filename}\n\nError: Could not extract text from Word document. Please ensure the file is not corrupted and contains readable text."
        else:
            content_str = f"Document: {file.filename}. Content type: {file.content_type}. Please use text files for full functionality."
            print(f" Unsupported file type processed with placeholder content")
            
    except Exception as e:
        print(f" Error reading file: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")
    
    # Create document record
    document_doc = {
        "filename": file.filename,
        "content_type": file.content_type,
        "size": len(content),
        "content": content_str,
        "uploaded_by": current_user["id"],
        "uploaded_at": datetime.utcnow(),
        "processing_status": "completed"
    }
    
    # Store in MongoDB
    result = db.documents.insert_one(document_doc)
    document_id = str(result.inserted_id)
    
    print(f" Document stored successfully with ID: {document_id}")
    
    # Create notification for file upload
    create_notification(
        user_id=current_user["id"],
        notification_type="file_upload",
        title="Document Uploaded",
        message=f"Your document '{file.filename}' has been successfully uploaded and processed.",
        metadata={
            "document_id": document_id,
            "filename": file.filename,
            "file_size": len(content),
            "content_type": file.content_type
        }
    )
    
    # Return response
    document_doc["id"] = document_id
    del document_doc["_id"]
    
    return document_doc

@app.post("/documents/upload-enhanced", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document_enhanced(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload a document with enhanced PDF extraction and automatic case/MCQ detection"""
    
    print(f" Enhanced document upload started: {file.filename} ({file.content_type})")
    
    # Validate file size (20MB limit)
    max_size = 20 * 1024 * 1024
    if file.size and file.size > max_size:
        raise HTTPException(
            status_code=400, 
            detail=f"File size ({file.size} bytes) exceeds 20MB limit"
        )
    
    # Validate file type
    allowed_types = [
        "text/plain", "application/pdf", "application/msword", 
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/markdown"
    ]
    
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=400, 
            detail=f"File type {file.content_type} not supported"
        )
    
    # Read file content
    try:
        content = await file.read()
        content_str = None
        extraction_metadata = {}
        
        if file.content_type in ["text/plain", "text/markdown"]:
            content_str = content.decode('utf-8')
            print(f" Text file processed: {len(content_str)} characters")
            
        elif file.content_type == "application/pdf":
            if PDF_AVAILABLE and PDF_LIBRARY == "pypdf":
                try:
                    extractor = PDFContentExtractor()
                    result = extractor.extract_text_from_pdf(content)
                    
                    content_str = result["full_text"]
                    
                    # Store extraction metadata
                    extraction_metadata = {
                        "total_pages": result["total_pages"],
                        "has_cases": result["has_cases"],
                        "has_mcqs": result["has_mcqs"],
                        "detected_cases_count": len(result["detected_cases"]),
                        "detected_mcqs_count": len(result["detected_mcqs"]),
                        "detected_cases": result["detected_cases"],
                        "detected_mcqs": result["detected_mcqs"]
                    }
                    
                    print(f" Enhanced PDF processing:")
                    print(f"   - Pages: {result['total_pages']}")
                    print(f"   - Cases detected: {len(result['detected_cases'])}")
                    print(f"   - MCQs detected: {len(result['detected_mcqs'])}")
                    print(f"   - Characters extracted: {len(content_str)}")
                    
                except Exception as e:
                    print(f" Enhanced PDF processing error: {e}")
                    # Fallback to basic extraction
                    pdf_reader = pypdf.PdfReader(io.BytesIO(content))
                    content_str = ""
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            content_str += f"\n--- Page {page_num + 1} ---\n"
                            content_str += page_text
            else:
                # Fallback content
                content_str = f"PDF file: {file.filename}\n\nPDF processing library not available."
                
        elif file.content_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            # Process Word documents (DOC and DOCX) with enhanced processing
            try:
                if file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    # DOCX file processing
                    import docx
                    doc = docx.Document(io.BytesIO(content))
                    content_str = ""
                    
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            content_str += paragraph.text + "\n"
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    content_str += cell.text + " "
                            content_str += "\n"
                    
                    if not content_str.strip():
                        content_str = f"DOCX file: {file.filename}\n\nNote: This DOCX file could not be processed for text extraction. Please ensure the document contains readable text."
                    
                    print(f" Enhanced DOCX processed: {len(content_str)} characters extracted")
                    
                elif file.content_type == "application/msword":
                    # DOC file processing using python-docx2txt
                    import docx2txt
                    content_str = docx2txt.process(io.BytesIO(content))
                    
                    if not content_str.strip():
                        content_str = f"DOC file: {file.filename}\n\nNote: This DOC file could not be processed for text extraction. Please ensure the document contains readable text."
                    
                    print(f" Enhanced DOC processed: {len(content_str)} characters extracted")
                    
                # For Word documents, we can also try to detect cases and MCQs
                if content_str and len(content_str) > 100:  # Only if we have substantial content
                    # Simple case detection for Word documents
                    case_indicators = ["case study", "patient", "diagnosis", "treatment", "symptoms", "medical history"]
                    mcq_indicators = ["question", "answer", "option", "a)", "b)", "c)", "d)", "multiple choice"]
                    
                    has_cases = any(indicator.lower() in content_str.lower() for indicator in case_indicators)
                    has_mcqs = any(indicator.lower() in content_str.lower() for indicator in mcq_indicators)
                    
                    extraction_metadata = {
                        "has_cases": has_cases,
                        "has_mcqs": has_mcqs,
                        "detected_cases_count": 1 if has_cases else 0,
                        "detected_mcqs_count": 1 if has_mcqs else 0,
                        "word_document_processed": True
                    }
                    
                    print(f" Enhanced Word document analysis:")
                    print(f"   - Cases detected: {has_cases}")
                    print(f"   - MCQs detected: {has_mcqs}")
                    print(f"   - Characters extracted: {len(content_str)}")
                    
            except ImportError as e:
                print(f" Missing library for Word document processing: {e}")
                content_str = f"Word document: {file.filename}\n\nError: Required libraries for Word document processing are not installed. Please install python-docx and docx2txt."
            except Exception as e:
                print(f" Error processing Word document: {e}")
                content_str = f"Word document: {file.filename}\n\nError: Could not extract text from Word document. Please ensure the file is not corrupted and contains readable text."
                
    except Exception as e:
        print(f" Error reading file: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")
    
    # Create document record
    document_doc = {
        "filename": file.filename,
        "content_type": file.content_type,
        "size": len(content),
        "content": content_str,
        "uploaded_by": current_user["id"],
        "uploaded_at": datetime.utcnow(),
        "processing_status": "completed",
        "extraction_metadata": extraction_metadata
    }
    
    # Store in MongoDB
    result = db.documents.insert_one(document_doc)
    document_id = str(result.inserted_id)
    
    print(f" Enhanced document stored successfully with ID: {document_id}")
    
    # Create notification for file upload
    create_notification(
        user_id=current_user["id"],
        notification_type="file_upload",
        title="Document Uploaded",
        message=f"Your document '{file.filename}' has been successfully uploaded and processed.",
        metadata={
            "document_id": document_id,
            "filename": file.filename,
            "file_size": len(content),
            "content_type": file.content_type,
            "extraction_metadata": extraction_metadata
        }
    )
    
    # Return response
    document_doc["id"] = document_id
    del document_doc["_id"]
    
    return document_doc

@app.get("/documents/{document_id}/extracted-content")
async def get_extracted_content(
    document_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get extracted cases and MCQs from a document"""
    
    try:
        document = db.documents.find_one({
            "_id": ObjectId(document_id),
            "uploaded_by": current_user["id"]
        })
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid document ID")
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    extraction_metadata = document.get("extraction_metadata", {})
    
    return {
        "document_id": document_id,
        "filename": document.get("filename"),
        "has_cases": extraction_metadata.get("has_cases", False),
        "has_mcqs": extraction_metadata.get("has_mcqs", False),
        "detected_cases": extraction_metadata.get("detected_cases", []),
        "detected_mcqs": extraction_metadata.get("detected_mcqs", []),
        "total_cases": extraction_metadata.get("detected_cases_count", 0),
        "total_mcqs": extraction_metadata.get("detected_mcqs_count", 0)
    }

@app.get("/test-openai")
async def test_openai():
    """Test OpenAI API and list available models"""
    try:
        if not openai_client:
            return {
                "api_key_configured": False,
                "api_key_prefix": None,
                "available_models": [],
                "error": "OpenAI API key not configured"
            }
        
        # List available models
        models = openai_client.models.list()
        available_models = [model.id for model in models.data if 'gpt' in model.id.lower()]
        
        return {
            "api_key_configured": bool(OPENAI_API_KEY),
            "api_key_prefix": OPENAI_API_KEY[:10] if OPENAI_API_KEY else None,
            "available_models": available_models
        }
    except Exception as e:
        return {
            "error": str(e),
            "error_type": type(e).__name__
        }

@app.post("/ai/generate-cases", response_model=CaseGenerationResponse)
async def generate_case_scenarios(
    request: CaseGenerationRequest,
    current_user: dict = Depends(get_current_user)
):
    """Generate case scenarios from a document using OpenAI GPT-4 mini"""
    
    if not OPENAI_API_KEY:
        raise HTTPException(status_code=500, detail="OpenAI API key not configured")
    
    # ---- Tiny difficulty config block (central place to tweak later) ----
    # Product spec: Easy / Moderate / Hard
    DIFFICULTY_SEQUENCE = ["Easy", "Moderate", "Hard"]
    
    # 1) Find the document
    try:
        document = db.documents.find_one({
            "_id": ObjectId(request.document_id),
            "uploaded_by": current_user["id"]
        })
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid document ID")
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")
    
    # Check if document has text content
    if not document.get("content"):
        raise HTTPException(status_code=400, detail="Document does not contain readable text content")
    
    try:
        # 2) Use OpenAI GPT-4 mini
        if not openai_client:
            raise Exception("OpenAI client not initialized")
        
        # Create the prompt for case generation
        system_prompt = f"""
        You are an expert medical case scenario generator specializing in creating comprehensive, educational medical cases for medical students. Based on the following document content and user prompt, generate {request.num_scenarios} realistic, detailed medical case scenarios.

        Document Content:
        {document['content'][:3000]}  # Increased context for better relevance

        User Prompt: {request.prompt}

        IMPORTANT INSTRUCTIONS:
        - Create realistic, clinically relevant cases
        - Include detailed patient presentations with specific symptoms, vital signs, and history
        - Provide comprehensive case descriptions (minimum 2-3 paragraphs each)
        - Include specific learning objectives and clinical reasoning points
        - Make cases challenging but educational
        - Include relevant diagnostic considerations and treatment approaches

        For each scenario, provide:
        1. A clear, descriptive title that indicates the main condition
        2. A detailed, comprehensive case description including:
           - Patient demographics and presenting complaint
           - Detailed history of present illness
           - Relevant past medical history
           - Physical examination findings
           - Initial diagnostic considerations
        3. 5-7 specific key learning points covering:
           - Pathophysiology
           - Diagnostic criteria
           - Differential diagnosis
           - Treatment options
           - Clinical pearls
        4. Appropriate difficulty level (Easy, Moderate, or Hard)

        Format your response as a JSON array with the following structure:
        [
            {{
                "title": "Specific Case Title (e.g., 'Acute Myocardial Infarction in a 55-year-old Male')",
                "description": "Comprehensive case description with detailed patient presentation, history, examination findings, and clinical context...",
                "key_points": ["Specific learning point 1", "Specific learning point 2", "Specific learning point 3", "Specific learning point 4", "Specific learning point 5"],
                "difficulty": "Easy/Moderate/Hard"
            }}
        ]
        """
        
        # Generate response from OpenAI
        response = openai_client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert medical case scenario generator. Always respond with valid JSON format."
                },
                {"role": "user", "content": system_prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        
        import json
        
        raw_content = response.choices[0].message.content or ""
        
        # 3) Small robustness: strip markdown fences if model wraps output
        text = raw_content.strip()
        if text.startswith("```"):
            # remove leading ``` or ```json
            text = text.strip("`")
            if text.lower().startswith("json"):
                text = text[4:].strip()
        
        try:
            scenarios_data = json.loads(text)
        except json.JSONDecodeError:
            # If JSON parsing fails, create a fallback response
            scenarios_data = [{
                "title": f"Generated Case {i+1}",
                "description": f"Case scenario based on: {request.prompt}",
                "key_points": [
                    "Key learning point 1",
                    "Key learning point 2",
                    "Key learning point 3"
                ],
                "difficulty": "Moderate"
            } for i in range(request.num_scenarios)]
        
        # 4) Normalize + enforce difficulty + safety defaults
        num_scenarios = min(len(scenarios_data), request.num_scenarios)
        normalized_scenarios = []
        for i in range(num_scenarios):
            raw = scenarios_data[i] or {}
            scenario = dict(raw)  # shallow copy
            
            # Force difficulty from our sequence instead of trusting the model
            forced_diff = DIFFICULTY_SEQUENCE[i % len(DIFFICULTY_SEQUENCE)]
            scenario["difficulty"] = forced_diff
            
            # Basic safety defaults so your Pydantic model doesn't explode
            if not scenario.get("title"):
                scenario["title"] = f"Generated Case {i+1}"
            if not scenario.get("description"):
                scenario["description"] = f"Case scenario based on: {request.prompt}"
            if not isinstance(scenario.get("key_points"), list) or not scenario["key_points"]:
                scenario["key_points"] = ["Key learning point derived from this case."]
            
            normalized_scenarios.append(scenario)
        
        # 5) Persist generated cases so Key Concepts / Explore Case can find them
        try:
            for scen in normalized_scenarios:
                db.generated_cases.insert_one(
                    {
                        "document_id": str(document["_id"]),   # IMPORTANT: store as string
                        "title": scen["title"],
                        "description": scen["description"],
                        "key_points": scen.get("key_points", []),
                        "difficulty": scen.get("difficulty", "Moderate"),
                        # You can add more here later if needed (e.g., age, gender)
                        "created_by": current_user["id"],
                        "created_at": datetime.utcnow(),
                    }
                )
        except Exception as e:
            # Non‑fatal: log but don't break the endpoint
            print(f"⚠️ Warning: failed to persist generated_cases: {e}")
        
        # 6) Convert to CaseScenario objects for the response
        scenarios = [CaseScenario(**scenario) for scenario in normalized_scenarios]
        
        return CaseGenerationResponse(
            document_id=request.document_id,
            prompt=request.prompt,
            scenarios=scenarios,
            generated_at=datetime.utcnow()
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating cases: {str(e)}")

@app.post("/ai/generate-case-titles", response_model=CaseTitlesResponse)
async def generate_case_titles(
    request: CaseTitleRequest,
    current_user: dict = Depends(get_current_user)
):
    """Generate case titles from a document using OpenAI GPT-4 mini"""
    
    if not OPENAI_API_KEY:
        raise HTTPException(status_code=500, detail="OpenAI API key not configured")
    
    # Find the document
    try:
        document = db.documents.find_one({
            "_id": ObjectId(request.document_id),
            "uploaded_by": current_user["id"]
        })
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid document ID")
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")
    
    # Check if document has text content
    if not document.get("content"):
        raise HTTPException(status_code=400, detail="Document does not contain readable text content")
    
    try:
        # Use OpenAI GPT-4 mini
        if not openai_client:
            raise Exception("OpenAI client not initialized")
        
        # Determine difficulty distribution
        num_cases = request.num_cases
        if num_cases == 5:
            difficulty_distribution = "EXACTLY 1 Moderate case (first), 2 Moderate cases (second and third), and 2 Hard cases (fourth and fifth) in this EXACT order"
            difficulty_requirements = """
        CRITICAL DIFFICULTY REQUIREMENTS (for 5 cases - MUST FOLLOW THIS EXACT ORDER):
        - Case 1 (FIRST): MUST be "Moderate" - Moderate complexity, some diagnostic challenges, requires intermediate clinical reasoning. Use moderately complex concepts from the document.
        - Case 2 (SECOND): MUST be "Moderate" - Moderate complexity, some diagnostic challenges, requires intermediate clinical reasoning. Use moderately complex concepts from the document.
        - Case 3 (THIRD): MUST be "Moderate" - Moderate complexity, some diagnostic challenges, requires intermediate clinical reasoning. Use different moderately complex concepts from the document.
        - Case 4 (FOURTH): MUST be "Hard" - Complex presentation, multiple differential diagnoses, requires advanced clinical reasoning and integration of multiple concepts. Use complex, nuanced concepts from the document.
        - Case 5 (FIFTH): MUST be "Hard" - Complex presentation, multiple differential diagnoses, requires advanced clinical reasoning and integration of multiple concepts. Use different complex, nuanced concepts from the document.
        
        The case description MUST match the difficulty level:
        - Easy: Clear, straightforward presentation with obvious clinical signs. Simple diagnostic path. Basic pathophysiology. Directly based on document content.
        - Moderate: Some complexity in presentation, requires connecting multiple findings. Moderate diagnostic challenge. Based on document content but requires some reasoning.
        - Hard: Complex, nuanced presentation with subtle findings. Multiple possible diagnoses. Requires advanced reasoning and knowledge integration. Based on complex concepts from document."""
        else:
            difficulty_distribution = f"Distribute difficulty levels appropriately across {num_cases} cases"
            difficulty_requirements = "        - Assign difficulty levels (Easy, Moderate, Hard) that match the complexity of each case"
        
        # Create the prompt for case title generation
        system_prompt = f"""
        You are an expert medical case generator. Based on the following document content, generate {num_cases} realistic medical case titles with brief descriptions that are DIRECTLY RELEVANT to the document content.

        Document Content:
        {document['content'][:3000]}

        CRITICAL RELEVANCE REQUIREMENTS:
        - Cases MUST be directly based on the medical concepts, conditions, and information in the document above
        - Use specific medical terminology, conditions, and details from the document
        - Do NOT generate generic cases - every case must reflect actual content from the document
        - Extract key medical concepts, diseases, symptoms, treatments, or procedures mentioned in the document
        - Ensure each case scenario incorporates specific information from the document content

        IMPORTANT INSTRUCTIONS:
        - Create realistic, clinically relevant case titles that are SPECIFICALLY RELATED to the document content
        - Each case should have a clear, descriptive title (1-2 lines)
        - DO NOT include case numbers (e.g., "Case 1:", "Case 2:") in the titles - just use descriptive titles like "Acute Myocardial Infarction in a 55-year-old Male"
        - CRITICAL: Each case description MUST be exactly 250-300 words (minimum 250, maximum 300 words). Count the words carefully. The description should provide comprehensive context including patient presentation, clinical findings, diagnostic considerations, treatment approaches, and important learning points.
        {difficulty_requirements}
        - Make cases diverse and educational
        - Focus on different aspects of the medical content

        CRITICAL: You must respond with ONLY a valid JSON array. Do not include any markdown formatting, explanations, or additional text. Start your response directly with [ and end with ].

        Format your response as a JSON array with the following structure:
        [
            {{
                "id": "case_1",
                "title": "Specific Case Title WITHOUT case numbers (e.g., 'Acute Myocardial Infarction in a 55-year-old Male' - NOT 'Case 1: Acute Myocardial Infarction')",
                "description": "A detailed description that is EXACTLY 250-300 words (count carefully) and MATCHES the difficulty level. Easy cases should be straightforward with clear findings. Hard cases should be complex with subtle findings and multiple differentials. Include comprehensive patient presentation, detailed clinical findings, diagnostic considerations, treatment approaches, and important learning points. The description must be between 250 and 300 words - no less, no more.",
                "difficulty": "Easy" or "Moderate" or "Hard" (must match case complexity)
            }}
        ]
        
        CRITICAL FOR 5 CASES: 
        - Generate exactly 5 unique cases in this EXACT order: [Easy, Moderate, Moderate, Hard, Hard]
        - Each case description must be comprehensive (250-300 words)
        - The difficulty level MUST match the complexity of the case description
        - ALL cases MUST be directly relevant to the document content above
        - Use specific medical concepts, conditions, and terminology from the document
        - Case 1 = Moderate, Case 2 = Moderate, Case 3 = Moderate, Case 4 = Hard, Case 5 = Hard
        """
        
        # Generate response from OpenAI
        system_message = "You are an expert medical case generator. Each case description MUST be 250-300 words. Always respond with valid JSON format."
        if request.num_cases == 5:
            system_message += " CRITICAL: When generating 5 cases, you MUST create exactly 3 Moderate cases (first, second, and third), and 2 Hard cases (fourth and fifth) in this EXACT order. The case descriptions must match their difficulty levels and be directly relevant to the document content."
        
        response = openai_client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": system_prompt}
            ],
            temperature=0.7,
            max_tokens=4000,  # Increased to accommodate 250-300 word descriptions for multiple cases
            timeout=120  # 2 minutes timeout for case title generation
        )
        
        # Parse the response
        import json
        print(f"AI Raw OpenAI response: {response.choices[0].message.content}")
        
        try:
            # Try to extract JSON from the response if it's wrapped in markdown
            response_text = response.choices[0].message.content.strip()
            if response_text.startswith('```json'):
                response_text = response_text.replace('```json', '').replace('```', '').strip()
            elif response_text.startswith('```'):
                response_text = response_text.replace('```', '').strip()
            
            cases_data = json.loads(response_text)
            print(f"SUCCESS: Successfully parsed JSON: {len(cases_data)} cases")
            
            # Validate and enforce difficulty distribution for 5 cases
            if request.num_cases == 5 and len(cases_data) >= 5:
                # Force correct distribution in exact order: Easy, Moderate, Moderate, Hard, Hard
                required_distribution = ["Easy", "Moderate", "Moderate", "Hard", "Hard"]
                
                # Check current distribution
                difficulty_counts = {"Easy": 0, "Moderate": 0, "Hard": 0}
                for case in cases_data[:5]:
                    diff = case.get("difficulty", "Moderate").strip()
                    # Normalize difficulty
                    if diff.lower() == "easy":
                        difficulty_counts["Easy"] += 1
                    elif diff.lower() == "moderate":
                        difficulty_counts["Moderate"] += 1
                    elif diff.lower() == "hard":
                        difficulty_counts["Hard"] += 1
                    else:
                        difficulty_counts["Moderate"] += 1  # Default to Moderate
                
                print(f"📊 Difficulty distribution before fix: Easy={difficulty_counts['Easy']}, Moderate={difficulty_counts['Moderate']}, Hard={difficulty_counts['Hard']}")
                
                # Always enforce correct distribution in order (even if already correct, ensure order is right)
                if difficulty_counts["Easy"] != 1 or difficulty_counts["Moderate"] != 2 or difficulty_counts["Hard"] != 2:
                    print(f"⚠️ WARNING: Difficulty distribution incorrect. Expected: 1 Easy, 2 Moderate, 2 Hard. Got: {difficulty_counts}")
                
                print(f"🔧 Enforcing correct distribution in order: [Easy, Moderate, Moderate, Hard, Hard]")
                
                # Force correct distribution in exact order (always, to ensure consistency)
                for i, case in enumerate(cases_data[:5]):
                    case["difficulty"] = required_distribution[i]
                    print(f"  Case {i+1}: Set difficulty to {required_distribution[i]}")
                
                # Verify final distribution
                final_counts = {"Easy": 0, "Moderate": 0, "Hard": 0}
                for case in cases_data[:5]:
                    final_counts[case.get("difficulty", "Moderate")] += 1
                print(f"✅ Final difficulty distribution: Easy={final_counts['Easy']}, Moderate={final_counts['Moderate']}, Hard={final_counts['Hard']}")
        except json.JSONDecodeError as e:
            print(f"ERROR: JSON parsing failed: {e}")
            print(f"Response text: {response.choices[0].message.content[:500]}...")
            # If JSON parsing fails, create a fallback response based on document content
            # Create a longer fallback description (250-300 words)
            fallback_description = f"""This medical case scenario is based on the uploaded document content and presents a comprehensive clinical scenario designed for medical education. The case includes a detailed patient presentation with comprehensive demographic information, presenting complaint, and relevant social history that provides important context for understanding the clinical situation. The patient's medical history is thoroughly documented, including past medical conditions, previous surgeries, family history, and any relevant genetic or environmental factors that may influence the current presentation. Physical examination findings are described in detail, including vital signs, general appearance, and system-by-system examination results with both positive and negative findings that are crucial for differential diagnosis. Diagnostic test results are provided with specific values, reference ranges, and clinical interpretation to help students understand how laboratory and imaging studies contribute to the diagnostic process. Treatment considerations are explored, including first-line and alternative therapeutic options, with discussion of mechanism of action, indications, contraindications, and potential side effects. The case also addresses important learning points related to pathophysiology, clinical reasoning, differential diagnosis, and evidence-based medicine principles. This comprehensive approach ensures that medical students gain a thorough understanding of the clinical scenario, develop critical thinking skills, and learn to apply medical knowledge in realistic patient care situations. The case is designed to be educational, clinically relevant, and aligned with current medical practice guidelines and standards of care."""

            cases_data = [{
                "id": f"case_{i+1}",
                "title": f"Case {i+1} from {document.get('filename', 'Document')}",
                "description": fallback_description,
                "difficulty": "Moderate"
            } for i in range(request.num_cases)]
        
        # Convert to CaseTitle objects
        cases = [CaseTitle(**case) for case in cases_data[:request.num_cases]]
        
        return CaseTitlesResponse(
            document_id=request.document_id,
            cases=cases,
            generated_at=datetime.utcnow()
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating case titles: {str(e)}")

@app.post("/ai/generate-mcqs", response_model=MCQResponse)
async def generate_mcqs(
    request: MCQRequest,
    current_user: dict = Depends(get_current_user)
):
    """Generate MCQ questions from a document or case using OpenAI GPT-4 mini"""
    
    print(f"🎯 MCQ Generation Request - Difficulty: {request.difficulty}, Case: {request.case_title}")
    
    if not OPENAI_API_KEY:
        raise HTTPException(status_code=500, detail="OpenAI API key not configured")
    
    document_context = ""
    
    def clean_document_content(content):
        """Remove metadata and non-medical content from document."""
        if not content:
            return ""
        
        lines = content.split('\n')
        cleaned_lines = []
        skip_section = False
        
        # Keywords that indicate metadata sections to skip
        metadata_indicators = [
            "author:", "authors:", "publication:", "published:", "journal:",
            "copyright", "doi:", "isbn:", "issn:", "reference", "citation",
            "bibliography", "acknowledgment", "affiliation", "institution:",
            "university:", "department:", "page", "chapter", "section",
            "abstract:", "keywords:", "corresponding author"
        ]
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Skip lines that are clearly metadata
            if any(indicator in line_lower for indicator in metadata_indicators):
                # Check if it's actually metadata (not a medical term)
                if any(term in line_lower for term in ["author", "publication", "journal", "copyright", 
                                                       "doi", "isbn", "issn", "reference", "citation",
                                                       "bibliography", "acknowledgment", "affiliation"]):
                    skip_section = True
                    continue
            
            # Skip empty lines if we're in a metadata section
            if skip_section and not line.strip():
                continue
            
            # Reset skip flag if we hit a new substantial line (likely medical content)
            if skip_section and len(line.strip()) > 20 and not any(indicator in line_lower for indicator in metadata_indicators):
                skip_section = False
            
            if not skip_section:
                cleaned_lines.append(line)
        
        cleaned = '\n'.join(cleaned_lines)
        
        # Remove common metadata patterns
        import re
        # Remove lines with email patterns
        cleaned = re.sub(r'[\w\.-]+@[\w\.-]+\.\w+', '', cleaned)
        # Remove lines with URL patterns
        cleaned = re.sub(r'https?://\S+', '', cleaned)
        # Remove lines that are just numbers (likely page numbers)
        cleaned = re.sub(r'^\s*\d+\s*$', '', cleaned, flags=re.MULTILINE)
        
        return cleaned.strip()
    
    # Get document context if document_id is provided
    if request.document_id:
        try:
            document = db.documents.find_one({
                "_id": ObjectId(request.document_id),
                "uploaded_by": current_user["id"]
            })
            if document and document.get("content"):
                raw_content = document['content']
                # Clean metadata before using
                cleaned_content = clean_document_content(raw_content)
                # If cleaning removed too much, use original but with warning
                if len(cleaned_content.strip()) < 100 and len(raw_content) > 500:
                    print(f"WARNING: Document cleaning removed too much content. Using original with metadata warning.")
                    document_context = raw_content[:800]
                else:
                    document_context = cleaned_content[:800]  # Increased after cleaning
        except Exception as e:
            print(f"ERROR: Failed to fetch/clean document: {e}")
            pass
    
    if not document_context or len(document_context.strip()) < 50:
        raise HTTPException(status_code=400, detail="No document content available for MCQ generation")
    
    # Fetch case details if case_title is provided
    case_details = None
    case_demographics = ""
    if request.case_title and request.document_id:
        import re
        title = request.case_title
        case_description = ""
        
        # First, try to extract demographics directly from the case title (most reliable)
        age_match = re.search(r'(\d+)[-\s]year[-\s]old', title, re.IGNORECASE)
        gender_match = re.search(r'\b(girl|boy|female|male|woman|man|child)\b', title, re.IGNORECASE)
        
        age = age_match.group(1) if age_match else None
        gender = gender_match.group(1).lower() if gender_match else None
        
        # Try to find the case in generated_cases collection for full description
        try:
            doc_id_str = str(request.document_id)
            # Try multiple query formats
            case_doc = db.generated_cases.find_one({
                "title": request.case_title,
                "document_id": doc_id_str
            })
            if not case_doc:
                # Try without document_id constraint (in case format differs)
                case_doc = db.generated_cases.find_one({
                    "title": request.case_title
                })
            
            if case_doc:
                case_details = case_doc
                case_description = case_doc.get("description", "")
                
                # Extract from description if not found in title
                if not age:
                    age_match = re.search(r'(\d+)[-\s]year[-\s]old', case_description, re.IGNORECASE)
                    age = age_match.group(1) if age_match else None
                if not gender:
                    gender_match = re.search(r'\b(girl|boy|female|male|woman|man|child)\b', case_description, re.IGNORECASE)
                    gender = gender_match.group(1).lower() if gender_match else None
        except Exception as e:
            print(f"⚠️ Warning: Could not fetch case details from database for MCQ generation: {e}")
            # Continue with title-based extraction
        
        # Build demographics string
        if age and gender:
            # Normalize gender
            if gender in ['girl', 'female', 'woman']:
                gender_normalized = 'female'
            elif gender in ['boy', 'male', 'man']:
                gender_normalized = 'male'
            else:
                gender_normalized = gender
            
            case_demographics = f"\n\nCRITICAL PATIENT DEMOGRAPHICS CONSISTENCY REQUIREMENT:\n- The patient in this case is a {age}-year-old {gender_normalized}\n- MCQ questions should be consistent with this case, but AVOID repeating the same demographics in every question\n- Vary the phrasing: Use 'this patient', 'the patient', 'the same patient', or reference clinical findings instead of repeating age/gender in every question\n- Only include demographics when clinically relevant to the specific question being asked\n- Focus on clinical scenarios, findings, or management rather than repeating demographics unnecessarily"
        elif case_description:
            # If we can't extract specific demographics, use the full case description
            case_demographics = f"\n\nCRITICAL CASE CONTEXT:\n- Full Case Description: {case_description[:500]}\n- ALL MCQ questions MUST be consistent with this specific case scenario\n- Maintain consistency with patient demographics, presentation, and clinical details from the case description above"
        elif request.case_title:
            # Last resort: use the case title itself
            case_demographics = f"\n\nCRITICAL CASE CONTEXT:\n- Case Title: {request.case_title}\n- ALL MCQ questions MUST be consistent with this specific case scenario\n- Extract and maintain consistency with patient demographics from the case title"
    
    try:
        # Use OpenAI GPT-4 mini
        if not openai_client:
            raise Exception("OpenAI client not initialized")
        
        # Create the prompt for MCQ generation
        case_context = ""
        if request.case_title:
            if case_demographics:
                case_context = f"\n\nSpecific Case Focus: {request.case_title}{case_demographics}\n\nGenerate MCQs specifically related to this case scenario. Maintain strict consistency with the patient demographics provided above."
            else:
                case_context = f"\n\nSpecific Case Focus: {request.case_title}\nGenerate MCQs specifically related to this case scenario."
        
        # Build difficulty instruction - make it very explicit
        difficulty_instruction = ""
        if request.difficulty:
            difficulty_instruction = f"""- Assign difficulty (Easy/Moderate/Hard)"""
        
        system_prompt = f"""TASK: GENERATE A SINGLE-BEST-ANSWER MULTIPLE-CHOICE QUESTION (MCQ) WITH RATIONALE

Based on the preceding medical case information, generate {request.num_questions} multiple-choice question(s) designed to test high-yield, essential clinical concepts related to this case.

CRITICAL DIVERSITY REQUIREMENT - VARY QUESTION TYPES AND FOCUS:
When generating {request.num_questions} questions, you MUST create VARIED question types with DIFFERENT focuses. Do NOT make all questions about the same patient or patient-specific scenarios. Instead, create a diverse mix:

1. **Patient-Specific Questions** (use sparingly, max 1-2 out of {request.num_questions}):
   - Can reference "this patient" or "the patient" without repeating demographics
   - Focus on specific clinical findings, lab results, or management decisions

2. **Case-Based Scenario Questions** (preferred):
   - "In this case scenario, what is the most likely mechanism..."
   - "Based on the clinical presentation described, which of the following..."
   - "Given the findings in this case, what would be the next best step..."

3. **General Concept Questions** (highly encouraged):
   - "What is the mechanism of action of [medication mentioned in case]?"
   - "Which of the following is a key pathophysiological feature of [condition in case]?"
   - "What is the most important diagnostic test for [condition related to case]?"
   - "Which medication class is first-line for [condition in case]?"

4. **Situational/Clinical Reasoning Questions**:
   - "A patient presents with [symptoms from case]. What is the most likely diagnosis?"
   - "In a patient with [condition from case], which finding would be most concerning?"
   - "What is the most appropriate management strategy for [situation from case]?"

5. **Mechanism/Pathophysiology Questions**:
   - "What is the underlying pathophysiological mechanism of [condition in case]?"
   - "How does [medication from case] exert its therapeutic effect?"
   - "Which molecular pathway is primarily involved in [disease from case]?"

6. **Diagnostic/Management Questions**:
   - "What is the most appropriate initial investigation for [condition from case]?"
   - "Which of the following is a contraindication to [treatment from case]?"
   - "What is the most important monitoring parameter for [medication from case]?"

MCQ Structure and Constraints:

Question Stem: Vary your question stems. Do NOT start every question with patient demographics. Use diverse approaches:
- "What is the mechanism of..."
- "Which of the following is the most likely..."
- "In patients with [condition], what is..."
- "What is the most appropriate..."
- "Which diagnostic test is most useful for..."
- "What is the pathophysiological basis of..." 

Options: Provide exactly five (5) answer options (A, B, C, D, E). Only one option must be unequivocally correct. 

Content Focus (Essential Knowledge): MUST test essential, high-yield clinical knowledge (e.g., differential diagnosis, next best step in management, critical pathophysiology, or drug mechanism). MAY test the name of a discoverer or a named disease/syndrome (e.g., Hashimoto's disease, Cushing's triad). MUST NOT test 'good-to-know' or trivial information.

CRITICAL CONTENT RESTRICTION - STRICTLY PROHIBITED:
You MUST NEVER create questions about document metadata, bibliographic information, or non-medical content. This includes but is not limited to:
- Author names, researcher names, or institutional affiliations
- Publication dates, years, or journal names
- Document titles, section headers, or chapter numbers (unless they are medical terms)
- File names, document IDs, or reference numbers
- Publisher information, copyright notices, or citation details
- Page numbers, line numbers, or formatting details
- Any administrative or bibliographic metadata

ALL questions MUST focus exclusively on:
- Medical concepts, pathophysiology, diagnosis, and treatment
- Clinical reasoning and decision-making
- Disease mechanisms, drug actions, and therapeutic principles
- Patient presentation, examination findings, and investigations
- Evidence-based medicine and clinical guidelines

If the source document contains metadata or non-medical information, IGNORE IT COMPLETELY. Only use the actual medical/clinical content for question generation.

Distractors (Incorrect Options): The four incorrect options (distractors) must be **plausible**. They should represent common misconceptions, less likely differential diagnoses, or inappropriate next steps in management to effectively test the student's clinical reasoning and differentiation skills. 

Mandatory Requirement: Answer Rationale

Explanation: Immediately after the correct answer identifier, provide a concise Explanation (Rationale) for the correct answer. This explanation must: 

a. Justify the Correct Answer: State clearly *why* the chosen option is correct. 

MEDICAL CONTENT (metadata removed - use only medical/clinical information):
{document_context}{case_context}

REMINDER: The content above has been cleaned of metadata. If you see any author names, publication dates, journal names, or bibliographic information, IGNORE IT. Only use medical concepts, pathophysiology, diagnosis, treatment, and clinical information.

CRITICAL: ALL questions must have EXACTLY the same difficulty level: "{request.difficulty or 'Moderate'}"

CRITICAL JSON FORMAT REQUIREMENT:
You MUST return a valid JSON array. Each question MUST have this EXACT structure:

[
  {{
    "id": "mcq_1",
    "question": "What is the mechanism of action of metformin in improving glycemic control?",
    "options": [
      {{"id": "A", "text": "Option A text here", "is_correct": false}},
      {{"id": "B", "text": "Option B text here (correct answer)", "is_correct": true}},
      {{"id": "C", "text": "Option C text here", "is_correct": false}},
      {{"id": "D", "text": "Option D text here", "is_correct": false}},
      {{"id": "E", "text": "Option E text here", "is_correct": false}}
    ],
    "explanation": "The correct answer is B because [detailed explanation]",
    "difficulty": "{request.difficulty or 'Moderate'}"
  }},
  {{
    "id": "mcq_2",
    "question": "In patients with type 2 diabetes and asthma, which of the following is the most important consideration?",
    "options": [
      {{"id": "A", "text": "Option A", "is_correct": false}},
      {{"id": "B", "text": "Option B", "is_correct": false}},
      {{"id": "C", "text": "Option C (correct)", "is_correct": true}},
      {{"id": "D", "text": "Option D", "is_correct": false}},
      {{"id": "E", "text": "Option E", "is_correct": false}}
    ],
    "explanation": "The correct answer is C because [explanation]",
    "difficulty": "{request.difficulty or 'Moderate'}"
  }}
]

CRITICAL REQUIREMENTS:
- Return ONLY a valid JSON array (no markdown, no code fences, no extra text)
- Each question MUST have exactly 5 options (A, B, C, D, E)
- Each option MUST be a dictionary with "id" (string), "text" (string), and "is_correct" (boolean)
- Exactly ONE option per question must have "is_correct": true
- ALL questions must have difficulty: "{request.difficulty or 'Moderate'}"
- Vary question types - do NOT make all questions about the same patient

Return JSON array only with ALL questions having difficulty: "{request.difficulty or 'Moderate'}" """
        
        # Generate response from OpenAI with optimized settings
        try:
            response = openai_client.chat.completions.create(
                model="gpt-4.1",
                messages=[
                    {"role": "system", "content": f"You are a medical educator. Generate single-best-answer MCQs with exactly 5 options (A-E). Test high-yield essential clinical concepts with plausible distractors. Include mandatory answer rationale. ALL questions must have EXACTLY the same difficulty: {request.difficulty or 'Moderate'}. Always respond with valid JSON format. CRITICAL DIVERSITY: When generating multiple questions, create VARIED question types - NOT all about the same patient. Mix patient-specific questions (max 1-2), case-based scenarios, general concept questions, mechanism/pathophysiology questions, diagnostic questions, and management questions. Do NOT repeat the same patient demographics or start every question with patient information. STRICTLY PROHIBITED: NEVER create questions about document metadata, author names, publication dates, journal names, file names, or any bibliographic/non-medical information. Only focus on medical concepts, pathophysiology, diagnosis, and treatment."},
                    {"role": "user", "content": system_prompt}
                ],
                temperature=0.3,  # Slightly higher for faster generation while maintaining quality
                max_tokens=4000,  # Increased to ensure complete questions with full options are generated
                timeout=120  # Increased timeout for better quality generation
            )
        except Exception as e:
            print(f"OpenAI API error: {e}")
            # Fallback: try with even more optimized settings
            response = openai_client.chat.completions.create(
                model="gpt-4.1",
                messages=[
                    {"role": "system", "content": f"Generate single-best-answer medical MCQs with 5 options (A-E) in JSON format. ALL questions must have difficulty: {request.difficulty or 'Moderate'}. Questions must test high-yield clinical concepts with plausible distractors. CRITICAL DIVERSITY: Create VARIED question types - NOT all about the same patient. Mix patient-specific, case-based, general concept, mechanism, diagnostic, and management questions. Do NOT repeat demographics or start every question with patient information. STRICTLY PROHIBITED: NEVER create questions about document metadata, author names, publication dates, journal names, file names, or any bibliographic/non-medical information. Only focus on medical concepts, pathophysiology, diagnosis, and treatment."},
                    {"role": "user", "content": f"""Create {request.num_questions} case-based medical MCQs with 5 options each from the following MEDICAL CONTENT ONLY (ignore any metadata, author names, publication info, or bibliographic details):

{document_context[:800]}

CRITICAL REQUIREMENTS:
- ALL questions must be {request.difficulty or 'Moderate'} difficulty
- Each question must test essential clinical knowledge with plausible distractors
- STRICTLY PROHIBITED: Do NOT create questions about author names, publication dates, journal names, document titles, file names, or any bibliographic metadata
- ONLY focus on medical concepts, pathophysiology, diagnosis, treatment, and clinical reasoning
- If you see any metadata in the content above, IGNORE IT COMPLETELY

Return valid JSON array with exactly 5 options per question."""}
                ],
                temperature=0.3,  # Slightly higher for faster generation
                max_tokens=3500,  # Increased for fallback generation
                timeout=90  # Increased timeout for fallback generation
            )
        
        # Parse the response
        import json
        print(f"AI Raw MCQ OpenAI response: {response.choices[0].message.content}")
        
        try:
            # Try to extract JSON from the response if it's wrapped in markdown
            response_text = response.choices[0].message.content.strip()
            
            # Remove markdown code blocks
            if response_text.startswith('```json'):
                response_text = response_text.replace('```json', '').replace('```', '').strip()
            elif response_text.startswith('```'):
                response_text = response_text.replace('```', '').strip()
            
            # Try to find JSON array boundaries if there's extra text
            start_idx = response_text.find('[')
            end_idx = response_text.rfind(']') + 1
            
            if start_idx != -1 and end_idx > start_idx:
                json_text = response_text[start_idx:end_idx]
                mcqs_data = json.loads(json_text)
            else:
                # Try parsing the whole text
                mcqs_data = json.loads(response_text)
            
            # Validate that we got actual questions, not placeholders
            if not isinstance(mcqs_data, list) or len(mcqs_data) == 0:
                raise ValueError("No valid questions in response")
            
            # Check for placeholder questions and validate structure
            for i, mcq in enumerate(mcqs_data):
                if not isinstance(mcq, dict):
                    raise ValueError(f"Question {i+1} is not a dictionary: {type(mcq)}")
                
                if not mcq.get("question") or "Case-based question" in mcq.get("question", "") or "based on document content" in mcq.get("question", "").lower():
                    raise ValueError("Placeholder questions detected in response")
                
                # Validate options structure
                if "options" not in mcq:
                    raise ValueError(f"Question {i+1} missing 'options' field")
                
                if not isinstance(mcq["options"], list):
                    raise ValueError(f"Question {i+1} has invalid options type: {type(mcq['options'])}")
                
                if len(mcq["options"]) == 0:
                    raise ValueError(f"Question {i+1} has no options")
                
                # Log first question structure for debugging
                if i == 0:
                    print(f"DEBUG: First question structure: id={mcq.get('id')}, question={mcq.get('question')[:50]}..., options_count={len(mcq.get('options', []))}")
                    if mcq.get("options"):
                        print(f"DEBUG: First option type: {type(mcq['options'][0])}")
                        print(f"DEBUG: First option value: {mcq['options'][0]}")
                        print(f"DEBUG: All options types: {[type(opt) for opt in mcq['options']]}")
                        print(f"DEBUG: All options: {mcq['options']}")
            
            print(f"SUCCESS: Successfully parsed MCQ JSON: {len(mcqs_data)} questions")
        except (json.JSONDecodeError, ValueError) as e:
            print(f"ERROR: MCQ JSON parsing/validation failed: {e}")
            print(f"Response text: {response.choices[0].message.content[:1000]}...")
            # Instead of creating placeholder questions, raise an error to trigger retry
            raise HTTPException(
                status_code=500, 
                detail=f"Failed to generate valid MCQs. The AI response could not be parsed or contained placeholder content. Please try again. Error: {str(e)}"
            )
        
        # Validate questions don't contain metadata/non-medical content
        metadata_keywords = [
            "author", "publication", "journal", "published", "copyright", 
            "publisher", "doi", "isbn", "issn", "reference", "citation",
            "bibliography", "abstract", "acknowledgment", "affiliation",
            "institution", "university", "department", "page", "chapter",
            "section header", "document id", "file name", "metadata"
        ]
        
        def contains_metadata(text):
            """Check if text contains metadata-related keywords in a medical context."""
            if not text:
                return False
            text_lower = text.lower()
            
            # Context-aware detection - check for metadata patterns, not just keywords
            # Patterns that indicate metadata (not medical terms)
            metadata_patterns = [
                r'\bauthor[s]?\s*[:=]\s*\w+',  # "Author: John Smith"
                r'\bpublication\s*[:=]',        # "Publication:"
                r'\bjournal\s*[:=]\s*\w+',     # "Journal: Nature"
                r'\bpublished\s+in\s+\d{4}',   # "Published in 2024"
                r'\bcopyright\s+\d{4}',        # "Copyright 2024"
                r'\bdoi\s*[:=]\s*10\.',        # "DOI: 10.1234/..."
                r'\bisbn\s*[:=]',              # "ISBN:"
                r'\bissn\s*[:=]',              # "ISSN:"
                r'\breference[s]?\s*[:=]\s*\d+', # "Reference: 1" or "References:"
                r'\bcitation[s]?\s*[:=]',      # "Citation:"
                r'\bbibliography\s*[:=]',      # "Bibliography:"
                r'\backnowledgment[s]?\s*[:=]', # "Acknowledgment:"
                r'\baffiliation[s]?\s*[:=]',   # "Affiliation:"
                r'\binstitution\s*[:=]\s*\w+', # "Institution: Harvard"
                r'\bpage\s+\d+\s+of\s+\d+',    # "Page 5 of 10"
                r'\bchapter\s+\d+\s*[:=]',     # "Chapter 3:"
                r'\bcorresponding\s+author',   # "Corresponding author"
            ]
            
            import re
            for pattern in metadata_patterns:
                if re.search(pattern, text_lower):
                    return True
            
            # Check for question/option that's clearly about metadata (not medical)
            # Only flag if it's asking about metadata explicitly
            if re.search(r'(what|which|who)\s+(is|was|are)\s+(the\s+)?(author|publication|journal|publisher)', text_lower):
                return True
            if re.search(r'(author|publication|journal|publisher)\s+(name|date|year|title)', text_lower):
                return True
            
            return False
        
        # Convert to MCQQuestion objects
        questions = []
        for mcq_data in mcqs_data[:request.num_questions]:
            try:
                # Check if question contains metadata
                question_text = mcq_data.get("question", "")
                if contains_metadata(question_text):
                    print(f"WARNING: Question {mcq_data.get('id', 'unknown')} contains metadata keywords: {question_text[:100]}")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Failed to generate MCQs: Question contains document metadata. Please try again."
                    )
                
                # Validate and normalize options
                if "options" not in mcq_data or not isinstance(mcq_data["options"], list):
                    print(f"ERROR: Invalid options format in MCQ {mcq_data.get('id', 'unknown')}")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Failed to generate MCQs: Invalid question format received. Please try again."
                    )
                
                options = []
                for opt_idx, opt in enumerate(mcq_data["options"]):
                    # Ensure option is a dictionary, not a string
                    if isinstance(opt, str):
                        print(f"ERROR: Option {opt_idx} is a string instead of dict: {opt[:100]}")
                        print(f"ERROR: Full options array: {mcq_data['options']}")
                        raise HTTPException(
                            status_code=500,
                            detail=f"Failed to generate MCQs: Invalid option format (option is a string). Please try again."
                        )
                    
                    if not isinstance(opt, dict):
                        print(f"ERROR: Option {opt_idx} is not a dict: type={type(opt)}, value={opt}")
                        print(f"ERROR: Full options array: {mcq_data['options']}")
                        raise HTTPException(
                            status_code=500,
                            detail=f"Failed to generate MCQs: Invalid option format (expected dictionary). Please try again."
                        )
                    
                    # Try to normalize option format - handle variations
                    normalized_opt = {}
                    
                    # Handle different field names
                    if "id" in opt:
                        normalized_opt["id"] = str(opt["id"])
                    elif "option_id" in opt:
                        normalized_opt["id"] = str(opt["option_id"])
                    elif "letter" in opt:
                        normalized_opt["id"] = str(opt["letter"])
                    else:
                        # Try to infer from position
                        normalized_opt["id"] = ["A", "B", "C", "D", "E"][opt_idx] if opt_idx < 5 else str(opt_idx)
                    
                    if "text" in opt:
                        normalized_opt["text"] = str(opt["text"])
                    elif "option_text" in opt:
                        normalized_opt["text"] = str(opt["option_text"])
                    elif "content" in opt:
                        normalized_opt["text"] = str(opt["content"])
                    else:
                        print(f"ERROR: Option {opt_idx} missing text field. Available keys: {list(opt.keys())}")
                        print(f"ERROR: Option value: {opt}")
                        raise HTTPException(
                            status_code=500,
                            detail=f"Failed to generate MCQs: Option missing text content. Please try again."
                        )
                    
                    # Check if option text contains metadata
                    if contains_metadata(normalized_opt["text"]):
                        print(f"WARNING: Option {opt_idx} in question {mcq_data.get('id', 'unknown')} contains metadata: {normalized_opt['text'][:100]}")
                        raise HTTPException(
                            status_code=500,
                            detail=f"Failed to generate MCQs: Option contains document metadata. Please try again."
                        )
                    
                    if "is_correct" in opt:
                        normalized_opt["is_correct"] = bool(opt["is_correct"])
                    elif "correct" in opt:
                        normalized_opt["is_correct"] = bool(opt["correct"])
                    elif "isCorrect" in opt:
                        normalized_opt["is_correct"] = bool(opt["isCorrect"])
                    else:
                        # Default to false if not specified
                        normalized_opt["is_correct"] = False
                        print(f"WARNING: Option {opt_idx} missing is_correct field, defaulting to False")
                    
                    try:
                        options.append(MCQOption(**normalized_opt))
                    except Exception as e:
                        print(f"ERROR: Failed to create MCQOption from: {opt}")
                        print(f"ERROR: Normalized to: {normalized_opt}")
                        print(f"ERROR: Exception: {e}")
                        raise HTTPException(
                            status_code=500,
                            detail=f"Failed to generate MCQs: Error processing question options. Please try again."
                        )
                
                # BULLETPROOF: Completely ignore AI's difficulty and force case difficulty
                if request.difficulty:
                    # Normalize difficulty to ensure proper capitalization (Easy, Moderate, Hard)
                    difficulty_lower = request.difficulty.lower().strip()
                    if difficulty_lower == "easy":
                        question_difficulty = "Easy"
                    elif difficulty_lower == "moderate":
                        question_difficulty = "Moderate"
                    elif difficulty_lower == "hard":
                        question_difficulty = "Hard"
                    else:
                        # If invalid, default to Moderate
                        question_difficulty = "Moderate"
                        print(f"⚠️ Invalid difficulty '{request.difficulty}', defaulting to 'Moderate'")
                    
                    # Always use the normalized case difficulty, completely ignore what AI returned
                    print(f"🔒 Forcing difficulty to case difficulty: '{question_difficulty}' (ignoring AI's '{mcq_data.get('difficulty', 'unknown')}')")
                else:
                    question_difficulty = mcq_data.get("difficulty", "Moderate")
                    # Normalize AI's difficulty too
                    difficulty_lower = question_difficulty.lower().strip()
                    if difficulty_lower == "easy":
                        question_difficulty = "Easy"
                    elif difficulty_lower == "moderate":
                        question_difficulty = "Moderate"
                    elif difficulty_lower == "hard":
                        question_difficulty = "Hard"
                    else:
                        question_difficulty = "Moderate"
                
                question = MCQQuestion(
                    id=mcq_data["id"],
                    question=mcq_data["question"],
                    options=options,
                    explanation=mcq_data["explanation"],
                    difficulty=question_difficulty  # This will ALWAYS be the case difficulty
                )
                questions.append(question)
            except HTTPException:
                # Re-raise HTTPExceptions as-is
                raise
            except Exception as e:
                # Catch any other unexpected errors
                print(f"ERROR: Unexpected error processing MCQ {mcq_data.get('id', 'unknown')}: {e}")
                raise HTTPException(
                    status_code=500,
                    detail=f"Failed to generate MCQs: Error processing question. Please try again."
                )
        
        # Final validation: Ensure ALL questions have the same difficulty
        if request.difficulty:
            # Normalize the expected difficulty
            difficulty_lower = request.difficulty.lower().strip()
            if difficulty_lower == "easy":
                expected_difficulty = "Easy"
            elif difficulty_lower == "moderate":
                expected_difficulty = "Moderate"
            elif difficulty_lower == "hard":
                expected_difficulty = "Hard"
            else:
                expected_difficulty = "Moderate"
            
            for i, question in enumerate(questions):
                if question.difficulty != expected_difficulty:
                    print(f"🚨 CRITICAL: Question {i} has wrong difficulty '{question.difficulty}', forcing to '{expected_difficulty}'")
                    question.difficulty = expected_difficulty
        
        # Log the final difficulties for debugging
        difficulties = [q.difficulty for q in questions]
        unique_difficulties = set(difficulties)
        print(f"✅ Final MCQ difficulties: {difficulties}")
        print(f"✅ Unique difficulties: {unique_difficulties}")
        
        if len(unique_difficulties) > 1:
            print(f"🚨 ERROR: Still have mixed difficulties: {unique_difficulties}")
        else:
            print(f"✅ SUCCESS: All questions have consistent difficulty: {list(unique_difficulties)[0]}")
        
        return MCQResponse(
            questions=questions,
            generated_at=datetime.utcnow()
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating MCQs: {str(e)}")

@app.post("/ai/identify-concepts", response_model=ConceptResponse)
async def identify_concepts(
    request: ConceptRequest,
    current_user: dict = Depends(get_current_user)
):
    """Identify key medical concepts from a document using OpenAI GPT-4 mini"""
    
    print(f"🔵🔵🔵 identify_concepts endpoint called - document_id: {request.document_id}, case_title: '{request.case_title}', num_concepts: {request.num_concepts}")
    
    if not OPENAI_API_KEY:
        raise HTTPException(status_code=500, detail="OpenAI API key not configured")
    
    # Find the document
    try:
        document = db.documents.find_one({
            "_id": ObjectId(request.document_id),
            "uploaded_by": current_user["id"]
        })
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid document ID")
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")
    
    # Check if document has text content
    if not document.get("content"):
        raise HTTPException(status_code=400, detail="Document does not contain readable text content")
    
    # Fetch case details if case_title is provided
    case_details = None
    case_demographics = ""
    case_difficulty = None
    
    print(f"🔍 Concept generation - case_title: '{request.case_title}'")
    
    if request.case_title:
        import re
        title = request.case_title
        case_description = ""
        
        # First, try to extract demographics directly from the case title (most reliable)
        age_match = re.search(r'(\d+)[-\s]year[-\s]old', title, re.IGNORECASE)
        gender_match = re.search(r'\b(girl|boy|female|male|woman|man|child)\b', title, re.IGNORECASE)
        
        age = age_match.group(1) if age_match else None
        gender = gender_match.group(1).lower() if gender_match else None
        
        # Try to find the case in generated_cases collection for full description
        case_doc = None
        try:
            doc_id_str = str(document["_id"])
            # Try multiple query formats
            case_doc = db.generated_cases.find_one({
                "title": request.case_title,
                "document_id": doc_id_str
            })
            if not case_doc:
                # Try without document_id constraint (in case format differs)
                case_doc = db.generated_cases.find_one({
                    "title": request.case_title
                })
            
            if case_doc:
                case_details = case_doc
                case_description = case_doc.get("description", "")
                case_difficulty = case_doc.get("difficulty", "").strip()
                
                # Extract from description if not found in title
                if not age:
                    age_match = re.search(r'(\d+)[-\s]year[-\s]old', case_description, re.IGNORECASE)
                    age = age_match.group(1) if age_match else None
                if not gender:
                    gender_match = re.search(r'\b(girl|boy|female|male|woman|man|child)\b', case_description, re.IGNORECASE)
                    gender = gender_match.group(1).lower() if gender_match else None
        except Exception as e:
            print(f"⚠️ Warning: Could not fetch case details from database: {e}")
            # Continue with title-based extraction
        
        print(f"🔍 Case details - case_difficulty: '{case_difficulty}' (type: {type(case_difficulty)})")
        if not case_difficulty:
            print(f"⚠️ WARNING: case_difficulty is None/empty - will use normal validation (first case is now Moderate)")
        
        # Build demographics string
        if age and gender:
            # Normalize gender
            if gender in ['girl', 'female', 'woman']:
                gender_normalized = 'female'
            elif gender in ['boy', 'male', 'man']:
                gender_normalized = 'male'
            else:
                gender_normalized = gender
            
            case_demographics = f"\n\nCRITICAL PATIENT DEMOGRAPHICS CONSISTENCY REQUIREMENT:\n- The patient in this case is a {age}-year-old {gender_normalized}\n- ALL generated content (patient profile, concepts, MCQs) MUST reflect this EXACT age and gender\n- DO NOT change the age or gender - maintain strict consistency\n- If the case title mentions '{age}-year-old {gender_normalized}', ensure all generated content matches this exactly"
        elif case_description:
            # If we can't extract specific demographics, use the full case description
            case_demographics = f"\n\nCRITICAL CASE CONTEXT:\n- Full Case Description: {case_description[:500]}\n- ALL generated content MUST be consistent with this specific case scenario\n- Maintain consistency with patient demographics, presentation, and clinical details from the case description above"
        elif request.case_title:
            # Last resort: use the case title itself
            case_demographics = f"\n\nCRITICAL CASE CONTEXT:\n- Case Title: {request.case_title}\n- ALL generated content MUST be consistent with this specific case scenario\n- Extract and maintain consistency with patient demographics from the case title"
    
    try:
        # Use OpenAI GPT-4 mini
        if not openai_client:
            raise Exception("OpenAI client not initialized")
        
        # Retry logic for concept generation - force real generation, no fallbacks
        max_retries = 5  # Increased retries to force real generation
        retry_count = 0
        concepts_data = None
        last_error = None
        
        while retry_count <= max_retries:
            try:
                # Build case JSON for the prompt
                import json
                case_json_data = {}
                case_id_str = ""
                case_title_str = request.case_title or "Medical Case"
                case_difficulty_str = case_difficulty or "Moderate"
                
                if case_doc:
                    case_id_str = str(case_doc.get("_id", ""))
                    case_title_str = case_doc.get("title", request.case_title or "Medical Case")
                    case_difficulty_str = case_doc.get("difficulty", case_difficulty or "Moderate")
                    case_json_data = {
                        "case_id": case_id_str,
                        "title": case_title_str,
                        "difficulty": case_difficulty_str,
                        "description": case_doc.get("description", ""),
                        "key_points": case_doc.get("key_points", [])
                    }
                else:
                    case_json_data = {
                        "case_id": "",
                        "title": case_title_str,
                        "difficulty": case_difficulty_str,
                        "description": "",
                        "key_points": []
                    }
                
                # Add demographics if available
                if age and gender:
                    gender_normalized = gender
                    if gender in ['girl', 'female', 'woman']:
                        gender_normalized = 'female'
                    elif gender in ['boy', 'male', 'man']:
                        gender_normalized = 'male'
                    case_json_data["age"] = age
                    case_json_data["gender"] = gender_normalized
                
                case_json = json.dumps(case_json_data, indent=2)
                
                # Extract key_concept from case if available
                key_concept = ""
                if case_doc and case_doc.get("key_points") and len(case_doc.get("key_points", [])) > 0:
                    key_concept = case_doc.get("key_points", [""])[0]
                
                # Include case description in prompt if available for better context
                case_description_text = ""
                if case_doc and case_doc.get("description"):
                    case_description_text = f"\n\nCASE DESCRIPTION:\n{case_doc.get('description')}\n"
                elif case_description:
                    case_description_text = f"\n\nCASE DESCRIPTION:\n{case_description}\n"
                
                # New structured prompt for key concepts
                identify_concepts_prompt = f"""
You are an expert medical educator. Based on the SINGLE, SPECIFIC clinical case provided below, generate a structured breakdown of the essential high‑yield concepts a medical student must master to understand THIS PARTICULAR CASE.

CRITICAL REQUIREMENT - UNIQUENESS:
- The concepts you generate MUST be UNIQUE to this specific case
- Focus on the SPECIFIC clinical scenario, patient demographics, and case details provided
- Do NOT generate generic concepts that could apply to any case
- Each concept must be tied to the SPECIFIC details of this case (age, gender, presentation, etc.)
- The concepts should reflect what makes THIS case unique and educational

==============================

INPUT CASE (JSON)

==============================

{case_json}{case_description_text}

Use ONLY the information from this SPECIFIC case. You may infer medically standard logic only when clearly implied. Generate concepts that are SPECIFIC to this case's unique clinical scenario.

==============================

OUTPUT FORMAT

==============================

Return a SINGLE JSON object:

{{
  "id": "concept_{case_id_str}",
  "case_id": "{case_id_str}",
  "title": "{case_title_str}",
  "difficulty": "{case_difficulty_str}",
  "key_concept": "{key_concept}",
  "key_concept_summary": "A comprehensive explanation (MUST be 150-250 words, no less) explaining the clinical concept behind THIS SPECIFIC case, referencing the case title and unique patient details. Write in a clear, university-student-friendly style.",
  "learning_objectives": [
    "Objective 1 (MUST be 150-250 words, no less) specific to this case's unique scenario, written for university students with clear explanations...",
    "Objective 2 (MUST be 150-250 words, no less) specific to this case's unique scenario, written for university students with clear explanations...",
    "Objective 3 (MUST be 150-250 words, no less) specific to this case's unique scenario, written for university students with clear explanations..."
  ],
  "core_pathophysiology": "A detailed explanation (MUST be 150-250 words, no less) of the mechanism or physiology relevant to THIS SPECIFIC case, tied to the specific vignette details (patient age, gender, presentation). Write in a clear, university-student-friendly style that explains the underlying mechanisms in an accessible way.",
  "clinical_reasoning_steps": [
    "Step 1 (MUST be 150-250 words, no less): Detailed reasoning tied to SPECIFIC case clues from this case, written for university students with clear explanations of the clinical thinking process...",
    "Step 2 (MUST be 150-250 words, no less): Detailed reasoning written for university students with clear explanations...",
    "Step 3 (MUST be 150-250 words, no less): Detailed reasoning written for university students with clear explanations..."
  ],
  "red_flags_and_pitfalls": [
    "Pitfall 1 (MUST be 150-250 words, no less): Detailed explanation written for university students, explaining why this is a common mistake and how to avoid it...",
    "Pitfall 2 (MUST be 150-250 words, no less): Detailed explanation written for university students, explaining why this is a common mistake and how to avoid it..."
  ],
  "differential_diagnosis_framework": [
    "Dx 1 (MUST be 150-250 words, no less): Detailed justification written for university students, explaining why this diagnosis is considered, what supports it, and what distinguishes it from other possibilities...",
    "Dx 2 (MUST be 150-250 words, no less): Detailed justification written for university students, explaining why this diagnosis is considered, what supports it, and what distinguishes it from other possibilities...",
    "Dx 3 (MUST be 150-250 words, no less): Detailed justification written for university students, explaining why this diagnosis is considered, what supports it, and what distinguishes it from other possibilities..."
  ],
  "important_labs_imaging_to_know": [
    "Lab/imaging 1 (MUST be 150-250 words, no less): Detailed explanation written for university students, explaining what this test is, why it matters for this case, what it shows, and its clinical significance...",
    "Lab/imaging 2 (MUST be 150-250 words, no less): Detailed explanation written for university students, explaining what this test is, why it matters for this case, what it shows, and its clinical significance..."
  ],
  "why_this_case_matters": "A comprehensive explanation (MUST be 150-250 words, no less) connecting the case to real-world practice and exam relevance, written in a clear, university-student-friendly style that explains why mastering this case is important for medical education and clinical practice."
}}

RULES:
- No markdown.
- No extra text.
- JSON ONLY.
- ALL content must be SPECIFIC to this case - reference the case title, patient demographics, and specific clinical details.
- Do NOT generate generic concepts that could apply to multiple cases.
- Tie every learning objective, reasoning step, and concept to the SPECIFIC details of this case.
- CRITICAL WORD COUNT REQUIREMENT: Each section (key_concept_summary, each learning_objective, core_pathophysiology, each clinical_reasoning_step, each red_flags_and_pitfalls item, each differential_diagnosis_framework item, each important_labs_imaging_to_know item, and why_this_case_matters) MUST contain a minimum of 150 words and should aim for 150-250 words. Count the words carefully - sections with less than 150 words are NOT acceptable.
- Give it to me like I am a university student - use accessible language while maintaining scientific accuracy.
"""
                
                system_prompt = identify_concepts_prompt
        
                # Generate response from OpenAI
                print(f"🔄 Attempt {retry_count + 1} of {max_retries + 1} to generate concepts...")
                response = openai_client.chat.completions.create(
                    model="gpt-4.1",
                    messages=[
                        {"role": "system", "content": "Medical educator. Generate key concepts for a medical case with clear sections. Return valid JSON only. CRITICAL: Each section MUST be 150-250 words minimum (no less than 150 words). Give it to me like I am a university student - write in a clear, accessible, and educational style appropriate for university-level medical education."},
                        {"role": "user", "content": system_prompt}
                    ],
                    temperature=0.1,  # Very low temperature for maximum accuracy and consistency
                    max_tokens=12000,  # Significantly increased for comprehensive detailed content (150-250 words per section)
                    timeout=240,  # Increased timeout for comprehensive detailed case generation
                    response_format={"type": "json_object"}  # Force JSON object format for better accuracy
                )
                
                # Parse the response
                import json
                print(f"AI Raw Concepts OpenAI response: {response.choices[0].message.content}")
                
                # Try to extract JSON from the response if it's wrapped in markdown
                response_text = response.choices[0].message.content.strip()
                print(f"🔍 Raw response text (first 500 chars): {response_text[:500]}")
                
                # Remove markdown code blocks if present
                if response_text.startswith('```json'):
                    response_text = response_text.replace('```json', '').replace('```', '').strip()
                elif response_text.startswith('```'):
                    response_text = response_text.replace('```', '').strip()
                
                # Try to find JSON object/array in the text
                # Look for first { or [
                start_idx = response_text.find('{')
                if start_idx == -1:
                    start_idx = response_text.find('[')
                
                if start_idx != -1:
                    # Find matching closing brace/bracket
                    brace_count = 0
                    bracket_count = 0
                    end_idx = start_idx
                    for i in range(start_idx, len(response_text)):
                        if response_text[i] == '{':
                            brace_count += 1
                        elif response_text[i] == '}':
                            brace_count -= 1
                        elif response_text[i] == '[':
                            bracket_count += 1
                        elif response_text[i] == ']':
                            bracket_count -= 1
                        
                        if brace_count == 0 and bracket_count == 0:
                            end_idx = i + 1
                            break
                    
                    if end_idx > start_idx:
                        response_text = response_text[start_idx:end_idx]
                
                print(f"🔍 Extracted JSON text: {response_text[:500]}")
                
                try:
                    concepts_data = json.loads(response_text)
                    print(f"✅ Successfully parsed Concepts JSON: {type(concepts_data)}")
                except json.JSONDecodeError as json_err:
                    print(f"❌ JSON parsing error: {json_err}")
                    print(f"❌ Full response text (first 1000 chars): {response_text[:1000]}")
                    raise ValueError(f"Invalid JSON format: {str(json_err)}")
                
                # Ensure we always return a single case as an array
                if not isinstance(concepts_data, list):
                    print(f"📦 Converting single object to array")
                    concepts_data = [concepts_data]
                
                # Validate that we have at least one concept
                if not concepts_data or len(concepts_data) == 0:
                    print(f"❌ ERROR: Empty concepts_data after parsing")
                    raise ValueError("No concepts found in parsed JSON")
                
                # Validate the response - common validation for all cases
                if concepts_data and len(concepts_data) > 0:
                    first_concept = concepts_data[0]
                    description = first_concept.get("description", "")
                    
                    # Check if structured fields are present (new format or old format)
                    has_structured_fields = bool(
                        first_concept.get("objective") or 
                        first_concept.get("patient_profile") or 
                        first_concept.get("history_of_present_illness") or
                        first_concept.get("past_medical_history") or
                        first_concept.get("medications") or
                        first_concept.get("examination") or
                        first_concept.get("initial_investigations") or
                        first_concept.get("case_progression") or
                        first_concept.get("final_diagnosis") or
                        # New format fields
                        first_concept.get("key_concept_summary") or
                        first_concept.get("learning_objectives") or
                        first_concept.get("core_pathophysiology") or
                        first_concept.get("clinical_reasoning_steps") or
                        first_concept.get("red_flags_and_pitfalls") or
                        first_concept.get("differential_diagnosis_framework") or
                        first_concept.get("important_labs_imaging_to_know") or
                        first_concept.get("why_this_case_matters")
                    )
                    
                    # For non-Easy/non-first cases, do normal validation
                    # If no structured fields and no description, that's a problem
                    if not has_structured_fields and not description:
                        print(f"⚠️ WARNING: No structured fields and no description, will retry...")
                        raise ValueError("No content found in response")
                    
                    # Check for placeholder text - only in description if it exists
                    # If structured fields are present, description might be empty (that's OK)
                    if description:  # Only check if description exists
                        placeholder_indicators = [
                            "[Patient demographics",
                            "[Detailed history]",
                            "[Relevant conditions]",
                            "[Current medications]",
                            "[Physical exam findings]",
                            "[Diagnostic tests and results]",
                            "[Case evolution]",
                            "[Diagnosis and learning objective]",
                            "Case breakdown not available"
                        ]
                        if any(indicator in description for indicator in placeholder_indicators):
                            print(f"⚠️ WARNING: Detected placeholder text in response, will retry...")
                            raise ValueError("Placeholder content detected")
                    
                    # Only check for generic descriptions if we don't have structured fields
                    # (structured fields indicate proper case breakdown format)
                    # Skip this check for Easy case to be more lenient
                    # Also check if it's the first case (might not have difficulty set yet)
                    is_easy_case = False
                    if case_difficulty:
                        is_easy_case = case_difficulty.lower().strip() == "easy"
                    # Also check if it's likely the first case (no difficulty set, or first in list)
                    # If case_difficulty is None, treat as Moderate (first case is now Moderate)
                    if not case_difficulty:
                        print(f"⚠️ No case_difficulty found - treating as Moderate (first case is now Moderate)")
                        is_easy_case = False  # First case is now Moderate, use normal validation
                    if not has_structured_fields and description and not is_easy_case:
                        generic_phrases = [
                            "systematic review investigates",
                            "this study examines",
                            "the research focuses on",
                            "this document discusses",
                            "the paper explores",
                            "this article presents",
                            "the study analyzes",
                            "this case explores key clinical concepts",
                            "based on the document content",
                            "this is an important medical concept",
                            "represents a fundamental principle"
                        ]
                        description_lower = description.lower()
                        # If description is mostly generic phrases without specific patient details, reject it
                        generic_count = sum(1 for phrase in generic_phrases if phrase in description_lower)
                        # Check for actual case details (patient, symptoms, diagnosis, etc.)
                        case_detail_indicators = [
                            "patient", "symptom", "diagnosis", "treatment", "examination", 
                            "history", "presenting", "complaint", "physical exam", "vital signs",
                            "laboratory", "imaging", "medication", "dose", "mg", "years old",
                            "presenting complaint", "chief complaint", "physical examination"
                        ]
                        detail_count = sum(1 for indicator in case_detail_indicators if indicator in description_lower)
                        
                        # If too many generic phrases and not enough case details, reject
                        # But be more lenient - require 3+ generic phrases AND < 2 details
                        if generic_count >= 3 and detail_count < 2:
                            print(f"⚠️ WARNING: Description too generic ({generic_count} generic phrases, {detail_count} case details), will retry...")
                            raise ValueError("Description is too generic and lacks specific case details")
                    
                    # Check if we have structured fields with substantial real content
                    # Support both old format (objective, patient_profile, etc.) and new format (key_concept_summary, learning_objectives, etc.)
                    content_length = 0
                    structured_fields_present = False
                    
                    # Check old format structured fields
                    old_format_fields = [
                        "objective", "patient_profile", "history_of_present_illness",
                        "past_medical_history", "medications", "examination",
                        "initial_investigations", "case_progression", "final_diagnosis"
                    ]
                    for field in old_format_fields:
                        if first_concept.get(field):
                            content_length += len(str(first_concept.get(field, "")))
                            structured_fields_present = True
                    
                    # Check new format structured fields
                    new_format_fields = [
                        "key_concept_summary", "core_pathophysiology", "why_this_case_matters"
                    ]
                    for field in new_format_fields:
                        if first_concept.get(field):
                            content_length += len(str(first_concept.get(field, "")))
                            structured_fields_present = True
                    
                    # Check array fields in new format
                    array_fields = [
                        "learning_objectives", "clinical_reasoning_steps",
                        "red_flags_and_pitfalls", "differential_diagnosis_framework",
                        "important_labs_imaging_to_know"
                    ]
                    for field in array_fields:
                        if first_concept.get(field) and isinstance(first_concept.get(field), list):
                            content_length += sum(len(str(item)) for item in first_concept.get(field, []))
                            if len(first_concept.get(field, [])) > 0:
                                structured_fields_present = True
                    
                    # Check description if no structured fields or as additional content
                    if description:
                        content_length += len(description)
                    
                    # Very lenient validation - just check we have some content
                    # For Easy cases, be even more lenient - accept if we have title or any field
                    has_title = bool(first_concept.get("title", "").strip())
                    
                    # Accept if we have any content at all - be very lenient
                    # No minimum content requirement - accept any content
                    if content_length > 0 or has_title or has_structured_fields:
                        print(f"✅ Validated content: {content_length} characters (structured fields: {structured_fields_present}, Easy case: {is_easy_case})")
                        print(f"✅ Final concepts_data length: {len(concepts_data)}")
                        print(f"✅ Concept title: {first_concept.get('title', 'N/A')}")
                        
                        # DETAILED LOGGING: Print all concept fields for debugging
                        print(f"\n🔍🔍🔍 DETAILED CONCEPT DATA for case_title: '{request.case_title}':")
                        print(f"📊 Total concepts: {len(concepts_data)}")
                        for idx, concept in enumerate(concepts_data):
                            print(f"\n📝 Concept {idx + 1}:")
                            print(f"   - title: {concept.get('title', 'NO TITLE')[:100]}")
                            print(f"   - title length: {len(concept.get('title', ''))}")
                            print(f"   - description: {concept.get('description', 'NO DESCRIPTION')[:200]}")
                            print(f"   - description length: {len(concept.get('description', ''))}")
                            print(f"   - objective: {concept.get('objective', 'NO OBJECTIVE')[:200]}")
                            print(f"   - objective length: {len(concept.get('objective', ''))}")
                            print(f"   - patient_profile: {concept.get('patient_profile', 'NO PATIENT_PROFILE')[:200]}")
                            print(f"   - patient_profile length: {len(concept.get('patient_profile', ''))}")
                            print(f"   - history_of_present_illness: {concept.get('history_of_present_illness', 'NO HISTORY')[:200]}")
                            print(f"   - history length: {len(concept.get('history_of_present_illness', ''))}")
                            print(f"   - examination: {concept.get('examination', 'NO EXAMINATION')[:200]}")
                            print(f"   - examination length: {len(concept.get('examination', ''))}")
                            print(f"   - final_diagnosis: {concept.get('final_diagnosis', 'NO DIAGNOSIS')[:200]}")
                            print(f"   - diagnosis length: {len(concept.get('final_diagnosis', ''))}")
                            print(f"   - case_title: {concept.get('case_title', 'NO CASE_TITLE')}")
                            print(f"   - Full concept keys: {list(concept.keys())}")
                        
                        break  # Success, exit retry loop
                    else:
                        print(f"⚠️ WARNING: No content found (content_length: {content_length}), will retry...")
                        raise ValueError("No content found in response")
                
            except (json.JSONDecodeError, ValueError, Exception) as e:
                last_error = e
                print(f"❌ ERROR on attempt {retry_count + 1}: {e}")
                if retry_count < max_retries:
                    retry_count += 1
                    # Exponential backoff: 1s, 2s, 4s, 8s, 16s
                    delay = min(2 ** retry_count, 16)
                    print(f"🔄 Retrying concept generation... ({retry_count}/{max_retries}) after {delay}s delay")
                    import time
                    time.sleep(delay)
                    continue
                else:
                    # All retries failed - raise error instead of using fallback
                    print(f"❌ All {max_retries + 1} attempts failed. Last error: {last_error}")
                    raise HTTPException(
                        status_code=500, 
                        detail=f"Failed to generate concepts after {max_retries + 1} attempts. Please try again. Error: {str(last_error)}"
                    )
        
        # Safety check - ensure concepts_data exists
        if concepts_data is None:
            print(f"❌ ERROR: concepts_data is None after retry loop")
            raise HTTPException(
                status_code=500,
                detail="Failed to generate concepts. No data was produced after all retries. Please try again."
            )
        
        # Ensure we always return a single case as an array
        if not isinstance(concepts_data, list):
            concepts_data = [concepts_data]
        
        # Validate and convert to Concept objects - no fallback
        if len(concepts_data) == 0:
            print(f"❌ ERROR: No concepts in response after all retries")
            raise HTTPException(
                status_code=500,
                detail="Failed to generate concepts. No valid content was produced. Please try again."
            )
        
        # Convert to Concept objects - take only the first case
        try:
            concept_data = concepts_data[0]
            
            # Ensure required fields exist with defaults
            if not concept_data.get("title"):
                concept_data["title"] = concept_data.get("id", "Medical Concept")
            if not concept_data.get("importance"):
                concept_data["importance"] = "High"
            
            # Handle different formats: new format (key_concept_summary, learning_objectives) or old format (objective, patient_profile)
            # Priority: new format first, then old format, then fallback to description
            if concept_data.get("key_concept_summary") or concept_data.get("learning_objectives"):
                # New format - build description from new structured fields
                description_parts = []
                if concept_data.get("key_concept_summary"):
                    description_parts.append(f"Key Concept Summary: {concept_data.get('key_concept_summary')}")
                if concept_data.get("learning_objectives") and isinstance(concept_data.get("learning_objectives"), list):
                    description_parts.append("\n\nLearning Objectives:")
                    for obj in concept_data.get("learning_objectives", []):
                        description_parts.append(f"\n- {obj}")
                if concept_data.get("core_pathophysiology"):
                    description_parts.append(f"\n\nCore Pathophysiology: {concept_data.get('core_pathophysiology')}")
                if concept_data.get("clinical_reasoning_steps") and isinstance(concept_data.get("clinical_reasoning_steps"), list):
                    description_parts.append("\n\nClinical Reasoning Steps:")
                    for step in concept_data.get("clinical_reasoning_steps", []):
                        description_parts.append(f"\n{step}")
                if concept_data.get("red_flags_and_pitfalls") and isinstance(concept_data.get("red_flags_and_pitfalls"), list):
                    description_parts.append("\n\nRed Flags and Pitfalls:")
                    for pitfall in concept_data.get("red_flags_and_pitfalls", []):
                        description_parts.append(f"\n- {pitfall}")
                if concept_data.get("differential_diagnosis_framework") and isinstance(concept_data.get("differential_diagnosis_framework"), list):
                    description_parts.append("\n\nDifferential Diagnosis Framework:")
                    for dx in concept_data.get("differential_diagnosis_framework", []):
                        description_parts.append(f"\n{dx}")
                if concept_data.get("important_labs_imaging_to_know") and isinstance(concept_data.get("important_labs_imaging_to_know"), list):
                    description_parts.append("\n\nImportant Labs/Imaging to Know:")
                    for lab in concept_data.get("important_labs_imaging_to_know", []):
                        description_parts.append(f"\n- {lab}")
                if concept_data.get("why_this_case_matters"):
                    description_parts.append(f"\n\nWhy This Case Matters: {concept_data.get('why_this_case_matters')}")
                
                concept_data["description"] = "".join(description_parts)
            elif concept_data.get("objective") or concept_data.get("patient_profile"):
                # Old format - build description from separate fields
                description_parts = []
                if concept_data.get("objective"):
                    description_parts.append(f"Objective: {concept_data.get('objective')}")
                if concept_data.get("patient_profile") or concept_data.get("history_of_present_illness"):
                    description_parts.append("\n\nCase Presentation:")
                    if concept_data.get("patient_profile"):
                        description_parts.append(f"\nPatient Profile: {concept_data.get('patient_profile')}")
                    if concept_data.get("history_of_present_illness"):
                        description_parts.append(f"\nHistory of Present Illness: {concept_data.get('history_of_present_illness')}")
                    if concept_data.get("past_medical_history"):
                        description_parts.append(f"\nPast Medical History: {concept_data.get('past_medical_history')}")
                    if concept_data.get("medications"):
                        description_parts.append(f"\nMedications: {concept_data.get('medications')}")
                    if concept_data.get("examination"):
                        description_parts.append(f"\nExamination: {concept_data.get('examination')}")
                if concept_data.get("initial_investigations"):
                    description_parts.append(f"\n\nInitial Investigations: {concept_data.get('initial_investigations')}")
                if concept_data.get("case_progression"):
                    description_parts.append(f"\n\nCase Progression/Intervention: {concept_data.get('case_progression')}")
                if concept_data.get("final_diagnosis"):
                    description_parts.append(f"\n\nFinal Diagnosis/Learning Anchor: {concept_data.get('final_diagnosis')}")
                
                concept_data["description"] = "".join(description_parts)
            elif not concept_data.get("description"):
                # If no description and no structured fields, create a minimal description
                concept_data["description"] = concept_data.get("title", "Medical concept from document")
                print(f"⚠️ WARNING: No description field found, using title as description")
            
            # Ensure case_title is set on the concept for proper storage and retrieval
            if not concept_data.get("case_title") and request.case_title:
                concept_data["case_title"] = request.case_title
                print(f"✅ Set case_title '{request.case_title}' on concept")
            
            concepts = [Concept(**concept_data)]
            print(f"✅ Successfully created Concept object: {concepts[0].title} for case: {concept_data.get('case_title', 'NO CASE_TITLE')}")
        except Exception as e:
            print(f"❌ ERROR: Failed to create Concept object: {e}")
            print(f"❌ Concept data: {concepts_data[0] if concepts_data else 'None'}")
            # No fallback - raise error to force retry
            raise HTTPException(
                status_code=500,
                detail=f"Failed to create concept object from generated data. Please try again. Error: {str(e)}"
            )
        
        return ConceptResponse(
            document_id=request.document_id,
            concepts=concepts,
            generated_at=datetime.utcnow()
        )
        
    except HTTPException:
        # Re-raise HTTP exceptions (these are intentional errors)
        raise
    except Exception as e:
        print(f"❌ CRITICAL ERROR in identify_concepts: {e}")
        import traceback
        print(f"❌ Traceback: {traceback.format_exc()}")
        # No fallback - raise error to force proper generation
        raise HTTPException(
            status_code=500, 
            detail=f"Error identifying concepts: {str(e)}. Please try again."
        )

@app.post("/ai/auto-generate", response_model=AutoGenerationResponse)
async def auto_generate_content(
    request: AutoGenerationRequest,
    current_user: dict = Depends(get_current_user),
    request_obj: Request = None
):
    """Automatically generate all content types from a document - optimized for speed"""
    
    # Rate limiting check
    if request_obj:
        client_ip = request_obj.client.host
        if not rate_limiter.is_allowed(client_ip):
            raise HTTPException(
                status_code=429, 
                detail="Rate limit exceeded. Please wait before making another request."
            )
    
    if not OPENAI_API_KEY:
        raise HTTPException(status_code=500, detail="OpenAI API key not configured")
    
    print(f" Auto-generation started for document: {request.document_id}")
    
    # Find the document
    try:
        document = db.documents.find_one({
            "_id": ObjectId(request.document_id),
            "uploaded_by": current_user["id"]
        })
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid document ID")
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")
    
    # Check if document has text content
    if not document.get("content"):
        raise HTTPException(status_code=400, detail="Document does not contain readable text content")
    
    response_data = {
        "document_id": request.document_id,
        "generated_at": datetime.utcnow(),
        "success": True,
        "message": "Auto-generation completed successfully"
    }
    
    try:
        # Use OpenAI GPT-4 mini
        if not openai_client:
            raise Exception("OpenAI client not initialized")
        
        # Generate Cases
        if request.generate_cases:
            print(f" Generating {request.num_cases} cases...")
            try:
                # Determine difficulty distribution
                num_cases = request.num_cases
                if num_cases == 5:
                    difficulty_distribution = "EXACTLY 1 Moderate case (first), 2 Moderate cases (second and third), and 2 Hard cases (fourth and fifth) in this EXACT order"
                    difficulty_requirements = """
CRITICAL DIFFICULTY REQUIREMENTS (for 5 cases - MUST FOLLOW THIS EXACT ORDER):
- Case 1 (FIRST): MUST be "Moderate" - Moderate complexity, some diagnostic challenges, requires intermediate clinical reasoning. Use moderately complex concepts from the document.
- Case 2 (SECOND): MUST be "Moderate" - Moderate complexity, some diagnostic challenges, requires intermediate clinical reasoning. Use moderately complex concepts from the document.
- Case 3 (THIRD): MUST be "Moderate" - Moderate complexity, some diagnostic challenges, requires intermediate clinical reasoning. Use different moderately complex concepts from the document.
- Case 4 (FOURTH): MUST be "Hard" - Complex presentation, multiple differential diagnoses, requires advanced clinical reasoning and integration of multiple concepts. Use complex, nuanced concepts from the document.
- Case 5 (FIFTH): MUST be "Hard" - Complex presentation, multiple differential diagnoses, requires advanced clinical reasoning and integration of multiple concepts. Use different complex, nuanced concepts from the document.

The case description MUST match the difficulty level:
- Easy: Clear, straightforward presentation with obvious clinical signs. Simple diagnostic path. Basic pathophysiology. Directly based on document content.
- Moderate: Some complexity in presentation, requires connecting multiple findings. Moderate diagnostic challenge. Based on document content but requires some reasoning.
- Hard: Complex, nuanced presentation with subtle findings. Multiple possible diagnoses. Requires advanced reasoning and knowledge integration. Based on complex concepts from document."""
                else:
                    difficulty_distribution = f"Distribute difficulty levels appropriately across {num_cases} cases"
                    difficulty_requirements = "- Assign difficulty levels (Easy, Moderate, Hard) that match the complexity of each case"
                
                cases_prompt = f"""You are an expert medical case scenario generator specializing in creating comprehensive, educational medical cases for medical students. Based on the following document content, generate {num_cases} realistic, detailed medical case scenarios that are DIRECTLY RELEVANT to the document content.

Document Content:
{document['content'][:3000]}

CRITICAL RELEVANCE REQUIREMENTS:
- Cases MUST be directly based on the medical concepts, conditions, and information in the document above
- Use specific medical terminology, conditions, and details from the document
- Do NOT generate generic cases - every case must reflect actual content from the document
- Extract key medical concepts, diseases, symptoms, treatments, or procedures mentioned in the document
- Ensure each case scenario incorporates specific information from the document content

IMPORTANT INSTRUCTIONS:
- Create realistic, clinically relevant cases that are SPECIFICALLY RELATED to the document content
- Include detailed patient presentations with specific symptoms, vital signs, and history
- Provide comprehensive case descriptions (minimum 2-3 paragraphs, 200-300 words each)
- Include specific learning objectives and clinical reasoning points
- Make cases challenging but educational
- Include relevant diagnostic considerations and treatment approaches

{difficulty_requirements}

For each scenario, provide:
1. A clear, descriptive title that indicates the main condition
2. A detailed, comprehensive case description that MATCHES the difficulty level:
   - Easy: Straightforward presentation, clear symptoms, obvious diagnosis path
   - Moderate: Some complexity, requires connecting findings, moderate diagnostic challenge
   - Hard: Complex presentation, subtle findings, multiple differentials, advanced reasoning required
   Include:
   - Patient demographics and presenting complaint
   - Detailed history of present illness
   - Relevant past medical history
   - Physical examination findings
   - Initial diagnostic considerations
3. 5-7 specific key learning points covering:
   - Pathophysiology
   - Diagnostic criteria
   - Differential diagnosis
   - Treatment options
   - Clinical pearls
4. Difficulty level that MATCHES the case complexity: "Easy", "Moderate", or "Hard"

Format your response as a JSON array with the following structure (MUST be in this exact order for 5 cases):
[
    {{
        "title": "Specific Case Title (e.g., 'Acute Myocardial Infarction in a 55-year-old Male')",
        "description": "Comprehensive case description (200-300 words) that MATCHES the difficulty level and is DIRECTLY RELEVANT to the document. Easy cases should be straightforward, Hard cases should be complex with subtle findings.",
        "key_points": ["Specific learning point 1", "Specific learning point 2", "Specific learning point 3", "Specific learning point 4", "Specific learning point 5"],
        "difficulty": "Moderate" (for first, second, and third cases), or "Hard" (for fourth and fifth cases)
    }}
]

CRITICAL FOR 5 CASES: 
- Generate exactly 5 unique cases in this EXACT order: [Easy, Moderate, Moderate, Hard, Hard]
- Each case description must be comprehensive (200-300 words minimum)
- The difficulty level MUST match the complexity of the case description
- ALL cases MUST be directly relevant to the document content above
- Use specific medical concepts, conditions, and terminology from the document"""
                
                system_message = "You are an expert medical case scenario generator. Always respond with valid JSON format. Each case description must be comprehensive (200-300 words minimum)."
                if request.num_cases == 5:
                    system_message += " CRITICAL: When generating 5 cases, you MUST create exactly 1 Easy case, 2 Moderate cases, and 2 Hard cases. The case descriptions must match their difficulty levels."
                
                cases_response = openai_client.chat.completions.create(
                    model="gpt-4.1",
                    messages=[
                        {"role": "system", "content": system_message},
                        {"role": "user", "content": cases_prompt}
                    ],
                    temperature=0.7,
                    max_tokens=6000
                )
                import json
                print(f"Raw cases response: {cases_response.choices[0].message.content[:500]}...")
                
                try:
                    # Clean the response text
                    response_text = cases_response.choices[0].message.content.strip()
                    
                    # Remove markdown formatting if present
                    if response_text.startswith('```json'):
                        response_text = response_text.replace('```json', '').replace('```', '').strip()
                    elif response_text.startswith('```'):
                        response_text = response_text.replace('```', '').strip()
                    
                    # Find JSON array boundaries
                    start_idx = response_text.find('[')
                    end_idx = response_text.rfind(']') + 1
                    
                    if start_idx != -1 and end_idx != -1:
                        json_text = response_text[start_idx:end_idx]
                        cases_data = json.loads(json_text)
                        
                        # Ensure we have the right number of cases
                        if len(cases_data) < request.num_cases:
                            print(f"Warning: Only got {len(cases_data)} cases, expected {request.num_cases}")
                            # If we got fewer cases than requested, try to generate more
                            if len(cases_data) == 0:
                                raise ValueError("No cases generated")
                        
                        # Take the requested number of cases, or all available if fewer
                        cases_to_use = cases_data[:request.num_cases]
                        
                        # Validate and enforce difficulty distribution for 5 cases
                        if request.num_cases == 5 and len(cases_to_use) == 5:
                            # Force correct distribution in exact order: Easy, Moderate, Moderate, Hard, Hard
                            required_distribution = ["Easy", "Moderate", "Moderate", "Hard", "Hard"]
                            
                            # Check current distribution
                            difficulty_counts = {"Easy": 0, "Moderate": 0, "Hard": 0}
                            for case in cases_to_use:
                                diff = case.get("difficulty", "Moderate").strip()
                                # Normalize difficulty
                                if diff.lower() == "easy":
                                    difficulty_counts["Easy"] += 1
                                elif diff.lower() == "moderate":
                                    difficulty_counts["Moderate"] += 1
                                elif diff.lower() == "hard":
                                    difficulty_counts["Hard"] += 1
                                else:
                                    difficulty_counts["Moderate"] += 1  # Default to Moderate
                            
                            print(f"📊 Difficulty distribution before fix: Easy={difficulty_counts['Easy']}, Moderate={difficulty_counts['Moderate']}, Hard={difficulty_counts['Hard']}")
                            
                            # Always enforce correct distribution in order (even if already correct, ensure order is right)
                            if difficulty_counts["Easy"] != 1 or difficulty_counts["Moderate"] != 2 or difficulty_counts["Hard"] != 2:
                                print(f"⚠️ WARNING: Difficulty distribution incorrect. Expected: 1 Easy, 2 Moderate, 2 Hard. Got: {difficulty_counts}")
                            
                            print(f"🔧 Enforcing correct distribution in order: [Easy, Moderate, Moderate, Hard, Hard]")
                            
                            # Force correct distribution in exact order (always, to ensure consistency)
                            for i, case in enumerate(cases_to_use):
                                case["difficulty"] = required_distribution[i]
                                print(f"  Case {i+1}: Set difficulty to {required_distribution[i]}")
                            
                            # Verify final distribution
                            final_counts = {"Easy": 0, "Moderate": 0, "Hard": 0}
                            for case in cases_to_use:
                                final_counts[case.get("difficulty", "Moderate")] += 1
                            print(f"✅ Final difficulty distribution: Easy={final_counts['Easy']}, Moderate={final_counts['Moderate']}, Hard={final_counts['Hard']}")
                        
                        response_data["cases"] = [CaseScenario(**case) for case in cases_to_use]
                        print(f" Generated {len(response_data['cases'])} cases successfully")
                    else:
                        raise ValueError("No valid JSON array found in response")
                        
                except json.JSONDecodeError as e:
                    print(f" Failed to parse cases JSON: {e}")
                    print(f" Response text: {cases_response.choices[0].message.content[:500]}...")
                    # Create fallback cases based on document content
                    fallback_cases = []
                    doc_content_preview = document.get('content', '')[:500] if document.get('content') else ''
                    for i in range(min(request.num_cases, 5)):  # Limit fallback to 5 cases
                        fallback_cases.append({
                            "title": f"Medical Case {i+1} from {document.get('filename', 'Document')}",
                            "description": f"""Based on the uploaded document '{document.get('filename', 'Document')}', this case presents a comprehensive medical scenario involving the key concepts discussed in the document. 

The case includes a detailed patient presentation with specific demographic information, presenting complaint, and relevant medical history. The patient's clinical presentation is thoroughly described, including vital signs, physical examination findings, and relevant diagnostic test results that align with the document content.

The case explores important diagnostic considerations, differential diagnosis approaches, and treatment strategies relevant to the medical concepts presented in the document. This scenario is designed to help medical students understand the clinical application of the theoretical knowledge discussed in the source material.""",
                            "key_points": [
                                "Key medical concept from document",
                                "Diagnostic approach based on document",
                                "Treatment considerations from document",
                                "Clinical reasoning points",
                                "Important learning objective",
                                "Pathophysiology considerations",
                                "Differential diagnosis approach"
                            ],
                            "difficulty": "Moderate"
                        })
                    response_data["cases"] = [CaseScenario(**case) for case in fallback_cases]
                    print(f" Generated {len(response_data['cases'])} fallback cases")
            except Exception as e:
                print(f" Case generation failed: {e}")
                # Create detailed fallback
                doc_filename = document.get('filename', 'Document')
                fallback_cases = [{
                    "title": f"Medical Case from {doc_filename}",
                    "description": f"""This medical case scenario is based on the uploaded document '{doc_filename}' and presents a comprehensive clinical scenario designed for medical education. 

The case includes a detailed patient presentation with comprehensive demographic information, presenting complaint, and relevant social history that provides important context for understanding the clinical situation. The patient's medical history is thoroughly documented, including past medical conditions, previous surgeries, family history, and any relevant genetic or environmental factors that may influence the current presentation.

Physical examination findings are described in detail, including vital signs, general appearance, and system-by-system examination results with both positive and negative findings that are crucial for differential diagnosis. Diagnostic test results are provided with specific values, reference ranges, and clinical interpretation to help students understand how laboratory and imaging studies contribute to the diagnostic process.""",
                    "key_points": [
                        "Document-based learning point 1",
                        "Document-based learning point 2", 
                        "Document-based learning point 3",
                        "Clinical reasoning and diagnostic approach",
                        "Treatment considerations and management strategies"
                    ],
                    "difficulty": "Moderate"
                }]
                response_data["cases"] = [CaseScenario(**case) for case in fallback_cases]  
        
        # Generate MCQs
        if request.generate_mcqs:
            print(f" Generating {request.num_mcqs} MCQs...")
            try:
                mcq_prompt = f"""Generate {request.num_mcqs} MCQ questions from this content:

{document['content'][:500]}

Return JSON array:
[{{"id": "mcq_1", "question": "Medical question?", "options": [{{"id": "A", "text": "Option A", "is_correct": false}}, {{"id": "B", "text": "Correct answer", "is_correct": true}}, {{"id": "C", "text": "Option C", "is_correct": false}}, {{"id": "D", "text": "Option D", "is_correct": false}}], "explanation": "Brief explanation", "difficulty": "Easy|Moderate|Hard"}}]

Generate exactly {request.num_mcqs} questions."""
                
                mcq_response = openai_client.chat.completions.create(
                    model="gpt-4.1",
                    messages=[
                        {"role": "system", "content": "You are an expert medical educator creating MCQ questions. Always respond with valid JSON format."},
                        {"role": "user", "content": mcq_prompt}
                    ],
                    temperature=0.7,
                    max_tokens=3000
                )
                print(f"Raw MCQ response: {mcq_response.choices[0].message.content[:500]}...")
                
                try:
                    # Clean the response text
                    response_text = mcq_response.choices[0].message.content.strip()
                    
                    # Remove markdown formatting if present
                    if response_text.startswith('```json'):
                        response_text = response_text.replace('```json', '').replace('```', '').strip()
                    elif response_text.startswith('```'):
                        response_text = response_text.replace('```', '').strip()
                    
                    # Find JSON array boundaries
                    start_idx = response_text.find('[')
                    end_idx = response_text.rfind(']') + 1
                    
                    if start_idx != -1 and end_idx != -1:
                        json_text = response_text[start_idx:end_idx]
                        mcq_data = json.loads(json_text)
                        
                        # Ensure we have the right number of MCQs
                        if len(mcq_data) < request.num_mcqs:
                            print(f"Warning: Only got {len(mcq_data)} MCQs, expected {request.num_mcqs}")
                            # If we got fewer MCQs than requested, try to generate more
                            if len(mcq_data) == 0:
                                raise ValueError("No MCQs generated")
                        
                        # Take the requested number of MCQs, or all available if fewer
                        mcqs_to_use = mcq_data[:request.num_mcqs]
                        questions = []
                        for mcq in mcqs_to_use:
                            options = [MCQOption(**option) for option in mcq["options"]]
                            question = MCQQuestion(
                                id=mcq["id"],
                                question=mcq["question"],
                                options=options,
                                explanation=mcq["explanation"],
                                difficulty=mcq["difficulty"]
                            )
                            questions.append(question)
                        response_data["mcqs"] = questions
                        print(f" Generated {len(response_data['mcqs'])} MCQs successfully")
                    else:
                        raise ValueError("No valid JSON array found in response")
                        
                except json.JSONDecodeError as e:
                    print(f" Failed to parse MCQs JSON: {e}")
                    print(f" Response text: {mcq_response.choices[0].message.content[:500]}...")
                    # Create fallback MCQs based on document content
                    fallback_mcqs = []
                    for i in range(min(request.num_mcqs, 5)):  # Limit fallback to 5 MCQs
                        fallback_mcqs.append({
                            "id": f"mcq_{i+1}",
                            "question": f"Based on the document '{document.get('filename', 'Document')}', which of the following is most accurate?",
                            "options": [
                                {"id": "A", "text": "Option A based on document content", "is_correct": False},
                                {"id": "B", "text": "Option B based on document content", "is_correct": True},
                                {"id": "C", "text": "Option C based on document content", "is_correct": False},
                                {"id": "D", "text": "Option D based on document content", "is_correct": False},
                                {"id": "E", "text": "Option E based on document content", "is_correct": False}
                            ],
                            "explanation": f"This question is based on the key concepts discussed in the uploaded document '{document.get('filename', 'Document')}'.",
                            "difficulty": "Moderate"
                        })
                    
                    questions = []
                    for mcq in fallback_mcqs:
                        options = [MCQOption(**option) for option in mcq["options"]]
                        question = MCQQuestion(
                            id=mcq["id"],
                            question=mcq["question"],
                            options=options,
                            explanation=mcq["explanation"],
                            difficulty=mcq["difficulty"]
                        )
                        questions.append(question)
                    response_data["mcqs"] = questions
                    print(f" Generated {len(response_data['mcqs'])} fallback MCQs")
            except Exception as e:
                print(f" MCQ generation failed: {e}")
                # Create minimal fallback
                fallback_mcq = {
                    "id": "mcq_1",
                    "question": f"Based on the document '{document.get('filename', 'Document')}', which statement is correct?",
                    "options": [
                        {"id": "A", "text": "Option A", "is_correct": False},
                        {"id": "B", "text": "Option B", "is_correct": True},
                        {"id": "C", "text": "Option C", "is_correct": False},
                        {"id": "D", "text": "Option D", "is_correct": False},
                        {"id": "E", "text": "Option E", "is_correct": False}
                    ],
                    "explanation": f"This question is based on the document content.",
                    "difficulty": "Moderate"
                }
                options = [MCQOption(**option) for option in fallback_mcq["options"]]
                question = MCQQuestion(
                    id=fallback_mcq["id"],
                    question=fallback_mcq["question"],
                    options=options,
                    explanation=fallback_mcq["explanation"],
                    difficulty=fallback_mcq["difficulty"]
                )
                response_data["mcqs"] = [question]
        
        # Generate Concepts
        if request.generate_concepts:
            print(f" Generating {request.num_concepts} concepts...")
            # Retry logic for concept generation - force real generation, no fallbacks
            max_retries = 5
            retry_count = 0
            concepts_data = None
            last_error = None
            
            while retry_count <= max_retries:
                try:
                    concepts_prompt = f"""Identify {request.num_concepts} key medical concepts from this content. The concepts MUST be DIRECTLY RELEVANT to the document content below.

Document Content:
{document['content'][:3000]}

CRITICAL REQUIREMENTS - 100% ACCURACY MANDATORY:
- ACCURACY IS PARAMOUNT: All medical information, facts, terminology, and clinical details MUST be 100% accurate and directly derived from the document content
- Concepts MUST be based on specific medical information from the document above
- Use ONLY specific medical terminology, conditions, and details that are explicitly stated in the document
- Do NOT generate generic concepts - every concept must reflect actual content from the document
- Do NOT invent, assume, or extrapolate information not present in the document
- Each concept description must be detailed and specific (minimum 100 characters)
- Verify every medical fact against the document before including it

Return JSON array:
[{{"id": "concept_1", "title": "Specific Medical Concept from Document", "description": "Detailed concept description based on document content (minimum 100 characters)", "importance": "High|Medium|Low"}}]

Identify exactly {request.num_concepts} concepts that are directly relevant to the document content."""
                    
                    concepts_response = openai_client.chat.completions.create(
                        model="gpt-4.1",
                        messages=[
                            {"role": "system", "content": "You are an expert medical educator identifying key concepts. Always respond with valid JSON array format. Concepts MUST be directly relevant to the document content provided. CRITICAL: Ensure 100% accuracy - all medical information must be factually correct and directly derived from the document. Do not generate generic or placeholder content."},
                            {"role": "user", "content": concepts_prompt}
                        ],
                        temperature=0.4,  # Very low temperature for maximum accuracy and consistency
                        max_tokens=3000
                    )
                    print(f"🔄 Attempt {retry_count + 1} of {max_retries + 1} to generate concepts...")
                    print(f"Raw concepts response: {concepts_response.choices[0].message.content[:500]}...")
                    
                    # Clean the response text
                    response_text = concepts_response.choices[0].message.content.strip()
                    
                    # Remove markdown formatting if present
                    if response_text.startswith('```json'):
                        response_text = response_text.replace('```json', '').replace('```', '').strip()
                    elif response_text.startswith('```'):
                        response_text = response_text.replace('```', '').strip()
                    
                    # Find JSON array boundaries
                    start_idx = response_text.find('[')
                    end_idx = response_text.rfind(']') + 1
                    
                    if start_idx != -1 and end_idx != -1:
                        json_text = response_text[start_idx:end_idx]
                        concepts_data = json.loads(json_text)
                        
                        # Validate content - check for real content, not placeholders
                        has_valid_content = False
                        for concept in concepts_data:
                            desc = concept.get("description", "")
                            title = concept.get("title", "")
                            # Check for placeholder text
                            if any(placeholder in desc.lower() or placeholder in title.lower() 
                                   for placeholder in ["[text]", "[patient", "[medical", "placeholder", "example"]):
                                raise ValueError("Generated content contains placeholders")
                            # Check minimum length
                            if len(desc) >= 100 and len(title) >= 10:
                                has_valid_content = True
                        
                        if not has_valid_content:
                            raise ValueError("Generated concepts lack sufficient detail")
                        
                        # Ensure we have the right number of concepts
                        if len(concepts_data) < request.num_concepts:
                            print(f"Warning: Only got {len(concepts_data)} concepts, expected {request.num_concepts}")
                            if len(concepts_data) == 0:
                                raise ValueError("No concepts generated")
                        
                        # Take the requested number of concepts, or all available if fewer
                        concepts_to_use = concepts_data[:request.num_concepts]
                        response_data["concepts"] = [Concept(**concept) for concept in concepts_to_use]
                        print(f"✅ Generated {len(response_data['concepts'])} concepts successfully")
                        break  # Success, exit retry loop
                    else:
                        raise ValueError("No valid JSON array found in response")
                        
                except (json.JSONDecodeError, ValueError, Exception) as e:
                    last_error = e
                    print(f"❌ ERROR on attempt {retry_count + 1}: {e}")
                    if retry_count < max_retries:
                        retry_count += 1
                        # Exponential backoff: 1s, 2s, 4s, 8s, 16s
                        delay = min(2 ** retry_count, 16)
                        print(f"🔄 Retrying concept generation... ({retry_count}/{max_retries}) after {delay}s delay")
                        import time
                        time.sleep(delay)
                        continue
                    else:
                        # All retries failed - raise error instead of using fallback
                        print(f"❌ All {max_retries + 1} attempts failed. Last error: {last_error}")
                        raise HTTPException(
                            status_code=500, 
                            detail=f"Failed to generate concepts after {max_retries + 1} attempts. Please try again. Error: {str(last_error)}"
                        )
            
            # If we get here without concepts_data, something went wrong
            if not concepts_data:
                raise HTTPException(
                    status_code=500,
                    detail="Failed to generate concepts. No valid content was produced. Please try again."
                )
        
        # Generate Case Titles
        if request.generate_titles:
            print(f" Generating {request.num_titles} case titles...")
            try:
                titles_prompt = f"""Generate {request.num_titles} case titles from this document:

{document['content'][:500]}

Return JSON array only:
[
    {{
        "id": "case_1",
        "title": "Case Title",
        "description": "Brief case description.",
        "difficulty": "Moderate"
    }}
]"""
                
                titles_response = openai_client.chat.completions.create(
                    model="gpt-4.1",
                    messages=[
                        {"role": "system", "content": "You are an expert medical case generator. Always respond with valid JSON format."},
                        {"role": "user", "content": titles_prompt}
                    ],
                    temperature=0.7,
                    max_tokens=2000
                )
                print(f"Raw titles response: {titles_response.choices[0].message.content[:500]}...")
                
                try:
                    # Clean the response text
                    response_text = titles_response.choices[0].message.content.strip()
                    
                    # Remove markdown formatting if present
                    if response_text.startswith('```json'):
                        response_text = response_text.replace('```json', '').replace('```', '').strip()
                    elif response_text.startswith('```'):
                        response_text = response_text.replace('```', '').strip()
                    
                    # Find JSON array boundaries
                    start_idx = response_text.find('[')
                    end_idx = response_text.rfind(']') + 1
                    
                    if start_idx != -1 and end_idx != -1:
                        json_text = response_text[start_idx:end_idx]
                        titles_data = json.loads(json_text)
                        
                        # Ensure we have the right number of titles
                        if len(titles_data) < request.num_titles:
                            print(f"Warning: Only got {len(titles_data)} titles, expected {request.num_titles}")
                        
                        response_data["titles"] = [CaseTitle(**title) for title in titles_data[:request.num_titles]]
                        print(f" Generated {len(response_data['titles'])} case titles successfully")
                    else:
                        raise ValueError("No valid JSON array found in response")
                        
                except json.JSONDecodeError as e:
                    print(f" Failed to parse titles JSON: {e}")
                    print(f" Response text: {titles_response.choices[0].message.content[:500]}...")
                    # Create fallback titles based on document content
                    fallback_titles = []
                    for i in range(min(request.num_titles, 5)):  # Limit fallback to 5 titles
                        fallback_titles.append({
                            "id": f"case_{i+1}",
                            "title": f"Medical Case {i+1} from {document.get('filename', 'Document')}",
                            "description": f"This case is based on the medical content discussed in the uploaded document '{document.get('filename', 'Document')}' and presents a realistic clinical scenario for learning.",
                            "difficulty": "Moderate"
                        })
                    response_data["titles"] = [CaseTitle(**title) for title in fallback_titles]
                    print(f" Generated {len(response_data['titles'])} fallback case titles")
            except Exception as e:
                print(f" Title generation failed: {e}")
                # Create minimal fallback
                fallback_title = {
                    "id": "case_1",
                    "title": f"Case from {document.get('filename', 'Document')}",
                    "description": f"Medical case based on the uploaded document content.",
                    "difficulty": "Moderate"
                }
                response_data["titles"] = [CaseTitle(**fallback_title)]
        
        print(f" Auto-generation completed successfully!")
        return AutoGenerationResponse(**response_data)
        
    except Exception as e:
        print(f" Auto-generation error: {str(e)}")
        return AutoGenerationResponse(
            document_id=request.document_id,
            generated_at=datetime.utcnow(),
            success=False,
            message=f"Auto-generation failed: {str(e)}"
        )

@app.post("/ai/quick-generate")
async def quick_generate_content(
    request: AutoGenerationRequest,
    current_user: dict = Depends(get_current_user)
):
    """Quick generation endpoint that returns immediately with basic content"""
    
    if not OPENAI_API_KEY:
        raise HTTPException(status_code=500, detail="OpenAI API key not configured")
    
    print(f" Quick generation started for document: {request.document_id}")
    
    # Find the document
    try:
        document = db.documents.find_one({
            "_id": ObjectId(request.document_id),
            "uploaded_by": current_user["id"]
        })
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid document ID")
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")
    
    # Check if document has text content
    if not document.get("content"):
        raise HTTPException(status_code=400, detail="Document does not contain readable text content")
    
    try:
        # Use OpenAI GPT-4 mini
        if not openai_client:
            raise Exception("OpenAI client not initialized")
        
        # Generate only MCQs first (most important for user experience)
        mcqs = []
        if request.generate_mcqs:
            print(f" Quick generating {request.num_mcqs} MCQs...")
            try:
                mcq_prompt = f"""Generate {request.num_mcqs} MCQ questions from this document:

{document['content'][:500]}

Return JSON array only:
[
    {{
        "id": "mcq_1",
        "question": "Question based on document content?",
        "options": [
            {{"id": "A", "text": "Option A", "is_correct": false}},
            {{"id": "B", "text": "Option B", "is_correct": true}},
            {{"id": "C", "text": "Option C", "is_correct": false}},
            {{"id": "D", "text": "Option D", "is_correct": false}},
            {{"id": "E", "text": "Option E", "is_correct": false}}
        ],
        "explanation": "Brief explanation of the correct answer.",
        "difficulty": "Moderate"
    }}
]"""
                
                mcq_response = openai_client.chat.completions.create(
                    model="gpt-4.1",
                    messages=[
                        {"role": "system", "content": "You are an expert medical educator creating MCQ questions. Always respond with valid JSON format."},
                        {"role": "user", "content": mcq_prompt}
                    ],
                    temperature=0.7,
                    max_tokens=2000
                )
                print(f"Quick MCQ response: {mcq_response.choices[0].message.content[:200]}...")
                
                try:
                    # Clean the response text
                    response_text = mcq_response.choices[0].message.content.strip()
                    
                    # Remove markdown formatting if present
                    if response_text.startswith('```json'):
                        response_text = response_text.replace('```json', '').replace('```', '').strip()
                    elif response_text.startswith('```'):
                        response_text = response_text.replace('```', '').strip()
                    
                    # Find JSON array boundaries
                    start_idx = response_text.find('[')
                    end_idx = response_text.rfind(']') + 1
                    
                    if start_idx != -1 and end_idx != -1:
                        json_text = response_text[start_idx:end_idx]
                        mcq_data = json.loads(json_text)
                        
                        # Convert to MCQQuestion objects
                        for mcq in mcq_data[:request.num_mcqs]:
                            options = [MCQOption(**option) for option in mcq["options"]]
                            question = MCQQuestion(
                                id=mcq["id"],
                                question=mcq["question"],
                                options=options,
                                explanation=mcq["explanation"],
                                difficulty=mcq["difficulty"]
                            )
                            mcqs.append(question)
                        print(f" Quick generated {len(mcqs)} MCQs successfully")
                    else:
                        raise ValueError("No valid JSON array found in response")
                        
                except json.JSONDecodeError as e:
                    print(f" Quick MCQ JSON parsing failed: {e}")
                    # Create fallback MCQs
                    for i in range(min(request.num_mcqs, 3)):
                        fallback_mcq = {
                            "id": f"mcq_{i+1}",
                            "question": f"Based on the document '{document.get('filename', 'Document')}', which statement is correct?",
                            "options": [
                                {"id": "A", "text": "Option A", "is_correct": False},
                                {"id": "B", "text": "Option B", "is_correct": True},
                                {"id": "C", "text": "Option C", "is_correct": False},
                                {"id": "D", "text": "Option D", "is_correct": False},
                                {"id": "E", "text": "Option E", "is_correct": False}
                            ],
                            "explanation": f"This question is based on the document content.",
                            "difficulty": "Moderate"
                        }
                        options = [MCQOption(**option) for option in fallback_mcq["options"]]
                        question = MCQQuestion(
                            id=fallback_mcq["id"],
                            question=fallback_mcq["question"],
                            options=options,
                            explanation=fallback_mcq["explanation"],
                            difficulty=fallback_mcq["difficulty"]
                        )
                        mcqs.append(question)
                    print(f" Quick generated {len(mcqs)} fallback MCQs")
            except Exception as e:
                print(f" Quick MCQ generation failed: {e}")
        
        # Return quick response with MCQs
        return {
            "document_id": request.document_id,
            "mcqs": mcqs,
            "generated_at": datetime.utcnow(),
            "success": True,
            "message": "Quick generation completed - MCQs ready, other content generating in background",
            "status": "partial"
        }
        
    except Exception as e:
        print(f" Quick generation error: {str(e)}")
        return {
            "document_id": request.document_id,
            "generated_at": datetime.utcnow(),
            "success": False,
            "message": f"Quick generation failed: {str(e)}",
            "status": "failed"
        }
    
class ChatRequest(BaseModel):
    """Chat request schema"""
    message: str = Field(..., min_length=1, max_length=1000)
    document_id: Optional[str] = None
    case_title: Optional[str] = None  # For explore cases mode

class ChatResponse(BaseModel):
    """Chat response schema"""
    response: str
    timestamp: datetime

class ChatSession(BaseModel):
    """Chat session schema"""
    id: str
    name: str
    created_at: datetime
    updated_at: datetime
    user_id: str
    document_id: Optional[str] = None
    message_count: int = 0
    document_filename: Optional[str] = None
    document_content_preview: Optional[str] = None
    # Context fields for case/concept specific chats
    case_title: Optional[str] = None
    concept_title: Optional[str] = None
    parent_chat_id: Optional[str] = None  # Reference to main chat

class CreateChatRequest(BaseModel):
    """Create chat request schema"""
    name: Optional[str] = None
    document_id: Optional[str] = None
    # Context fields for case/concept specific chats
    case_title: Optional[str] = None
    concept_title: Optional[str] = None
    parent_chat_id: Optional[str] = None

class ChatMessage(BaseModel):
    """Chat message schema"""
    id: str
    chat_id: str
    message: str
    response: str
    timestamp: datetime
    document_id: Optional[str] = None

class GeneratedCase(BaseModel):
    """Generated case schema for storage"""
    id: str
    chat_id: str
    document_id: str
    title: str
    description: str
    key_points: List[str]
    difficulty: str
    created_at: datetime

class GeneratedMCQ(BaseModel):
    """Generated MCQ schema for storage"""
    id: str
    chat_id: str
    document_id: str
    case_title: Optional[str] = None
    question: str
    options: List[MCQOption]
    explanation: str
    difficulty: str
    created_at: datetime

class GeneratedConcept(BaseModel):
    """Generated concept schema for storage"""
    id: str
    chat_id: str
    document_id: str
    case_title: Optional[str] = None
    title: str
    description: str
    importance: str
    created_at: datetime

@app.post("/ai/chat", response_model=ChatResponse)
async def chat_with_ai(
    request: ChatRequest,
    current_user: dict = Depends(get_current_user)
):
    """Chat with AI about uploaded documents"""
    
    print(f"CHAT ENDPOINT CALLED")
    print(f"User: {current_user.get('email', 'unknown')}")
    print(f"Message: {request.message}")
    print(f"Document ID: {request.document_id}")
    
    if not OPENAI_API_KEY:
        print("ERROR: OPENAI_API_KEY not configured")
        raise HTTPException(status_code=500, detail="OpenAI API key not configured")
    
    # If document_id is provided, get the document context
    document_context = ""
    if request.document_id:
        try:
            document = db.documents.find_one({
                "_id": ObjectId(request.document_id),
                "uploaded_by": current_user["id"]
            })
            if document and document.get("content"):
                document_context = f"\n\nDocument Context:\n{document['content'][:500]}"  # Limit context
        except Exception:
            pass  # Continue without document context if there's an error
    
    try:
        print("AI: Initializing OpenAI client...")
        # Use OpenAI GPT-4 mini
        if not openai_client:
            raise Exception("OpenAI client not initialized")
        
        print("AI: Creating prompt...")
        
        # Check if this is an explore cases request (when case_title is present in request)
        # For explore cases, use conversational format with full case details
        case_title = request.case_title
        case_difficulty = None
        case_key_concept = None
        case_description = None
        
        if case_title:
            # Fetch full case details to make the prompt unique to this specific case
            try:
                doc_id_str = str(document["_id"]) if document else None
                case_doc = None
                if doc_id_str:
                    case_doc = db.generated_cases.find_one({
                        "title": case_title,
                        "document_id": doc_id_str
                    })
                if not case_doc:
                    case_doc = db.generated_cases.find_one({
                        "title": case_title
                    })
                
                if case_doc:
                    case_difficulty = case_doc.get("difficulty", "Moderate")
                    case_key_concept = ""
                    if case_doc.get("key_points") and len(case_doc.get("key_points", [])) > 0:
                        case_key_concept = case_doc.get("key_points", [""])[0]
                    case_description = case_doc.get("description", "")
            except Exception as e:
                print(f"⚠️ Warning: Could not fetch case details: {e}")
                # Continue with just case_title if fetch fails
            
            # Use default values if case details not found
            if not case_difficulty:
                case_difficulty = "Moderate"
            if not case_key_concept:
                case_key_concept = ""
            if not case_description:
                case_description = ""
            
            # Explore Cases Mode: Conversational format with case-specific context
            explore_case_system_prompt = f"""
You are an expert Medical Educator having a natural, back‑and‑forth conversation with a medical student about a clinical case.

==============================

CASE CONTEXT

==============================

Case Title: {case_title}

Difficulty: {case_difficulty}

Key Concept: {case_key_concept}

Case Description: {case_description}

==============================

BEHAVIOR RULES

==============================

Always:

- Directly answer the student's question conversationally.

- Then expand with relevant clinical reasoning.

- Speak warmly and naturally — no lecturing tone.

- Refer to details from THIS case.

- Interpret vague questions generously.

- No bullet points, no numbered lists, no headings.

- Use flowing, natural paragraphs only.

Your goal:

Help the student understand the case as if you're sitting beside them, guiding their reasoning.

Input:

Student Question → {request.message}

{document_context}
"""
            
            system_prompt = explore_case_system_prompt
            system_role = "You are an expert Medical Educator having a natural, back-and-forth conversation with a medical student about a clinical case. Respond in flowing paragraphs without bullet points or structured formatting - just like a natural conversation."
        else:
            # General chat mode: standard medical assistant
            system_prompt = f"""Medical AI assistant. Provide concise, accurate answers.

Question: {request.message}
{document_context}

Instructions:
- Give direct, clear answers (2-3 paragraphs max)
- Use medical terms with brief explanations
- Focus on key points: diagnosis, treatment, clinical significance
- Use bullet points for clarity
- Be precise, avoid unnecessary detail"""
            
            system_role = "Medical AI assistant. Provide concise, accurate answers."
        
        print("AI: Generating response from OpenAI...")
        # Generate response from OpenAI
        response = openai_client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": system_role},
                {"role": "user", "content": system_prompt}
            ],
            temperature=0.5,
            max_tokens=800,
            timeout=60
        )
        
        print("SUCCESS: Response generated successfully")
        return ChatResponse(
            response=response.choices[0].message.content,
            timestamp=datetime.now()
        )
        
    except Exception as e:
        print(f"ERROR: Error in chat endpoint: {str(e)}")
        print(f"ERROR: Error type: {type(e).__name__}")
        import traceback
        print(f"ERROR: Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Error generating AI response: {str(e)}")

def generate_chat_name(document_id: Optional[str] = None) -> str:
    """Generate a dynamic chat name"""
    if document_id:
        try:
            document = db.documents.find_one({"_id": ObjectId(document_id)})
            if document:
                # Use document filename as base for chat name
                filename = document.get("filename", "Document")
                # Remove file extension and clean up the name
                name_without_ext = filename.split('.')[0]
                # Truncate if too long
                if len(name_without_ext) > 30:
                    name_without_ext = name_without_ext[:30] + "..."
                return name_without_ext
        except Exception as e:
            print(f"Error generating chat name from document: {e}")
            pass
    
    # Generate timestamp-based name
    timestamp = datetime.now().strftime("%m/%d %H:%M")
    return f"Chat {timestamp}"

@app.post("/chats", response_model=ChatSession, status_code=status.HTTP_201_CREATED)
async def create_chat(
    request: CreateChatRequest,
    current_user: dict = Depends(get_current_user)
):
    """Create a new chat session"""
    
    try:
        print(f"CREATE CHAT ENDPOINT CALLED")
        print(f"Request: {request}")
        print(f"Document ID: {request.document_id}")
        print(f"Current user: {current_user.get('id', 'No user ID')}")
        
        # Generate chat name if not provided
        chat_name = request.name or generate_chat_name(request.document_id)
        print(f"Generated chat name: {chat_name}")
        
        # Get document information if document_id is provided
        document_filename = None
        document_content_preview = None
        
        if request.document_id:
            try:
                document = db.documents.find_one({
                    "_id": ObjectId(request.document_id),
                    "uploaded_by": current_user["id"]
                })
                if document:
                    document_filename = document.get("filename", "Unknown Document")
                    # Store first 200 characters as preview
                    document_content_preview = document.get("content", "")[:200] + "..." if len(document.get("content", "")) > 200 else document.get("content", "")
            except Exception as e:
                print(f"Error fetching document info: {e}")
        
        # Create chat session
        chat_doc = {
            "name": chat_name,
            "user_id": current_user["id"],
            "document_id": request.document_id,
            "document_filename": document_filename,
            "document_content_preview": document_content_preview,
            "message_count": 0,
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow(),
            # Context fields
            "case_title": request.case_title,
            "concept_title": request.concept_title,
            "parent_chat_id": request.parent_chat_id
        }
        
        print(f"Chat document to insert: {chat_doc}")
        result = db.chats.insert_one(chat_doc)
        print(f"Chat created with ID: {result.inserted_id}")
        
        chat_doc["id"] = str(result.inserted_id)
        del chat_doc["_id"]
        
        print(f"Returning chat: {chat_doc}")
        return chat_doc
        
    except Exception as e:
        print(f"ERROR: Error creating chat: {str(e)}")
        print(f"ERROR: Error type: {type(e).__name__}")
        import traceback
        print(f"ERROR: Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Error creating chat: {str(e)}")

@app.get("/chats", response_model=List[ChatSession])
async def get_user_chats(current_user: dict = Depends(get_current_user)):
    """Get all chat sessions for the current user"""
    
    print(f"🔍 GET_CHATS: User ID: {current_user.get('id', 'No ID')}")
    print(f"🔍 GET_CHATS: User email: {current_user.get('email', 'No email')}")
    
    chats = list(db.chats.find(
        {"user_id": current_user["id"]}
    ).sort("updated_at", -1))
    
    print(f"🔍 GET_CHATS: Found {len(chats)} chats in database")
    
    for chat in chats:
        chat["id"] = str(chat["_id"])
        del chat["_id"]
        print(f"🔍 GET_CHATS: Chat {chat['id']} - {chat.get('name', 'No name')}")
    
    print(f"🔍 GET_CHATS: Returning {len(chats)} chats")
    return chats

@app.get("/chats/{chat_id}", response_model=ChatSession)
async def get_chat(
    chat_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get a specific chat session"""
    
    try:
        chat = db.chats.find_one({
            "_id": ObjectId(chat_id),
            "user_id": current_user["id"]
        })
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid chat ID")
    
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")
    
    chat["id"] = str(chat["_id"])
    del chat["_id"]
    
    return chat

@app.delete("/chats/{chat_id}")
async def delete_chat(
    chat_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete a chat session and all its messages"""
    
    try:
        # Verify chat belongs to user
        try:
            chat = db.chats.find_one({
                "_id": ObjectId(chat_id),
                "user_id": current_user["id"]
            })
        except Exception as e:
            print(f"❌ Error finding chat: {e}")
            raise HTTPException(status_code=400, detail="Invalid chat ID")
        
        if not chat:
            raise HTTPException(status_code=404, detail="Chat not found")
        
        # Delete chat and all its associated data
        try:
            db.chats.delete_one({"_id": ObjectId(chat_id)})
            db.chat_messages.delete_many({"chat_id": chat_id})
            
            # Delete all generated content for this chat
            db.generated_cases.delete_many({"chat_id": chat_id})
            db.generated_mcqs.delete_many({"chat_id": chat_id})
            db.generated_concepts.delete_many({"chat_id": chat_id})
            
            print(f"✅ Deleted chat {chat_id} and all associated content")
            return {"message": "Chat deleted successfully"}
        except Exception as e:
            print(f"❌ Error deleting chat data: {e}")
            raise HTTPException(status_code=500, detail=f"Error deleting chat: {str(e)}")
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Unexpected error in delete_chat: {e}")
        raise HTTPException(status_code=500, detail=f"Error deleting chat: {str(e)}")

@app.post("/chats/{chat_id}/messages", response_model=ChatMessage, status_code=status.HTTP_201_CREATED)
async def send_chat_message(
    chat_id: str,
    request: ChatRequest,
    current_user: dict = Depends(get_current_user),
    request_obj: Request = None
):
    """Send a message in a specific chat session"""
    
    # Rate limiting check
    if request_obj:
        client_ip = request_obj.client.host
        if not rate_limiter.is_allowed(client_ip):
            raise HTTPException(
                status_code=429, 
                detail="Rate limit exceeded. Please wait before making another request."
            )
    
    # Verify chat belongs to user
    try:
        chat = db.chats.find_one({
            "_id": ObjectId(chat_id),
            "user_id": current_user["id"]
        })
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid chat ID")
    
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")
    
        # Get AI response using existing chat endpoint logic
        if not OPENAI_API_KEY:
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")
    
    # If document_id is provided, get the document context
    document_context = ""
    document_id = request.document_id or chat.get("document_id")
    if document_id:
        try:
            document = db.documents.find_one({
                "_id": ObjectId(document_id),
                "uploaded_by": current_user["id"]
            })
            if document and document.get("content"):
                document_context = f"\n\nDocument Context:\n{document['content'][:500]}"
        except Exception:
            pass
    
    # Check if this is an explore cases request (when case_title is present in request or chat)
    case_title = request.case_title or chat.get("case_title")
    case_difficulty = None
    case_key_concept = None
    case_description = None
    
    # Fetch full case details if case_title is provided
    if case_title:
        try:
            doc_id_str = document_id if document_id else None
            case_doc = None
            if doc_id_str:
                case_doc = db.generated_cases.find_one({
                    "title": case_title,
                    "document_id": doc_id_str
                })
            if not case_doc:
                case_doc = db.generated_cases.find_one({
                    "title": case_title
                })
            
            if case_doc:
                case_difficulty = case_doc.get("difficulty", "Moderate")
                case_key_concept = ""
                if case_doc.get("key_points") and len(case_doc.get("key_points", [])) > 0:
                    case_key_concept = case_doc.get("key_points", [""])[0]
                case_description = case_doc.get("description", "")
        except Exception as e:
            print(f"⚠️ Warning: Could not fetch case details: {e}")
            # Continue with just case_title if fetch fails
        
        # Use default values if case details not found
        if not case_difficulty:
            case_difficulty = "Moderate"
        if not case_key_concept:
            case_key_concept = ""
        if not case_description:
            case_description = ""
    
    try:
        # Use OpenAI GPT-4 mini
        if not openai_client:
            raise Exception("OpenAI client not initialized")
        
        if case_title:
            # Explore Cases Mode: Conversational format with case-specific context
            explore_case_system_prompt = f"""
You are an expert Medical Educator having a natural, back‑and‑forth conversation with a medical student about a clinical case.

==============================

CASE CONTEXT

==============================

Case Title: {case_title}

Difficulty: {case_difficulty}

Key Concept: {case_key_concept}

Case Description: {case_description}

==============================

BEHAVIOR RULES

==============================

Always:

- Directly answer the student's question conversationally.

- Then expand with relevant clinical reasoning.

- Speak warmly and naturally — no lecturing tone.

- Refer to details from THIS case.

- Interpret vague questions generously.

- No bullet points, no numbered lists, no headings.

- Use flowing, natural paragraphs only.

Your goal:

Help the student understand the case as if you're sitting beside them, guiding their reasoning.

Input:

Student Question → {request.message}

{document_context}
"""
            
            system_prompt = explore_case_system_prompt
            system_role = "You are an expert Medical Educator having a natural, back-and-forth conversation with a medical student about a clinical case. Respond in flowing paragraphs without bullet points or structured formatting - just like a natural conversation."
        else:
            # General chat mode: standard medical assistant
            system_prompt = f"""Medical AI assistant. Provide concise, accurate answers.

Question: {request.message}
{document_context}

Instructions:
- Give direct, clear answers (2-3 paragraphs max)
- Use medical terms with brief explanations
- Focus on key points: diagnosis, treatment, clinical significance
- Use bullet points for clarity
- Be precise, avoid unnecessary detail"""
            
            system_role = "Medical AI assistant. Provide concise, accurate answers."
        
        # Generate response from OpenAI
        response = openai_client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": system_role},
                {"role": "user", "content": system_prompt}
            ],
            temperature=0.5,
            max_tokens=800,
            timeout=60
        )
        ai_response = response.choices[0].message.content
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating AI response: {str(e)}")
    
    # Save message to database
    message_doc = {
        "chat_id": chat_id,
        "message": request.message,
        "response": ai_response,
        "document_id": document_id,
        "timestamp": datetime.now()
    }
    
    result = db.chat_messages.insert_one(message_doc)
    
    # Update chat message count and timestamp
    db.chats.update_one(
        {"_id": ObjectId(chat_id)},
        {
            "$inc": {"message_count": 1},
            "$set": {"updated_at": datetime.now()}
        }
    )
    
    message_doc["id"] = str(result.inserted_id)
    del message_doc["_id"]
    
    return message_doc

@app.put("/chats/{chat_id}/document", response_model=ChatSession)
async def update_chat_document(
    chat_id: str,
    document_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Update a chat with document information"""
    
    try:
        # Verify chat belongs to user
        chat = db.chats.find_one({
            "_id": ObjectId(chat_id),
            "user_id": current_user["id"]
        })
        
        if not chat:
            raise HTTPException(status_code=404, detail="Chat not found")
        
        # Get document information
        document = db.documents.find_one({
            "_id": ObjectId(document_id),
            "uploaded_by": current_user["id"]
        })
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Update chat with document information
        document_filename = document.get("filename", "Unknown Document")
        chat_name = re.sub(r'\.[^/.]+$', '', document_filename) if document_filename else "Chat"
        
        update_data = {
            "document_id": document_id,
            "document_filename": document_filename,
            "document_content_preview": document.get("content", "")[:200] + "..." if len(document.get("content", "")) > 200 else document.get("content", ""),
            "name": chat_name,  # Update chat name to match document filename
            "updated_at": datetime.now()
        }
        
        result = db.chats.update_one(
            {"_id": ObjectId(chat_id)},
            {"$set": update_data}
        )
        
        if result.modified_count == 0:
            raise HTTPException(status_code=400, detail="Failed to update chat")
        
        # Return updated chat
        updated_chat = db.chats.find_one({"_id": ObjectId(chat_id)})
        updated_chat["id"] = str(updated_chat["_id"])
        del updated_chat["_id"]
        
        print(f"✅ Chat {chat_id} updated with document {document_id}")
        return updated_chat
        
    except Exception as e:
        print(f"❌ Error updating chat with document: {e}")
        raise HTTPException(status_code=500, detail=f"Error updating chat: {str(e)}")

@app.get("/chats/{chat_id}/messages", response_model=List[ChatMessage])
async def get_chat_messages(
    chat_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get all messages for a specific chat session"""
    
    # Verify chat belongs to user
    try:
        chat = db.chats.find_one({
            "_id": ObjectId(chat_id),
            "user_id": current_user["id"]
        })
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid chat ID")
    
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")
    
    messages = list(db.chat_messages.find(
        {"chat_id": chat_id}
    ).sort("timestamp", 1))
    
    for message in messages:
        message["id"] = str(message["_id"])
        del message["_id"]
    
    return messages

@app.post("/chats/{chat_id}/content/save")
async def save_generated_content(
    chat_id: str,
    content: dict,
    current_user: dict = Depends(get_current_user)
):
    """Save generated content (cases, MCQs, concepts) for a chat"""
    
    try:
        # Verify chat belongs to user
        chat = db.chats.find_one({
            "_id": ObjectId(chat_id),
            "user_id": current_user["id"]
        })
        
        if not chat:
            raise HTTPException(status_code=404, detail="Chat not found")
        
        document_id = chat.get("document_id")
        if not document_id:
            raise HTTPException(status_code=400, detail="No document associated with this chat")
        
        # Save cases
        if content.get("cases"):
            for case in content["cases"]:
                case_doc = {
                    "chat_id": chat_id,
                    "document_id": document_id,
                    "title": case["title"],
                    "description": case["description"],
                    "key_points": case["key_points"],
                    "difficulty": case["difficulty"],
                    "created_at": datetime.now()
                }
                db.generated_cases.insert_one(case_doc)
        
        # Save MCQs
        if content.get("mcqs"):
            for mcq in content["mcqs"]:
                mcq_doc = {
                    "chat_id": chat_id,
                    "document_id": document_id,
                    "case_title": content.get("case_title"),
                    "question": mcq["question"],
                    "options": mcq["options"],
                    "explanation": mcq["explanation"],
                    "difficulty": mcq["difficulty"],
                    "created_at": datetime.now()
                }
                db.generated_mcqs.insert_one(mcq_doc)
        
        # Save concepts - delete old concepts for this case_title first to prevent duplicates
        if content.get("concepts"):
            case_title = content.get("case_title")
            if case_title:
                # Delete existing concepts for this case_title to prevent duplicates
                delete_result = db.generated_concepts.delete_many({
                    "chat_id": chat_id,
                    "case_title": case_title
                })
                print(f"🗑️ Deleted {delete_result.deleted_count} old concepts for case_title: '{case_title}'")
            
            for concept in content["concepts"]:
                concept_doc = {
                    "chat_id": chat_id,
                    "document_id": document_id,
                    "case_title": content.get("case_title"),  # Ensure case_title is saved
                    "title": concept.get("title", ""),
                    "description": concept.get("description", ""),
                    "importance": concept.get("importance", "medium"),
                    # Save old format structured fields if present
                    "objective": concept.get("objective"),
                    "patient_profile": concept.get("patient_profile"),
                    "history_of_present_illness": concept.get("history_of_present_illness"),
                    "past_medical_history": concept.get("past_medical_history"),
                    "medications": concept.get("medications"),
                    "examination": concept.get("examination"),
                    "initial_investigations": concept.get("initial_investigations"),
                    "case_progression": concept.get("case_progression"),
                    "final_diagnosis": concept.get("final_diagnosis"),
                    # Save new format structured fields if present
                    "case_id": concept.get("case_id"),
                    "key_concept": concept.get("key_concept"),
                    "key_concept_summary": concept.get("key_concept_summary"),
                    "learning_objectives": concept.get("learning_objectives"),
                    "core_pathophysiology": concept.get("core_pathophysiology"),
                    "clinical_reasoning_steps": concept.get("clinical_reasoning_steps"),
                    "red_flags_and_pitfalls": concept.get("red_flags_and_pitfalls"),
                    "differential_diagnosis_framework": concept.get("differential_diagnosis_framework"),
                    "important_labs_imaging_to_know": concept.get("important_labs_imaging_to_know"),
                    "why_this_case_matters": concept.get("why_this_case_matters"),
                    "created_at": datetime.now()
                }
                db.generated_concepts.insert_one(concept_doc)
            print(f"✅ Saved {len(content['concepts'])} concepts for case_title: '{case_title}'")
        
        print(f"✅ Saved generated content for chat {chat_id}")
        return {"message": "Content saved successfully"}
        
    except Exception as e:
        print(f"❌ Error saving generated content: {e}")
        raise HTTPException(status_code=500, detail=f"Error saving content: {str(e)}")

@app.get("/chats/{chat_id}/content")
async def get_generated_content(
    chat_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get all generated content for a chat"""
    
    try:
        # Verify chat belongs to user
        chat = db.chats.find_one({
            "_id": ObjectId(chat_id),
            "user_id": current_user["id"]
        })
        
        if not chat:
            raise HTTPException(status_code=404, detail="Chat not found")
        
        # Get cases
        cases = list(db.generated_cases.find({"chat_id": chat_id}))
        for case in cases:
            case["id"] = str(case["_id"])
            del case["_id"]
        
        # Get MCQs
        mcqs = list(db.generated_mcqs.find({"chat_id": chat_id}))
        for mcq in mcqs:
            mcq["id"] = str(mcq["_id"])
            del mcq["_id"]
        
        # Get concepts
        concepts = list(db.generated_concepts.find({"chat_id": chat_id}))
        for concept in concepts:
            concept["id"] = str(concept["_id"])
            del concept["_id"]
        
        print(f"✅ Retrieved content for chat {chat_id}: {len(cases)} cases, {len(mcqs)} MCQs, {len(concepts)} concepts")
        
        return {
            "cases": cases,
            "mcqs": mcqs,
            "concepts": concepts
        }
        
    except Exception as e:
        print(f"❌ Error retrieving generated content: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving content: {str(e)}")

@app.post("/chats/context", response_model=ChatSession, status_code=status.HTTP_201_CREATED)
async def get_or_create_context_chat(
    request: CreateChatRequest,
    current_user: dict = Depends(get_current_user)
):
    """Get or create a context-specific chat (case/concept) - Medical AI Casewise"""
    
    try:
        print(f"CONTEXT CHAT: Request: {request}")
        print(f"CONTEXT CHAT: User: {current_user.get('id')}")
        
        # Look for existing context chat
        query = {
            "user_id": current_user["id"],
            "document_id": request.document_id
        }
        
        if request.case_title:
            query["case_title"] = request.case_title
        if request.concept_title:
            query["concept_title"] = request.concept_title
        if request.parent_chat_id:
            query["parent_chat_id"] = request.parent_chat_id
        
        existing_chat = db.chats.find_one(query)
        
        if existing_chat:
            print(f"CONTEXT CHAT: Found existing chat: {existing_chat['_id']}")
            existing_chat["id"] = str(existing_chat["_id"])
            del existing_chat["_id"]
            return existing_chat
        
        # Create new context chat
        print(f"CONTEXT CHAT: Creating new context chat")
        
        # Generate context-specific name
        context_name = ""
        if request.concept_title:
            context_name = f"Concept: {request.concept_title}"
        elif request.case_title:
            context_name = f"Case: {request.case_title}"
        else:
            context_name = f"Chat {datetime.now().strftime('%m/%d %H:%M')}"
        
        # Get document info if available
        document_filename = None
        document_content_preview = None
        if request.document_id:
            try:
                document = db.documents.find_one({
                    "_id": ObjectId(request.document_id),
                    "uploaded_by": current_user["id"]
                })
                if document:
                    document_filename = document.get("filename", "Unknown Document")
                    document_content_preview = document.get("content", "")[:200] + "..." if len(document.get("content", "")) > 200 else document.get("content", "")
            except Exception as e:
                print(f"Error fetching document info: {e}")
        
        # Create context chat
        chat_doc = {
            "name": context_name,
            "user_id": current_user["id"],
            "document_id": request.document_id,
            "document_filename": document_filename,
            "document_content_preview": document_content_preview,
            "message_count": 0,
            "created_at": datetime.now(),
            "updated_at": datetime.now(),
            "case_title": request.case_title,
            "concept_title": request.concept_title,
            "parent_chat_id": request.parent_chat_id
        }
        
        result = db.chats.insert_one(chat_doc)
        chat_doc["id"] = str(result.inserted_id)
        del chat_doc["_id"]
        
        print(f"CONTEXT CHAT: Created new chat: {chat_doc['id']}")
        return chat_doc
        
    except Exception as e:
        print(f"❌ Error creating context chat: {e}")
        raise HTTPException(status_code=500, detail=f"Error creating context chat: {str(e)}")


# Dynamic hint generation endpoint
class DynamicHintRequest(BaseModel):
    question_id: str
    question_text: str
    options: List[MCQOption]
    user_attempts: int = 1
    document_context: Optional[str] = None

class DynamicHintResponse(BaseModel):
    hint: str
    generated_at: str

@app.post("/ai/generate-dynamic-hint", response_model=DynamicHintResponse)
async def generate_dynamic_hint(
    request: DynamicHintRequest,
    current_user: dict = Depends(get_current_user)
):
    """
    Generate a dynamic hint for an MCQ question based on user attempts and context.
    """
    try:
        if not openai_client:
            raise HTTPException(status_code=500, detail="OpenAI client not initialized")
        
        # Get document context if available
        document_context = ""
        if request.document_context:
            document_context = f"\n\nDocument Context: {request.document_context[:500]}"
        
        # Create context-aware hint based on attempt number
        attempt_context = ""
        if request.user_attempts == 1:
            attempt_context = "This is the first attempt. Provide a gentle hint that guides toward the correct answer without revealing it."
        elif request.user_attempts == 2:
            attempt_context = "This is the second attempt. Provide a more specific hint that narrows down the options."
        else:
            attempt_context = "This is the third attempt. Provide a detailed hint that strongly guides toward the correct answer."
        
        # Build the prompt for dynamic hint generation
        system_prompt = f"""You are a medical educator. Generate a helpful hint for this MCQ question.

Question: {request.question_text}

Options:
{chr(10).join([f"{opt.id}. {opt.text}" for opt in request.options])}

Context: {attempt_context}{document_context}

Requirements:
- Provide a hint that guides the student toward the correct answer
- Make the hint more specific based on the attempt number
- Do not reveal the correct answer directly
- Keep the hint concise (1-2 sentences)
- Focus on key medical concepts or reasoning

Generate only the hint text, no additional formatting."""

        # Generate response from OpenAI
        response = openai_client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": "You are a medical educator. Provide helpful, educational hints."},
                {"role": "user", "content": system_prompt}
            ],
            temperature=0.7,
            max_tokens=200,
            timeout=30  # 30 seconds timeout for hint generation
        )
        
        hint_text = response.choices[0].message.content.strip()
        
        return DynamicHintResponse(
            hint=hint_text,
            generated_at=datetime.now().isoformat()
        )
        
    except Exception as e:
        print(f"Dynamic hint generation error: {e}")
        # Fallback hint
        fallback_hints = [
            "Consider the key medical concepts mentioned in the question.",
            "Think about the most common clinical presentation or treatment approach.",
            "Review the fundamental principles related to this medical topic."
        ]
        fallback_hint = fallback_hints[min(request.user_attempts - 1, len(fallback_hints) - 1)]
        
        return DynamicHintResponse(
            hint=fallback_hint,
            generated_at=datetime.now().isoformat()
        )


@app.get("/")
def root():
    return {
        "message": "Medical AI - Casewise Backend", 
        "status": "running",
        "database": "Casewise",
        "version": "2.0.0",
        "features": [
            "Document Upload & Storage",
            "Automatic Content Generation",
            "Medical Case Generation",
            "MCQ Question Generation", 
            "Concept Identification",
            "AI Chat with Documents",
            "Chat Session Management"
        ],
        "endpoints": [
            "/signup", 
            "/login", 
            "/auth/google",
            "/me",
            "/documents/upload",
            "/ai/auto-generate",
            "/ai/generate-cases",
            "/ai/generate-case-titles",
            "/ai/generate-mcqs",
            "/ai/identify-concepts",
            "/ai/chat",
            "/chats",
            "/chats/{chat_id}",
            "/chats/{chat_id}/messages"
        ]
    }



if __name__ == "__main__":
    import uvicorn
    # Use PORT environment variable for deployment platforms (Render, Railway, etc.)
    port = int(os.getenv("PORT", 8000))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=True)